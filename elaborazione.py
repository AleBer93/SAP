import random
import os
import time
from collections import Counter

import excel2img
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import psutil
import win32com.client
from docx import Document, shared
from openpyxl import load_workbook
from openpyxl.chart import PieChart, Reference
from openpyxl.chart.label import DataLabelList
from openpyxl.chart.layout import Layout, ManualLayout
from openpyxl.chart.marker import DataPoint
from openpyxl.chart.text import RichText
from openpyxl.drawing.text import (CharacterProperties, Paragraph,
                                   ParagraphProperties)
from openpyxl.styles import (Alignment, Border, Font,  # Per cambiare lo stile
                             PatternFill, Side)
from openpyxl.styles.numbers import (FORMAT_NUMBER_COMMA_SEPARATED1,
                                     FORMAT_PERCENTAGE_00)
from openpyxl.utils import get_column_letter  # Per lavorare sulle colonne
from PIL import ImageGrab

from SAP import Portfolio


class Elaborazione(Portfolio):
    # TODO : crea un context manager per entrare ed uscire dai documenti excel con win32
    """Elabora un portafoglio. Inherits from Portfolio."""

    def __init__(self, intermediario, file_portafoglio=None):
        """
        Initialize the class.

        Parameters:
            intermediario {str} = intermediario per cui fare l'analisi
            file_portafoglio {str} = nome del file da analizzare
        """
        super().__init__(intermediario=intermediario, file_portafoglio=file_portafoglio)
        # Microsoft Excel
        if file_portafoglio is None: # se non viene inserito l'argument file_portafoglio
            ptf = self.file_portafoglio.name
        else:
            ptf = self.file_portafoglio
        self.file_elaborato = ptf[:-5] + '_elaborato.xlsx'
        self.wb = load_workbook(ptf) # apre il file originale
      
    def agglomerato(self):
        """
        Crea un agglomerato del portafoglio diviso per tipo di strumento. Distribuisce gli strumenti in una sola pagina.
        """
        # Dataframe del portfolio
        df = self.df_portfolio
        controvalori = {strumento : df.loc[df['strumento']==strumento, 'controvalore_in_euro'].sum() for strumento in df['strumento'].unique()}
        # Lista degli strumenti possibili nel file di input tradotti in italiano per il commento
        self.dict_str_comm = {'cash' : 'liquidità', 'gov_bond' : 'obbligazioni governative', 'corp_bond' : 'obbligazioni societarie',
            'certificate' : 'certificati', 'equity' : 'azioni', 'etf' : 'etf', 'fund' : 'fondi', 'real_estate' : 'real estate',
            'hedge_fund' : 'fondi hedge', 'private_equity' : 'private equity', 'venture_capital' : 'venture capital', 'private_debt' : 'private debt',
            'insurance' : 'polizze', 'gp' : 'gestioni patrimoniali', 'pip' : 'fondi pensione', 'alternative' : 'altro'}
        # Dizionario che associa ai tipi di strumenti presenti in portafoglio un loro nome in italiano.
        strumenti_dict = {key : value.upper() for key, value in self.dict_str_comm.items()}
        # Lista della numerosità degli strumenti in portafoglio
        c = Counter(list(df.loc[:, 'strumento']))
        # print(c)
        # Header
        ws = self.wb.create_sheet('agglomerato')
        ws = self.wb['agglomerato']
        self.wb.active = ws
        header = ['ISIN', 'Descrizione', 'Quantità', 'Controvalore iniziale', 'Prezzo di carico', 'Divisa', 'Prezzo di mercato in euro', 'Rateo', 'Valore di mercato in euro']
        len_header = len(header)
        for col in ws.iter_cols(min_row=1, max_row=1, min_col=1, max_col=len_header):
            ws[col[0].coordinate].value = header[0]
            del header[0]
            ws[col[0].coordinate].alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            ws[col[0].coordinate].font = Font(name='Century Gothic', size=18, color='FFFFFF', bold=True)
            ws[col[0].coordinate].border = Border(top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), left=Side(border_style='dotted', color='000000'))
            ws[col[0].coordinate].fill = PatternFill(fill_type='solid', fgColor='808080')
            ws.row_dimensions[col[0].row].height = 92.55
        # Body
        min_row = 2
        max_row = 1
        min_col = 1
        max_col = len_header
        strumenti_in_ptf = [strumento for strumento in self.strumenti if c[strumento] > 0]
        for strumento in strumenti_in_ptf:
            max_row = max_row + c[strumento] + 1
            for row in ws.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
                # Etichetta
                if row[0].row == min_row:
                    ws[row[0].coordinate].value = strumenti_dict[strumento]
                    ws[row[0].coordinate].font = Font(name='Century Gothic', size=18, color='808080', bold=True)
                    ws[row[0].coordinate].fill = PatternFill(fill_type='solid', fgColor='F2F2F2')
                    ws[row[0].coordinate].alignment = Alignment(horizontal='center', vertical='center')
                    ws[row[0].coordinate].border = Border(top=Side(border_style='double', color='000000'), left=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
                    ws.merge_cells(start_row=row[0].row, end_row=row[0].row, start_column=min_col, end_column=max_col-1)
                    ws[row[max_col-1].coordinate].value = controvalori[strumento]
                    ws[row[max_col-1].coordinate].font = Font(name='Century Gothic', size=18, color='808080', bold=True)
                    ws[row[max_col-1].coordinate].fill = PatternFill(fill_type='solid', fgColor='F2F2F2')
                    ws[row[max_col-1].coordinate].alignment = Alignment(horizontal='right', vertical='center')
                    ws[row[max_col-1].coordinate].border = Border(top=Side(border_style='double', color='000000'), bottom=Side(border_style='thin', color='000000'), right=Side(border_style='thin', color='000000'))
                    ws[row[max_col-1].coordinate].number_format = '€ #,0.00'
                    ws.row_dimensions[row[0].row].height = 27
                # Strumenti
                else:
                    for _ in range(0, c[strumento]):
                        ws[row[0].offset(row=_, column=len_header-9).coordinate].value = df.loc[df['strumento']==strumento, 'ISIN'].values[_]
                        ws[row[0].offset(row=_, column=len_header-9).coordinate].font = Font(name='Century Gothic', size=18, color='000000')
                        ws[row[0].offset(row=_, column=len_header-9).coordinate].alignment = Alignment(horizontal='left', vertical='center')
                        ws[row[0].offset(row=_, column=len_header-9).coordinate].border = Border(top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), right=Side(border_style='dotted', color='000000')) if strumento not in ['cash', 'insurance', 'gp', 'pip'] else Border(top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), right=Side(border_style='dotted', color='FFFFFF'))
                        ws[row[0].offset(row=_, column=len_header-8).coordinate].value = df.loc[df['strumento']==strumento, 'nome'].values[_]
                        ws[row[0].offset(row=_, column=len_header-8).coordinate].font = Font(name='Century Gothic', size=18, color='000000')
                        ws[row[0].offset(row=_, column=len_header-8).coordinate].alignment = Alignment(horizontal='left', vertical='center')
                        ws[row[0].offset(row=_, column=len_header-8).coordinate].border =Border(top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), left=Side(border_style='dotted', color='000000'), right=Side(border_style='dotted', color='000000')) if strumento not in ['cash', 'insurance', 'gp', 'pip'] else Border(top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), left=Side(border_style='dotted', color='FFFFFF'), right=Side(border_style='dotted', color='FFFFFF'))
                        ws[row[0].offset(row=_, column=len_header-7).coordinate].value = df.loc[df['strumento']==strumento, 'quantità'].values[_]
                        ws[row[0].offset(row=_, column=len_header-7).coordinate].font = Font(name='Century Gothic', size=18, color='000000')
                        ws[row[0].offset(row=_, column=len_header-7).coordinate].alignment = Alignment(horizontal='right', vertical='center')
                        ws[row[0].offset(row=_, column=len_header-7).coordinate].border = Border(top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), left=Side(border_style='dotted', color='000000'), right=Side(border_style='dotted', color='000000')) if strumento not in ['cash', 'insurance', 'gp', 'pip'] else Border(top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), left=Side(border_style='dotted', color='FFFFFF'), right=Side(border_style='dotted', color='FFFFFF'))
                        ws[row[0].offset(row=_, column=len_header-7).coordinate].number_format = '#,0.00'
                        ws[row[0].offset(row=_, column=len_header-6).coordinate].value = df.loc[df['strumento']==strumento, 'controvalore_iniziale_in_euro'].values[_]
                        ws[row[0].offset(row=_, column=len_header-6).coordinate].font = Font(name='Century Gothic', size=18, color='000000')
                        ws[row[0].offset(row=_, column=len_header-6).coordinate].alignment = Alignment(horizontal='right', vertical='center')
                        ws[row[0].offset(row=_, column=len_header-6).coordinate].border = Border(top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), left=Side(border_style='dotted', color='000000'), right=Side(border_style='dotted', color='000000')) if strumento not in ['cash', 'insurance', 'gp', 'pip'] else Border(top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), left=Side(border_style='dotted', color='FFFFFF'), right=Side(border_style='dotted', color='FFFFFF'))
                        ws[row[0].offset(row=_, column=len_header-6).coordinate].number_format = '#,0.00'
                        ws[row[0].offset(row=_, column=len_header-5).coordinate].value = df.loc[df['strumento']==strumento, 'prezzo_di_carico'].values[_]
                        ws[row[0].offset(row=_, column=len_header-5).coordinate].font = Font(name='Century Gothic', size=18, color='000000')
                        ws[row[0].offset(row=_, column=len_header-5).coordinate].alignment = Alignment(horizontal='right', vertical='center')
                        ws[row[0].offset(row=_, column=len_header-5).coordinate].border = Border(top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), left=Side(border_style='dotted', color='000000'), right=Side(border_style='dotted', color='000000')) if strumento not in ['cash', 'insurance', 'gp', 'pip'] else Border(top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), left=Side(border_style='dotted', color='FFFFFF'), right=Side(border_style='dotted', color='FFFFFF'))
                        ws[row[0].offset(row=_, column=len_header-5).coordinate].number_format = '#,0.00'
                        ws[row[0].offset(row=_, column=len_header-4).coordinate].value = df.loc[df['strumento']==strumento, 'divisa'].values[_]
                        ws[row[0].offset(row=_, column=len_header-4).coordinate].font = Font(name='Century Gothic', size=18, color='000000')
                        ws[row[0].offset(row=_, column=len_header-4).coordinate].alignment = Alignment(horizontal='center', vertical='center')
                        ws[row[0].offset(row=_, column=len_header-4).coordinate].border = Border(top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), left=Side(border_style='dotted', color='000000'), right=Side(border_style='dotted', color='000000')) if strumento not in ['cash', 'insurance', 'gp', 'pip'] else Border(top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), left=Side(border_style='dotted', color='FFFFFF'), right=Side(border_style='dotted', color='FFFFFF'))
                        ws[row[0].offset(row=_, column=len_header-3).coordinate].value = df.loc[df['strumento']==strumento, 'prezzo'].values[_]
                        ws[row[0].offset(row=_, column=len_header-3).coordinate].font = Font(name='Century Gothic', size=18, color='000000')
                        ws[row[0].offset(row=_, column=len_header-3).coordinate].alignment = Alignment(horizontal='right', vertical='center')
                        ws[row[0].offset(row=_, column=len_header-3).coordinate].border = Border(top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), left=Side(border_style='dotted', color='000000'), right=Side(border_style='dotted', color='000000')) if strumento not in ['cash', 'insurance', 'gp', 'pip'] else Border(top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), left=Side(border_style='dotted', color='FFFFFF'), right=Side(border_style='dotted', color='FFFFFF'))
                        ws[row[0].offset(row=_, column=len_header-3).coordinate].number_format = '#,0.00'
                        ws[row[0].offset(row=_, column=len_header-2).coordinate].value = df.loc[df['strumento']==strumento, 'rateo'].values[_]
                        ws[row[0].offset(row=_, column=len_header-2).coordinate].font = Font(name='Century Gothic', size=18, color='000000')
                        ws[row[0].offset(row=_, column=len_header-2).coordinate].alignment = Alignment(horizontal='right', vertical='center')
                        ws[row[0].offset(row=_, column=len_header-2).coordinate].border = Border(top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), left=Side(border_style='dotted', color='000000'), right=Side(border_style='dotted', color='000000')) if strumento not in ['cash', 'insurance', 'gp', 'pip'] else Border(top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), left=Side(border_style='dotted', color='FFFFFF'), right=Side(border_style='dotted', color='FFFFFF'))
                        ws[row[0].offset(row=_, column=len_header-2).coordinate].number_format = '#,0.00'
                        ws[row[0].offset(row=_, column=len_header-1).coordinate].value = df.loc[df['strumento']==strumento, 'controvalore_in_euro'].values[_]
                        ws[row[0].offset(row=_, column=len_header-1).coordinate].font = Font(name='Century Gothic', size=18, color='000000')
                        ws[row[0].offset(row=_, column=len_header-1).coordinate].alignment = Alignment(horizontal='right', vertical='center')
                        ws[row[0].offset(row=_, column=len_header-1).coordinate].border = Border(top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), left=Side(border_style='dotted', color='000000'), right=Side(border_style='thin', color='000000')) if strumento not in ['cash', 'insurance', 'gp', 'pip'] else Border(top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), left=Side(border_style='dotted', color='FFFFFF'), right=Side(border_style='thin', color='000000'))
                        ws[row[0].offset(row=_, column=len_header-1).coordinate].number_format = '#,0.00'
                        ws.row_dimensions[row[0].row].height = 23.25
                    min_row = max_row + 1
                    break
        # Footer
        max_row = min_row
        for col in ws.iter_cols(min_row=max_row, max_row=max_row, min_col=1, max_col=len_header):
            if col[0].column == min_col:
                ws[col[0].coordinate].value = 'TOTALE PORTAFOGLIO'
                ws[col[0].coordinate].font = Font(name='Century Gothic', size=18, color='FFFFFF', bold=True)
                ws[col[0].coordinate].fill = PatternFill(fill_type='solid', fgColor='808080')
                ws[col[0].coordinate].alignment = Alignment(horizontal='center', vertical='center')
                ws[col[0].coordinate].border = Border(top=Side(border_style='double', color='000000'), left=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), right=Side(border_style='thin', color='000000'))
            if col[0].column == len_header:
                ws[col[0].coordinate].value = sum(controvalori.values())
                ws[col[0].coordinate].font = Font(name='Century Gothic', size=18, color='FFFFFF', bold=True)
                ws[col[0].coordinate].fill = PatternFill(fill_type='solid', fgColor='808080')
                ws[col[0].coordinate].alignment = Alignment(horizontal='right', vertical='center')
                ws[col[0].coordinate].border = Border(top=Side(border_style='double', color='000000'), left=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'), right=Side(border_style='thin', color='000000'))
                ws[col[0].coordinate].number_format = '€ #,0.00'
                ws.row_dimensions[col[0].row].height = 27
            ws.merge_cells(start_row=col[0].row, end_row=col[0].row, start_column=1, end_column=len_header-1)

    def figure(self):
        """
        Crea le tabelle e le figure delle micro categorie, delle macro categorie, degli strumenti e delle valute.
        """

        SCARTO = 2

        # Creazione foglio figure
        ws_figure = self.wb.create_sheet('figure')
        ws_figure = self.wb['figure']
        self.wb.active = ws_figure

        #---Macro asset class---#
        dict_peso_macro = self.peso_macro()

        # Tabella macro asset class #
        if len(self.fonts_macro) < len(self.macro_asset_class):
            raise Exception(f"Il numero di font delle macro ({len(self.fonts_macro)}) è inferiore al numero delle macro ({len(self.macro_asset_class)}).")
        elif len(self.fonts_macro) > len(self.macro_asset_class):
            raise Exception(f"Il numero di font delle macro ({len(self.fonts_macro)}) è superiore al numero delle macro ({len(self.macro_asset_class)}).")
        # Header
        header_macro = ['MACRO ASSET CLASS', '', 'Peso']
        dim_macro = [3.4, 47, 9.5]
        min_row, max_row = 1, 1
        min_col = 1
        max_col = min_col + len(header_macro) - 1
        for col in ws_figure.iter_cols(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
            ws_figure[col[0].coordinate].value = header_macro[col[0].column-min_col]
            ws_figure[col[0].coordinate].alignment = Alignment(horizontal='center', vertical='center')
            ws_figure[col[0].coordinate].font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
            ws_figure[col[0].coordinate].fill = PatternFill(fill_type='solid', fgColor='595959')
            ws_figure[col[0].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure.column_dimensions[ws_figure[col[0].coordinate].column_letter].width = dim_macro[col[0].column-min_col]
        ws_figure.merge_cells(start_row=min_row, end_row=max_row, start_column=min_col, end_column=min_col+1)
        # Body
        min_row = min_row + 1
        max_row = min_row + len(self.macro_asset_class) - 1
        for row in ws_figure.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
            ws_figure[row[0].coordinate].fill = PatternFill(fill_type='solid', fgColor=self.fonts_macro[row[0].row-min_row])
            ws_figure[row[0].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure[row[1].coordinate].value = self.macro_asset_class[row[0].row-min_row]
            ws_figure[row[1].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure[row[2].coordinate].value = dict_peso_macro[ws_figure[row[1].coordinate].value]
            ws_figure[row[2].coordinate].number_format = '0.0%'
            ws_figure[row[2].coordinate].alignment = Alignment(horizontal='center')
            ws_figure[row[2].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
        # Footer
        max_row = max_row + 1
        ws_figure.cell(max_row, min_col, value='TOTALE')
        ws_figure.cell(max_row, min_col).alignment = Alignment(horizontal='center', vertical='center')
        ws_figure.cell(max_row, min_col).font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
        ws_figure.cell(max_row, min_col).fill = PatternFill(fill_type='solid', fgColor='595959')
        ws_figure.cell(max_row, min_col).border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
        ws_figure.merge_cells(start_row=max_row, end_row=max_row, start_column=min_col, end_column=max_col-1)
        ws_figure.cell(max_row, max_col, value=sum(dict_peso_macro.values()))
        assert sum(ws_figure.cell(i, max_col).value for i in range(min_row, max_row)) == ws_figure.cell(max_row, max_col).value
        ws_figure.cell(max_row, max_col).alignment = Alignment(horizontal='center', vertical='center')
        ws_figure.cell(max_row, max_col).font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
        ws_figure.cell(max_row, max_col).fill = PatternFill(fill_type='solid', fgColor='595959')
        ws_figure.cell(max_row, max_col).border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
        ws_figure.cell(max_row, max_col).number_format = FORMAT_PERCENTAGE_00

        # Grafico macro openpyxl
        chart = PieChart()
        chart.height = 4.77
        chart.width = 6.77
        labels = Reference(ws_figure, min_col=min_col+1, max_col=min_col+1, min_row=min_row, max_row=max_row-1)
        data = Reference(ws_figure, min_col=min_col+2, max_col=min_col+2, min_row=min_row, max_row=max_row-1)
        chart.add_data(data, titles_from_data=False)
        chart.set_categories(labels)
        chart.dataLabels = DataLabelList(dLblPos='bestFit')
        chart.dataLabels.showVal = True
        chart.dataLabels.textProperties = RichText(p=[Paragraph(pPr=ParagraphProperties(defRPr=CharacterProperties(sz=1100, b=True)), endParaRPr=CharacterProperties(sz=1100, b=True))])
        chart.legend = None
        for _ in range(0,4):
            series = chart.series[0]
            pt = DataPoint(idx=_)
            pt.graphicalProperties.solidFill = self.fonts_macro[_]
            series.dPt.append(pt)
        chart.layout = Layout(manualLayout=ManualLayout(x=0.5, y=0.5, h=1, w=1)) # posizione e dimensione figura
        ws_figure.add_chart(chart, get_column_letter(min_col) + str(max_row + SCARTO))
        
        # Grafico macro matplotlib
        plt.subplots(figsize=(4,4))
        try:
            plt.pie([dict_peso_macro[_] for _ in self.macro_asset_class], labels=[str(round((value*100),2)).replace('.',',')+'%' if value > 0.01 else '' for value in dict_peso_macro.values()], radius=1.2, colors=['#'+font for font in self.fonts_macro], pctdistance=0.1, labeldistance=0.5, textprops={'fontsize':14, 'name':'Century Gothic'}, normalize=False)
        except ValueError:
            plt.pie([dict_peso_macro[_] for _ in self.macro_asset_class], labels=[str(round((value*100),2)).replace('.',',')+'%' if value > 0.01 else '' for value in dict_peso_macro.values()], radius=1.2, colors=['#'+font for font in self.fonts_macro], pctdistance=0.1, labeldistance=0.5, textprops={'fontsize':14, 'name':'Century Gothic'}, normalize=True)
        finally:
            plt.savefig('img/macro_pie.png', bbox_inches='tight', pad_inches=0)

        #---Micro asset class---#
        dict_peso_micro = self.peso_micro()
        # Durations
        durations = self.duration()
        
        # Tabella micro asset class #
        if len(self.fonts_micro) < len(self.micro_asset_class):
            raise Exception(f"Il numero di font delle micro ({len(self.fonts_micro)}) è inferiore al numero delle micro ({len(self.micro_asset_class)}).")
        elif len(self.fonts_micro) > len(self.micro_asset_class):
            raise Exception(f"Il numero di font delle micro ({len(self.fonts_micro)}) è superiore al numero delle micro ({len(self.micro_asset_class)}).")
        # Header
        header_micro = ['', 'ASSET CLASS', 'Indice', 'Peso', 'Warning', 'Duration']
        dim_micro = [3.4, 16, 57, 9.5, 9.5, 9.5]
        min_row, max_row = 1, 1
        min_col = max_col + SCARTO
        max_col = min_col + len(header_micro) - 1
        for col in ws_figure.iter_cols(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
            ws_figure[col[0].coordinate].value = header_micro[col[0].column-min_col]
            ws_figure[col[0].coordinate].alignment = Alignment(horizontal='center', vertical='center')
            ws_figure[col[0].coordinate].font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
            ws_figure[col[0].coordinate].fill = PatternFill(fill_type='solid', fgColor='595959')
            ws_figure[col[0].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure.column_dimensions[ws_figure[col[0].coordinate].column_letter].width = dim_micro[col[0].column-min_col]
        # Body
        list_macro = ['Monetario', '', '', 'Obbligazionario', '', '', '', '', '', 'Azionario', '', '', '', 'Commodities']
        indici_micro = ['The BofA Merrill Lynch 0-1 Year Euro Government Index', 'The BofA Merrill Lynch 0-1 Year US Treasury Index', 
            'The BofA Merrill Lynch 0-1 Year G7 Government Index', 'The BofA Merrill Lynch Euro Broad Market Index', 'The BofA Merrill Lynch Euro Large Cap Corporate Index',
            'The BofA Merrill Lynch Euro High Yield Index', 'The BofA Merrill Lynch Global Broad Market Index', 'The BofA Merrill Lynch Global EM Sovereign & Credit Plus Index',
            'The BofA Merrill Lynch Global High Yield Index', 'MSCI Europe', 'MSCI North America', 'MSCI Pacific', 'MSCI Emerging Markets', 'Thomson Reuters/CoreCommodity CRB Commodity Index',
            ]
        min_row = min_row + 1
        max_row = min_row + len(self.micro_asset_class) - 1
        list_peso_micro = list(dict_peso_micro.values())
        for row in ws_figure.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
            ws_figure[row[0].coordinate].fill = PatternFill(fill_type='solid', fgColor=self.fonts_micro[row[0].row-min_row])
            ws_figure[row[0].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure[row[1].coordinate].value = list_macro[row[0].row-min_row]
            ws_figure[row[1].coordinate].alignment = Alignment(horizontal='center', vertical='center')
            ws_figure[row[1].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure[row[2].coordinate].value = indici_micro[row[0].row-min_row]
            ws_figure[row[2].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure[row[3].coordinate].value = tuple(dict_peso_micro.values())[row[0].row-min_row]
            ws_figure[row[3].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure[row[3].coordinate].number_format = '0.0%'
            ws_figure[row[3].coordinate].alignment = Alignment(horizontal='center')
            # warnings
            if ws_figure[row[2].coordinate].value == 'The BofA Merrill Lynch Euro Large Cap Corporate Index':
                if list_peso_micro[4]/dict_peso_macro['Obbligazionario'] > 0.6:
                    ws_figure[row[4].coordinate].value = '!!!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
                elif list_peso_micro[4]/dict_peso_macro['Obbligazionario'] > 0.5:
                    ws_figure[row[4].coordinate].value = '!!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
                elif list_peso_micro[4]/dict_peso_macro['Obbligazionario'] > 0.4:
                    ws_figure[row[4].coordinate].value = '!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
            elif ws_figure[row[2].coordinate].value == 'The BofA Merrill Lynch Euro High Yield Index' or ws_figure[row[2].coordinate].value == 'The BofA Merrill Lynch Global High Yield Index':
                if (list_peso_micro[5]+list_peso_micro[8])/dict_peso_macro['Obbligazionario'] > 0.4:
                    ws_figure[row[4].coordinate].value = '!!!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
                elif (list_peso_micro[5]+list_peso_micro[8])/dict_peso_macro['Obbligazionario'] > 0.3:
                    ws_figure[row[4].coordinate].value = '!!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
                elif (list_peso_micro[5]+list_peso_micro[8])/dict_peso_macro['Obbligazionario'] > 0.2:
                    ws_figure[row[4].coordinate].value = '!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
            elif ws_figure[row[2].coordinate].value == 'The BofA Merrill Lynch Global EM Sovereign & Credit Plus Index':
                if list_peso_micro[7]/dict_peso_macro['Obbligazionario'] > 0.4:
                    ws_figure[row[4].coordinate].value = '!!!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
                elif list_peso_micro[7]/dict_peso_macro['Obbligazionario'] > 0.3:
                    ws_figure[row[4].coordinate].value = '!!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
                elif list_peso_micro[7]/dict_peso_macro['Obbligazionario'] > 0.2:
                    ws_figure[row[4].coordinate].value = '!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
            elif ws_figure[row[2].coordinate].value == 'MSCI Europe':
                if list_peso_micro[9]/dict_peso_macro['Azionario'] > 0.8:
                    ws_figure[row[4].coordinate].value = '!!!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
                elif list_peso_micro[9]/dict_peso_macro['Azionario'] > 0.7:
                    ws_figure[row[4].coordinate].value = '!!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
                elif list_peso_micro[9]/dict_peso_macro['Azionario'] > 0.6:
                    ws_figure[row[4].coordinate].value = '!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
            elif ws_figure[row[2].coordinate].value == 'MSCI North America':
                if list_peso_micro[10]/dict_peso_macro['Azionario'] > 0.8:
                    ws_figure[row[4].coordinate].value = '!!!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
                elif list_peso_micro[10]/dict_peso_macro['Azionario'] > 0.7:
                    ws_figure[row[4].coordinate].value = '!!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
                elif list_peso_micro[10]/dict_peso_macro['Azionario'] > 0.6:
                    ws_figure[row[4].coordinate].value = '!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
            elif ws_figure[row[2].coordinate].value == 'MSCI Pacific':
                if list_peso_micro[11]/dict_peso_macro['Azionario'] > 0.4:
                    ws_figure[row[4].coordinate].value = '!!!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
                elif list_peso_micro[11]/dict_peso_macro['Azionario'] > 0.3:
                    ws_figure[row[4].coordinate].value = '!!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
                elif list_peso_micro[11]/dict_peso_macro['Azionario'] > 0.2:
                    ws_figure[row[4].coordinate].value = '!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
            elif ws_figure[row[2].coordinate].value == 'MSCI Emerging Markets':
                if list_peso_micro[12]/dict_peso_macro['Azionario'] > 0.3:
                    ws_figure[row[4].coordinate].value = '!!!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
                elif list_peso_micro[12]/dict_peso_macro['Azionario'] > 0.2:
                    ws_figure[row[4].coordinate].value = '!!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
                elif list_peso_micro[12]/dict_peso_macro['Azionario'] > 0.1:
                    ws_figure[row[4].coordinate].value = '!C'
                    ws_figure[row[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
            ws_figure[row[4].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure[row[4].coordinate].alignment = Alignment(horizontal='center')
            ws_figure[row[4].coordinate].font = Font(color='000000', bold=True)
            if not self.df_portfolio[(self.df_portfolio['strumento']=='gov_bond') | (self.df_portfolio['strumento']=='corp_bond')].empty:
                if ws_figure[row[2].coordinate].value == 'The BofA Merrill Lynch Euro Broad Market Index':
                    ws_figure[row[5].coordinate].value = round(durations['Obbligazionario Euro Governativo All Maturities'], 2) if durations['Obbligazionario Euro Governativo All Maturities'] > 0.00 else None
                elif ws_figure[row[2].coordinate].value == 'The BofA Merrill Lynch Euro Large Cap Corporate Index':
                    ws_figure[row[5].coordinate].value = round(durations['Obbligazionario Euro Corporate'], 2) if durations['Obbligazionario Euro Corporate'] > 0.00 else None
                elif ws_figure[row[2].coordinate].value == 'The BofA Merrill Lynch Euro High Yield Index':
                    ws_figure[row[5].coordinate].value = round(durations['Obbligazionario Euro High Yield'], 2) if durations['Obbligazionario Euro High Yield'] > 0.00 else None
                elif ws_figure[row[2].coordinate].value == 'The BofA Merrill Lynch Global Broad Market Index':
                    ws_figure[row[5].coordinate].value = round(durations['Obbligazionario Globale Aggregate'], 2) if durations['Obbligazionario Globale Aggregate'] > 0.00 else None
                elif ws_figure[row[2].coordinate].value == 'The BofA Merrill Lynch Global EM Sovereign & Credit Plus Index':
                    ws_figure[row[5].coordinate].value = round(durations['Obbligazionario Paesi Emergenti'], 2) if durations['Obbligazionario Paesi Emergenti'] > 0.00 else None
                elif ws_figure[row[2].coordinate].value == 'The BofA Merrill Lynch Global High Yield Index':
                    ws_figure[row[5].coordinate].value = round(durations['Obbligazionario Globale High Yield'], 2) if durations['Obbligazionario Globale High Yield'] > 0.00 else None
            ws_figure[row[5].coordinate].alignment = Alignment(horizontal='center')
            ws_figure[row[5].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
        # hard coding : deve ripeterlo tante volte quanti sono gli oggetti in self.dict_macro_micro
        start_row = min_row
        end_row = min_row + len(self.dict_macro_micro[ws_figure.cell(row=start_row, column=min_col+1).value]) - 1
        ws_figure.merge_cells(start_row=start_row, end_row=end_row, start_column=min_col+1, end_column=min_col+1)
        start_row = end_row + 1
        end_row = start_row + len(self.dict_macro_micro[ws_figure.cell(row=start_row, column=min_col+1).value]) - 1
        ws_figure.merge_cells(start_row=start_row, end_row=end_row, start_column=min_col+1, end_column=min_col+1)
        start_row = end_row + 1
        end_row = start_row + len(self.dict_macro_micro[ws_figure.cell(row=start_row, column=min_col+1).value]) - 1
        ws_figure.merge_cells(start_row=start_row, end_row=end_row, start_column=min_col+1, end_column=min_col+1)
        start_row = end_row + 1
        end_row = start_row + len(self.dict_macro_micro[ws_figure.cell(row=start_row, column=min_col+1).value]) - 1
        ws_figure.merge_cells(start_row=start_row, end_row=end_row, start_column=min_col+1, end_column=min_col+1)
        # Footer
        max_row = max_row + 1
        for col in ws_figure.iter_rows(min_row=max_row, max_row=max_row, min_col=min_col, max_col=max_col):
            ws_figure[col[0].coordinate].value = 'TOTALE'
            ws_figure[col[0].coordinate].fill = PatternFill(fill_type='solid', fgColor='595959')
            ws_figure[col[0].coordinate].alignment = Alignment(horizontal='center')
            ws_figure[col[0].coordinate].font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
            ws_figure[col[3].coordinate].value = sum(dict_peso_micro.values())
            ws_figure[col[3].coordinate].alignment = Alignment(horizontal='center')
            ws_figure[col[3].coordinate].font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
            ws_figure[col[3].coordinate].fill = PatternFill(fill_type='solid', fgColor='595959')
            ws_figure[col[3].coordinate].number_format = FORMAT_PERCENTAGE_00
            ws_figure[col[4].coordinate].fill = PatternFill(fill_type='solid', fgColor='595959')
            ws_figure[col[5].coordinate].fill = PatternFill(fill_type='solid', fgColor='595959')
        ws_figure.merge_cells(start_row=max_row, end_row=max_row, start_column=min_col, end_column=min_col+2)
        assert sum(ws_figure.cell(i, max_col-2).value for i in range(min_row, max_row)) == ws_figure.cell(max_row, max_col-2).value

        # Grafico micro openpyxl
        chart = PieChart()
        chart.height = 4.77
        chart.width = 6.77
        labels = Reference(ws_figure, min_col=min_col+2, max_col=min_col+2, min_row=min_row, max_row=max_row-1)
        data = Reference(ws_figure, min_col=min_col+3, max_col=min_col+3, min_row=min_row, max_row=max_row-1)
        chart.add_data(data, titles_from_data=False)
        chart.set_categories(labels)
        chart.dataLabels = DataLabelList(dLblPos='bestFit')
        chart.dataLabels.showVal = True
        chart.dataLabels.textProperties = RichText(p=[Paragraph(pPr=ParagraphProperties(defRPr=CharacterProperties(sz=1100, b=True)), endParaRPr=CharacterProperties(sz=1100, b=True))])
        chart.legend = None
        for _ in range(0,14):
            series = chart.series[0]
            pt = DataPoint(idx=_)
            pt.graphicalProperties.solidFill = self.fonts_micro[_]
            series.dPt.append(pt)
        chart.layout = Layout(manualLayout=ManualLayout(x=0.5, y=0.5, h=1, w=1)) # posizione e dimensione figura
        ws_figure.add_chart(chart, get_column_letter(min_col)+str(max_row + SCARTO))
        
        # Grafico micro matplotlib
        plt.subplots(figsize=(4,4))
        try:
            plt.pie([dict_peso_micro[self.micro_asset_class[_]] for _ in range(0, len(self.micro_asset_class))], labels=[str(round((value*100),2)).replace('.',',')+'%' if value > 0.05 else '' for key, value in dict_peso_micro.items()], radius=1.2, colors=['#'+font for font in self.fonts_micro], pctdistance=0.2, labeldistance=0.6, rotatelabels =True, textprops={'fontsize':14, 'name':'Century Gothic', 'rotation_mode':'anchor', 'va':'center', 'ha':'center'}, normalize=False)
        except ValueError:
            plt.pie([dict_peso_micro[self.micro_asset_class[_]] for _ in range(0, len(self.micro_asset_class))], labels=[str(round((value*100),2)).replace('.',',')+'%' if value > 0.05 else '' for key, value in dict_peso_micro.items()], radius=1.2, colors=['#'+font for font in self.fonts_micro], pctdistance=0.2, labeldistance=0.6, rotatelabels =True, textprops={'fontsize':14, 'name':'Century Gothic', 'rotation_mode':'anchor', 'va':'center', 'ha':'center'}, normalize=True)
        finally:
            plt.savefig('img/micro_pie.png', bbox_inches='tight', pad_inches=0)
        # Grafico micro bar matplotlib
        plt.subplots(figsize=(18,5))
        plt.bar(x=[_.replace('Altre Valute', 'Altro').replace('Obbligazionario', 'Obb').replace('Governativo', 'Gov').replace('All Maturities', '').replace('Aggregate', '').replace('North America', 'Nord america').replace('Pacific', 'Pacifico').replace('Emerging Markets', 'Emergenti') for _ in self.micro_asset_class], height=[dict_peso_micro[self.micro_asset_class[_]] for _ in range(0, len(self.micro_asset_class))], width=1, color=['#'+font for font in self.fonts_micro])
        plt.xticks(rotation=25)
        plt.savefig('img/micro_bar.png', bbox_inches='tight', pad_inches=0)

        #---Strumenti---#
        dict_strumenti = self.peso_strumenti()
        df_peso_strumenti = pd.DataFrame.from_dict(dict_strumenti, orient='index', columns=['peso_strumento'])
        # Lista degli strumenti possibili nel file di input tradotti in italiano per la figura
        dict_str_fig = {'cash' : 'Conto corrente', 'gov_bond' : 'Obbligazioni', 'corp_bond' : 'Obbligazioni',
            'certificate' : 'Obbligazioni strutturate / Certificates', 'equity' : 'Azioni', 'etf' : 'ETF/ETC',
            'fund' : 'Fondi comuni/Sicav', 'real_estate' : 'Real Estate', 'hedge_fund' : 'Hedge funds', 'private_equity' : 'Private Equity',
            'venture_capital' : 'Venture Capital', 'private_debt' : 'Private Debt', 'insurance' : 'Polizze', 'gp' : 'Gestioni patrimoniali',
            'pip' : 'Fondi pensione', 'alternative' : 'Altro'}
        df_peso_strumenti.rename(dict_str_fig, inplace=True)
        df_peso_strumenti = df_peso_strumenti.groupby(df_peso_strumenti.index, sort=False).agg({'peso_strumento' : sum})
        series_peso_strumenti = df_peso_strumenti['peso_strumento'].squeeze()
        dict_peso_strumenti = series_peso_strumenti.to_dict()
        
        # Tabella strumenti #
        if len(self.fonts_strumenti) < len(self.strumenti)-1: # obb_gov e obb_corp sono unite in un'unica voce
            raise Exception(f"Il numero di font degli strumenti ({len(self.fonts_strumenti)}) è inferiore al numero degli strumenti ({len(self.strumenti)-1})\nRicorda che obb_gov e obb_corp vengono uniti in un unica voce, quindi serve un colore in meno.")
        elif len(self.fonts_strumenti) > len(self.strumenti)-1:
            raise Exception(f"Il numero di font degli strumenti ({len(self.fonts_strumenti)}) è superiore al numero degli strumenti ({len(self.strumenti)-1})\nRicorda che obb_gov e obb_corp vengono uniti in un unica voce, quindi serve un colore in meno.")
        # Header
        header_strumenti = ['STRUMENTI', '', 'Peso', 'Warning']
        dim_strumenti = [3.4, 47, 9.5, 9.5]
        min_row, max_row = 1, 1
        min_col = max_col + SCARTO
        max_col = min_col + len(header_strumenti) - 1
        for col in ws_figure.iter_cols(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
            ws_figure[col[0].coordinate].value = header_strumenti[col[0].column-min_col]
            ws_figure[col[0].coordinate].alignment = Alignment(horizontal='center', vertical='center')
            ws_figure[col[0].coordinate].font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
            ws_figure[col[0].coordinate].fill = PatternFill(fill_type='solid', fgColor='595959')
            ws_figure[col[0].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure.column_dimensions[ws_figure[col[0].coordinate].column_letter].width = dim_strumenti[col[0].column-min_col]
        ws_figure.merge_cells(start_row=min_row, end_row=max_row, start_column=min_col, end_column=min_col+1)
        # Body
        min_row = min_row + 1
        max_row = min_row + len(dict_peso_strumenti.keys()) - 1
        for row in ws_figure.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
            ws_figure[row[0].coordinate].fill = PatternFill(fill_type='solid', fgColor=self.fonts_strumenti[row[0].row-min_row])
            ws_figure[row[0].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure[row[1].coordinate].value = list(dict_peso_strumenti.keys())[row[0].row-min_row]
            ws_figure[row[1].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure[row[2].coordinate].value = dict_peso_strumenti[ws_figure[row[1].coordinate].value]
            ws_figure[row[2].coordinate].number_format = '0.0%'
            ws_figure[row[2].coordinate].alignment = Alignment(horizontal='center')
            # warnings
            ws_figure[row[2].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            if ws_figure[row[1].coordinate].value == 'Obbligazioni strutturate / Certificates' and dict_peso_strumenti.get('Obbligazioni strutturate / Certificates', 0.00) > 0.10:
                ws_figure[row[3].coordinate].value = '!C'
                ws_figure[row[3].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
            elif ws_figure[row[1].coordinate].value == 'Hedge funds' and dict_peso_strumenti.get('Hedge funds', 0.00) > 0.25:
                ws_figure[row[3].coordinate].value = '!C'
                ws_figure[row[3].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
            ws_figure[row[3].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure[row[3].coordinate].alignment = Alignment(horizontal='center')
            ws_figure[row[3].coordinate].font = Font(color='000000', bold=True)
        # Footer
        max_row = max_row + 1
        ws_figure.cell(max_row, min_col, value='TOTALE')
        ws_figure.cell(max_row, min_col).alignment = Alignment(horizontal='center', vertical='center')
        ws_figure.cell(max_row, min_col).font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
        ws_figure.cell(max_row, min_col).fill = PatternFill(fill_type='solid', fgColor='595959')
        ws_figure.cell(max_row, min_col).border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
        ws_figure.merge_cells(start_row=max_row, end_row=max_row, start_column=min_col, end_column=min_col+1)
        ws_figure.cell(max_row, min_col+2, value=sum(dict_peso_strumenti.values()))
        assert sum(ws_figure.cell(i, min_col+2).value for i in range(min_row, max_row)) == ws_figure.cell(max_row, min_col+2).value
        ws_figure.cell(max_row, min_col+2).alignment = Alignment(horizontal='center', vertical='center')
        ws_figure.cell(max_row, min_col+2).font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
        ws_figure.cell(max_row, min_col+2).fill = PatternFill(fill_type='solid', fgColor='595959')
        ws_figure.cell(max_row, min_col+2).border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
        ws_figure.cell(max_row, min_col+2).number_format = FORMAT_PERCENTAGE_00
        ws_figure.cell(max_row, min_col+3).font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
        ws_figure.cell(max_row, min_col+3).fill = PatternFill(fill_type='solid', fgColor='595959')
        ws_figure.cell(max_row, min_col+3).border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
        
        # Grafico strumenti openpyxl
        chart = PieChart()
        chart.height = 4.77
        chart.width = 6.77
        labels = Reference(ws_figure, min_col=min_col+1, max_col=min_col+1, min_row=min_row, max_row=max_row-1)
        data = Reference(ws_figure, min_col=min_col+2, max_col=min_col+2, min_row=min_row, max_row=max_row-1)
        chart.add_data(data, titles_from_data=False)
        chart.set_categories(labels)
        chart.dataLabels = DataLabelList(dLblPos='bestFit')
        chart.dataLabels.showVal = True
        chart.dataLabels.textProperties = RichText(p=[Paragraph(pPr=ParagraphProperties(defRPr=CharacterProperties(sz=1100, b=True)), endParaRPr=CharacterProperties(sz=1100, b=True))])
        chart.legend = None
        # cambia colori delle fette
        for _ in range(0,11):
            series = chart.series[0]
            pt = DataPoint(idx=_)
            pt.graphicalProperties.solidFill = self.fonts_strumenti[_]
            series.dPt.append(pt)
        # posizione e dimensione figura
        chart.layout = Layout(manualLayout=ManualLayout(x=0.5, y=0.5, h=1, w=1))
        ws_figure.add_chart(chart, get_column_letter(min_col) + str(max_row + SCARTO))

        # Grafico strumenti matplotlib
        plt.subplots(figsize=(4,4))
        try:
            plt.pie([value for value in dict_peso_strumenti.values()], labels=[str(round((value*100),2)).replace('.',',')+'%' if value > 0.03 else '' for value in dict_peso_strumenti.values()], radius=1.2, colors=['#'+font for font in self.fonts_strumenti], pctdistance=0.2, labeldistance=0.6, rotatelabels =True, textprops={'fontsize':14, 'name':'Century Gothic', 'rotation_mode':'anchor', 'va':'center', 'ha':'center'}, normalize=False)
        except ValueError:
            plt.pie([value for value in dict_peso_strumenti.values()], labels=[str(round((value*100),2)).replace('.',',')+'%' if value > 0.03 else '' for value in dict_peso_strumenti.values()], radius=1.2, colors=['#'+font for font in self.fonts_strumenti], pctdistance=0.2, labeldistance=0.6, rotatelabels =True, textprops={'fontsize':14, 'name':'Century Gothic', 'rotation_mode':'anchor', 'va':'center', 'ha':'center'}, normalize=True)
        finally:
            plt.savefig('img/strumenti_pie.png', bbox_inches='tight', pad_inches=0)

        #---Valute---#
        dict_peso_valute = self.peso_valuta()

        # Tabella valute #
        if len(self.fonts_valute) < len(self.valute):
            raise Exception(f"Il numero di font delle valute ({len(self.fonts_valute)}) è inferiore al numero delle valute ({len(self.valute)}).")
        elif len(self.fonts_valute) > len(self.valute):
            raise Exception(f"Il numero di font delle valute ({len(self.fonts_valute)}) è superiore al numero delle valute ({len(self.valute)}).")
        # Header
        header_valute = ['', 'VALUTE', 'Peso', 'Warning']
        dim_valute = [3.4, 9.5, 9.5, 9.5]
        min_row, max_row = 1, 1
        min_col = max_col + SCARTO
        max_col = min_col + len(header_valute) - 1
        for col in ws_figure.iter_cols(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
            ws_figure[col[0].coordinate].value = header_valute[col[0].column-min_col]
            ws_figure[col[0].coordinate].alignment = Alignment(horizontal='center', vertical='center')
            ws_figure[col[0].coordinate].font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
            ws_figure[col[0].coordinate].fill = PatternFill(fill_type='solid', fgColor='595959')
            ws_figure[col[0].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure.column_dimensions[ws_figure[col[0].coordinate].column_letter].width = dim_valute[col[0].column-min_col]
        # Body
        min_row = min_row + 1
        max_row = min_row + len(dict_peso_valute) -1
        for row in ws_figure.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
            ws_figure[row[0].coordinate].fill = PatternFill(fill_type='solid', fgColor=self.fonts_valute[row[0].row-min_row])
            ws_figure[row[0].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure[row[1].coordinate].value = list(dict_peso_valute.keys())[row[0].row-min_row]
            ws_figure[row[1].coordinate].alignment = Alignment(horizontal='center')
            ws_figure[row[1].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure[row[2].coordinate].value = dict_peso_valute[ws_figure[row[1].coordinate].value]
            ws_figure[row[2].coordinate].number_format = '0.0%'
            ws_figure[row[2].coordinate].alignment = Alignment(horizontal='center')
            ws_figure[row[2].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            # Warnings
            if ws_figure[row[1].coordinate].value == 'EUR' and dict_peso_valute.get('EUR', 0.00) < 0.40:
                ws_figure[row[3].coordinate].value = '!!C'
                ws_figure[row[3].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
            elif ws_figure[row[1].coordinate].value == 'EUR' and dict_peso_valute.get('EUR', 0.00) < 0.50:
                ws_figure[row[3].coordinate].value = '!!C'
                ws_figure[row[3].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
            elif ws_figure[row[1].coordinate].value == 'EUR' and dict_peso_valute.get('EUR', 0.00) < 0.60:
                ws_figure[row[3].coordinate].value = '!C'
                ws_figure[row[3].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
            ws_figure[row[3].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure[row[3].coordinate].alignment = Alignment(horizontal='center')
            ws_figure[row[3].coordinate].font = Font(color='000000', bold=True)
        # Footer
        max_row = max_row + 1
        ws_figure.cell(max_row, min_col, value='TOTALE')
        ws_figure.cell(max_row, min_col).alignment = Alignment(horizontal='center', vertical='center')
        ws_figure.cell(max_row, min_col).font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
        ws_figure.cell(max_row, min_col).fill = PatternFill(fill_type='solid', fgColor='595959')
        ws_figure.cell(max_row, min_col).border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
        ws_figure.merge_cells(start_row=max_row, end_row=max_row, start_column=min_col, end_column=min_col+1)
        ws_figure.cell(max_row, min_col+2).value = sum(dict_peso_valute.values())
        assert sum(ws_figure.cell(i, min_col+2).value for i in range(min_row, max_row)) == ws_figure.cell(max_row, min_col+2).value
        ws_figure.cell(max_row, min_col+2).alignment = Alignment(horizontal='center', vertical='center')
        ws_figure.cell(max_row, min_col+2).font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
        ws_figure.cell(max_row, min_col+2).fill = PatternFill(fill_type='solid', fgColor='595959')
        ws_figure.cell(max_row, min_col+2).border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
        ws_figure.cell(max_row, min_col+2).number_format = FORMAT_PERCENTAGE_00
        ws_figure.cell(max_row, min_col+3).font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
        ws_figure.cell(max_row, min_col+3).fill = PatternFill(fill_type='solid', fgColor='595959')
        ws_figure.cell(max_row, min_col+3).border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))

        # Grafico valute openpyxl
        chart = PieChart()
        chart.height = 4.77
        chart.width = 6.77
        labels = Reference(ws_figure, min_col=min_col+1, max_col=min_col+1, min_row=min_row, max_row=max_row-1)
        data = Reference(ws_figure, min_col=min_col+2, max_col=min_col+2, min_row=min_row, max_row=max_row-1)
        chart.add_data(data, titles_from_data=False)
        chart.set_categories(labels)
        chart.dataLabels = DataLabelList(dLblPos='bestFit')
        chart.dataLabels.showVal = True
        chart.dataLabels.textProperties = RichText(p=[Paragraph(pPr=ParagraphProperties(defRPr=CharacterProperties(sz=1100, b=True)), endParaRPr=CharacterProperties(sz=1100, b=True))])
        chart.legend = None
        for _ in range(0,7):
            series = chart.series[0]
            pt = DataPoint(idx=_)
            pt.graphicalProperties.solidFill = self.fonts_valute[_]
            series.dPt.append(pt)
        chart.layout = Layout(manualLayout=ManualLayout(x=0.5, y=0.5, h=1, w=1)) # posizione e dimensione figura
        ws_figure.add_chart(chart, get_column_letter(min_col) + str(max_row + SCARTO))
        
        # Grafico valute matplotlib
        plt.subplots(figsize=(4,4))
        try:
            plt.pie([value for value in dict_peso_valute.values()], labels=[str(round((value*100),2)).replace('.',',')+'%' if value > 0.05 else '' for value in dict_peso_valute.values()], radius=1.2, colors=['#'+font for font in self.fonts_valute], pctdistance=0.2, labeldistance=0.6, rotatelabels =True, textprops={'fontsize':14, 'name':'Century Gothic', 'rotation_mode':'anchor', 'va':'center', 'ha':'center'}, normalize=False)
        except ValueError:
            plt.pie([value for value in dict_peso_valute.values()], labels=[str(round((value*100),2)).replace('.',',')+'%' if value > 0.05 else '' for value in dict_peso_valute.values()], radius=1.2, colors=['#'+font for font in self.fonts_valute], pctdistance=0.2, labeldistance=0.6, rotatelabels =True, textprops={'fontsize':14, 'name':'Century Gothic', 'rotation_mode':'anchor', 'va':'center', 'ha':'center'}, normalize=True)
        finally:
            plt.savefig('img/valute_pie.png', bbox_inches='tight', pad_inches=0)

        #---Risparmio---#
        dict_risparmio = {'amministrato' : 0, 'gestito' : 0}
        for strumento, peso in dict_strumenti.items():
            if strumento in self.amministrato:
                dict_risparmio['amministrato'] += peso
            else:
                dict_risparmio['gestito'] += peso
        fonts_risparmio = ['072FF9', '072FF9', '072FF9', '072FF9',
            'EB1515', 'EB1515', 'EB1515', 'EB1515', 'EB1515', 'EB1515', 'EB1515', 'EB1515', 'EB1515', 'EB1515', 'EB1515']

        # Tabella risparmio #

        # Header
        header_risparmio = ['', 'STRUMENTI', 'Peso']
        dim_risparmio = [3.4, 47, 9.5]
        min_row, max_row = 1, 1
        min_col = max_col + SCARTO
        max_col = min_col + len(header_risparmio) - 1
        for col in ws_figure.iter_cols(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
            ws_figure[col[0].coordinate].value = header_risparmio[col[0].column-min_col]
            ws_figure[col[0].coordinate].alignment = Alignment(horizontal='center', vertical='center')
            ws_figure[col[0].coordinate].font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
            ws_figure[col[0].coordinate].fill = PatternFill(fill_type='solid', fgColor='595959')
            ws_figure[col[0].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure.column_dimensions[ws_figure[col[0].coordinate].column_letter].width = dim_risparmio[col[0].column-min_col]
        # Body
        min_row = min_row + 1
        max_row = min_row + len(dict_peso_strumenti.keys()) - 1
        for row in ws_figure.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
            ws_figure[row[0].coordinate].fill = PatternFill(fill_type='solid', fgColor=fonts_risparmio[row[0].row-min_row])
            ws_figure[row[0].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure[row[1].coordinate].value = list(dict_peso_strumenti.keys())[row[0].row-min_row]
            ws_figure[row[1].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure[row[2].coordinate].value = dict_peso_strumenti[ws_figure[row[1].coordinate].value]
            ws_figure[row[2].coordinate].number_format = '0.0%'
            ws_figure[row[2].coordinate].alignment = Alignment(horizontal='center')
            ws_figure[row[2].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
        # Footer
        max_row = max_row + 1
        ws_figure.cell(max_row, min_col, value='TOTALE')
        ws_figure.cell(max_row, min_col).alignment = Alignment(horizontal='center', vertical='center')
        ws_figure.cell(max_row, min_col).font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
        ws_figure.cell(max_row, min_col).fill = PatternFill(fill_type='solid', fgColor='595959')
        ws_figure.cell(max_row, min_col).border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
        ws_figure.merge_cells(start_row=max_row, end_row=max_row, start_column=min_col, end_column=min_col+1)
        ws_figure.cell(max_row, min_col+2, value=sum(dict_peso_strumenti.values()))
        assert sum(ws_figure.cell(i, min_col+2).value for i in range(min_row, max_row)) == ws_figure.cell(max_row, min_col+2).value
        ws_figure.cell(max_row, min_col+2).alignment = Alignment(horizontal='center', vertical='center')
        ws_figure.cell(max_row, min_col+2).font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
        ws_figure.cell(max_row, min_col+2).fill = PatternFill(fill_type='solid', fgColor='595959')
        ws_figure.cell(max_row, min_col+2).border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
        ws_figure.cell(max_row, min_col+2).number_format = FORMAT_PERCENTAGE_00
        
        # # Grafico risparmio openpyxl
        # chart = PieChart()
        # chart.height = 4.77
        # chart.width = 6.77
        # labels = Reference(ws_figure, min_col=min_col+1, max_col=min_col+1, min_row=min_row, max_row=max_row-1)
        # data = Reference(ws_figure, min_col=min_col+2, max_col=min_col+2, min_row=min_row, max_row=max_row-1)
        # chart.add_data(data, titles_from_data=False)
        # chart.set_categories(labels)
        # chart.dataLabels = DataLabelList(dLblPos='bestFit')
        # chart.dataLabels.showVal = True
        # chart.dataLabels.textProperties = RichText(p=[Paragraph(pPr=ParagraphProperties(defRPr=CharacterProperties(sz=1100, b=True)), endParaRPr=CharacterProperties(sz=1100, b=True))])
        # chart.legend = None
        # # cambia colori delle fette
        # for _ in range(0,11):
        #     series = chart.series[0]
        #     pt = DataPoint(idx=_)
        #     pt.graphicalProperties.solidFill = self.fonts_strumenti[_]
        #     series.dPt.append(pt)
        # # posizione e dimensione figura
        # chart.layout = Layout(manualLayout=ManualLayout(x=0.5, y=0.5, h=1, w=1))
        # ws_figure.add_chart(chart, get_column_letter(min_col) + str(max_row + SCARTO))

        # Grafico risparmio matplotlib
        fig, ax = plt.subplots(figsize=(4,4))
        try:
            wedges, texts = ax.pie([value for value in dict_risparmio.values()], labels=[str(round((value*100),2)).replace('.',',')+'%' if value > 0.03 else '' for value in dict_risparmio.values()], radius=1.2, colors=['#072FF9', '#EB1515'], pctdistance=0.2, labeldistance=0.6, rotatelabels =True, textprops={'fontsize':14, 'name':'Century Gothic', 'rotation_mode':'anchor', 'va':'center', 'ha':'center'}, normalize=False)
            ax.legend(wedges, ['amministrato', 'gestito'], loc="best")
        except ValueError:
            plt.pie([value for value in dict_risparmio.values()], labels=[str(round((value*100),2)).replace('.',',')+'%' if value > 0.03 else '' for value in dict_risparmio.values()], radius=1.2, colors=['#072FF9', '#EB1515'], pctdistance=0.2, labeldistance=0.6, rotatelabels =True, textprops={'fontsize':14, 'name':'Century Gothic', 'rotation_mode':'anchor', 'va':'center', 'ha':'center'}, normalize=True)
            ax.legend(['#072FF9', '#EB1515'], ['amministrato', 'gestito'], loc="best")
        finally:
            plt.savefig('img/risparmio_pie.png', bbox_inches='tight', pad_inches=0)


        #---Emittenti---#
        # Se gli emittenti sono tanti la tabella e il grafico diventano illeggibili.
        # Meglio fare il top 20 emittenti nel portafoglio.
        top = 20
        dict_emittenti = self.peso_emittente()
        if dict_emittenti is None:
            pass
        else:
            # Sort dictionary by issuer's weigth
            sorted_dict_emittenti = dict(sorted(dict_emittenti.items(), key=lambda x:x[1], reverse=True))
            # Get top 20 issuers
            # list_emittenti_top = [emittente for num, emittente in enumerate(list(sorted_dict_emittenti.keys())) if num < 20]
            # dict_emittenti_top = {key : value for key, value in sorted_dict_emittenti.items() if key in list_emittenti_top20}
            dict_emittenti_top = {key : value for key, value in list(sorted_dict_emittenti.items())[0:top]}
            fonts_emittenti = {emittente : str(hex(random.randint(0, 16777215)).replace('0x', '').zfill(6)) for emittente in dict_emittenti_top.keys()}
            # Tabella emittenti #

            # Header
            header_emittenti = ['', 'TOP '+str(top)+' EMITTENTI', 'Peso']
            dim_emittenti = [3.4, 47, 9.5]
            min_row, max_row = 1, 1
            min_col = max_col + SCARTO
            max_col = min_col + len(header_emittenti) - 1
            for col in ws_figure.iter_cols(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
                ws_figure[col[0].coordinate].value = header_emittenti[col[0].column-min_col]
                ws_figure[col[0].coordinate].alignment = Alignment(horizontal='center', vertical='center')
                ws_figure[col[0].coordinate].font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
                ws_figure[col[0].coordinate].fill = PatternFill(fill_type='solid', fgColor='595959')
                ws_figure[col[0].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
                ws_figure.column_dimensions[ws_figure[col[0].coordinate].column_letter].width = dim_emittenti[col[0].column-min_col]
            # Body
            min_row = min_row + 1
            max_row = min_row + len(dict_emittenti_top.keys()) - 1
            for row in ws_figure.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
                ws_figure[row[0].coordinate].fill = PatternFill(fill_type='solid', fgColor=fonts_emittenti[list(dict_emittenti_top.keys())[row[0].row-min_row]])
                ws_figure[row[0].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
                ws_figure[row[1].coordinate].value = list(dict_emittenti_top.keys())[row[0].row-min_row]
                ws_figure[row[1].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
                ws_figure[row[2].coordinate].value = dict_emittenti_top[ws_figure[row[1].coordinate].value]
                ws_figure[row[2].coordinate].number_format = '0.0%'
                ws_figure[row[2].coordinate].alignment = Alignment(horizontal='center')
                ws_figure[row[2].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            # Footer
            max_row = max_row + 1
            ws_figure.cell(max_row, min_col, value='TOTALE')
            ws_figure.cell(max_row, min_col).alignment = Alignment(horizontal='center', vertical='center')
            ws_figure.cell(max_row, min_col).font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
            ws_figure.cell(max_row, min_col).fill = PatternFill(fill_type='solid', fgColor='595959')
            ws_figure.cell(max_row, min_col).border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure.merge_cells(start_row=max_row, end_row=max_row, start_column=min_col, end_column=min_col+1)
            ws_figure.cell(max_row, min_col+2, value=sum(dict_emittenti_top.values()))
            assert sum(ws_figure.cell(i, min_col+2).value for i in range(min_row, max_row)) == ws_figure.cell(max_row, min_col+2).value
            ws_figure.cell(max_row, min_col+2).alignment = Alignment(horizontal='center', vertical='center')
            ws_figure.cell(max_row, min_col+2).font = Font(name='Arial', size=11, color='FFFFFF', bold=True)
            ws_figure.cell(max_row, min_col+2).fill = PatternFill(fill_type='solid', fgColor='595959')
            ws_figure.cell(max_row, min_col+2).border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_figure.cell(max_row, min_col+2).number_format = FORMAT_PERCENTAGE_00

            # # Grafico risparmio openpyxl
            # chart = PieChart()
            # chart.height = 4.77
            # chart.width = 6.77
            # labels = Reference(ws_figure, min_col=min_col+1, max_col=min_col+1, min_row=min_row, max_row=max_row-1)
            # data = Reference(ws_figure, min_col=min_col+2, max_col=min_col+2, min_row=min_row, max_row=max_row-1)
            # chart.add_data(data, titles_from_data=False)
            # chart.set_categories(labels)
            # chart.dataLabels = DataLabelList(dLblPos='bestFit')
            # chart.dataLabels.showVal = True
            # chart.dataLabels.textProperties = RichText(p=[Paragraph(pPr=ParagraphProperties(defRPr=CharacterProperties(sz=1100, b=True)), endParaRPr=CharacterProperties(sz=1100, b=True))])
            # chart.legend = None
            # # cambia colori delle fette
            # for _ in range(0,11):
            #     series = chart.series[0]
            #     pt = DataPoint(idx=_)
            #     pt.graphicalProperties.solidFill = fonts_emittenti[_]
            #     series.dPt.append(pt)
            # # posizione e dimensione figura
            # chart.layout = Layout(manualLayout=ManualLayout(x=0.5, y=0.5, h=1, w=1))
            # ws_figure.add_chart(chart, get_column_letter(min_col) + str(max_row + SCARTO))

            # Grafico torta emittente matplotlib
            plt.subplots(figsize=(4,4))
            try:
                plt.pie([value for value in dict_emittenti_top.values()], labels=[str(round((value*100),2)).replace('.',',')+'%' if value > 0.03 else '' for value in dict_emittenti_top.values()], radius=1.2, colors=['#'+font for font in fonts_emittenti.values()], pctdistance=0.2, labeldistance=0.6, rotatelabels =True, textprops={'fontsize':14, 'name':'Century Gothic', 'rotation_mode':'anchor', 'va':'center', 'ha':'center'}, normalize=False)
            except ValueError:
                plt.pie([value for value in dict_emittenti_top.values()], labels=[str(round((value*100),2)).replace('.',',')+'%' if value > 0.03 else '' for value in dict_emittenti_top.values()], radius=1.2, colors=['#'+font for font in fonts_emittenti.values()], pctdistance=0.2, labeldistance=0.6, rotatelabels =True, textprops={'fontsize':14, 'name':'Century Gothic', 'rotation_mode':'anchor', 'va':'center', 'ha':'center'}, normalize=True)
            finally:
                plt.savefig('img/emittenti_pie.png', bbox_inches='tight', pad_inches=0)

            # Grafico barre emittente matplotlib
            dict_emittenti_top_reversed = reversed(list(dict_emittenti_top.keys()))
            plt.subplots(figsize=(18,10))
            plt.barh(y=[_ for _ in reversed(list(dict_emittenti_top.keys()))], width=[round(_*100, 2) for _ in reversed(list(dict_emittenti_top.values()))], height=0.8, color=['#'+fonts_emittenti.get(emittente) for emittente in reversed(list(dict_emittenti_top.keys()))])
            plt.xticks(np.arange(0, round(max(list(dict_emittenti_top.values()))*100, 2)+1.0, step=5), rotation=0)
            plt.grid(linewidth=0.2)
            plt.savefig('img/emittenti_bar.png', bbox_inches='tight', pad_inches=0)

    def mappatura_fondi(self):
        """
        Crea la tabella e il grafico a barre della mappatura dei fondi.
        """
        # Carica il foglio fondi   
        fondi = self.wb['fondi']
        self.wb.active = fondi
        # Fondi
        df_portfolio = self.df_portfolio
        prodotti_gestiti = df_portfolio.loc[df_portfolio['strumento']=='fund']
        numero_prodotti_gestiti = prodotti_gestiti.nome.count()
        if numero_prodotti_gestiti > 0:
            # Mappatura dei fondi
            df_mappatura = self.df_mappatura
            df_mappatura_fondi = df_mappatura.loc[df_mappatura['ISIN'].isin(prodotti_gestiti['ISIN'])]
            # Header
            header = list(['ISIN', 'Nome']) + self.micro_asset_class
            dimensions = [23, 70.7, 23, 23, 23, 23, 23, 23, 23, 23, 23, 23, 23, 23, 23, 23]
            min_row = 1
            max_row = 1
            min_col = 11
            max_col = min_col + len(self.micro_asset_class) + 1
            for col in fondi.iter_cols(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
                fondi[col[0].coordinate].value = header[col[0].column-min_col]
                fondi[col[0].coordinate].alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                fondi[col[0].coordinate].font = Font(name='Calibri', size=16, color='FFFFFF', bold=True)
                fondi[col[0].coordinate].fill = PatternFill(fill_type='solid', fgColor='808080')
                fondi[col[0].coordinate].border = Border(top=Side(border_style='thin', color='000000'), right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
                fondi.column_dimensions[col[0].column_letter].width = dimensions[col[0].column-min_col]
            # Body
            min_row = min_row + 1
            max_row = max_row + numero_prodotti_gestiti
            for row in fondi.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
                fondi[row[0].coordinate].value = df_mappatura_fondi['ISIN'].values[row[0].row-min_row]
                fondi[row[0].coordinate].font = Font(name='Calibri', size=18, color='000000')
                fondi[row[0].coordinate].border = Border(top=Side(border_style='thin', color='000000'), right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
                fondi[row[1].coordinate].value = df_mappatura_fondi.loc[df_mappatura_fondi['ISIN']==fondi[row[0].coordinate].value, 'nome'].values[0]
                fondi[row[1].coordinate].font = Font(name='Calibri', size=18, color='000000')
                fondi[row[1].coordinate].border = Border(top=Side(border_style='thin', color='000000'), right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
                for _ in range(2, 16):
                    fondi[row[_].coordinate].value = df_mappatura_fondi.loc[df_mappatura_fondi['ISIN']==fondi[row[0].coordinate].value, fondi.cell(row=min_row-1, column=row[_].column).value].values[0]
                    fondi[row[_].coordinate].font = Font(name='Calibri', size=18, color='000000')
                    fondi[row[_].coordinate].border = Border(top=Side(border_style='thin', color='000000'), right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
                    fondi[row[_].coordinate].alignment = Alignment(horizontal='center')
                    fondi[row[_].coordinate].number_format = '0%'
            # Footer
            min_row = max_row + 1
            max_row = max_row + 2
            footer = ['CONTROVALORE TOTALE', 'PESO PERCENTUALE']
            for row in fondi.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
                fondi[row[1].coordinate].value = footer[row[0].row-min_row]
                fondi[row[1].coordinate].font = Font(name='Calibri', size=18, color='000000', bold=True)
                fondi[row[1].coordinate].border = Border(top=Side(border_style='thin', color='000000'), right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
                if row[0].row == min_row:
                    for _ in range(2, 16):
                        fondi[row[_].coordinate].value = np.nan_to_num(np.array(df_mappatura_fondi[fondi.cell(row=min_row-numero_prodotti_gestiti-1, column=row[_].column).value]), nan=0.0) @ np.array(prodotti_gestiti['controvalore_in_euro'])
                        fondi[row[_].coordinate].font = Font(name='Calibri', size=18, color='000000')
                        fondi[row[_].coordinate].border = Border(top=Side(border_style='thin', color='000000'), right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
                        fondi[row[_].coordinate].alignment = Alignment(horizontal='center')
                        fondi[row[_].coordinate].number_format = FORMAT_NUMBER_COMMA_SEPARATED1
                if row[0].row == max_row:
                    for _ in range(2, 16):
                        fondi[row[_].coordinate].value = fondi.cell(row=max_row-1, column=row[_].column).value / sum(prodotti_gestiti['controvalore_in_euro'])
                        fondi[row[_].coordinate].font = Font(name='Calibri', size=18, color='000000')
                        fondi[row[_].coordinate].border = Border(top=Side(border_style='thin', color='000000'), right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
                        fondi[row[_].coordinate].alignment = Alignment(horizontal='center')
                        fondi[row[_].coordinate].number_format = '0%'
            # Grafico micro bar matplotlib
            plt.subplots(figsize=(18,5))
            plt.bar(x=[_.replace('Altre Valute', 'Altro').replace('Obbligazionario', 'Obb').replace('Governativo', 'Gov').replace('All Maturities', '').replace('Aggregate', '').replace('North America', 'Nord america').replace('Pacific', 'Pacifico').replace('Emerging Markets', 'Emergenti') for _ in self.micro_asset_class], height=[fondi.cell(row=max_row, column=_).value for _ in range(min_col+2, max_col+1)], width=1, color=['#E4DFEC', '#CCC0DA', '#B1A0C7', '#92CDDC', '#00B0F0', '#0033CC', '#0070C0', '#1F497D', '#000080', '#F79646', '#FFCC66', '#DA5300', '#F62F00', '#EDF06A'])
            plt.xticks(rotation=25)
            plt.savefig('img/map_fondi_bar.png', bbox_inches='tight', pad_inches=0)

    def volatilità(self):
        """Calcola la volatilità del portafoglio"""
        vol = self.risk()
        ws_rischio = self.wb['rischio']
        self.wb.active = ws_rischio
        ws_rischio.cell(row=1, column=1, value=vol).number_format = FORMAT_PERCENTAGE_00

    def sintesi(self):
        """Crea la tabella da piazzare in fondo alla presentazione."""
        self.wb.create_sheet('sintesi')
        ws_sintesi = self.wb['sintesi']
        self.wb.active = ws_sintesi
        df_p = self.df_portfolio
        df_m = self.df_mappatura.drop(['TOTALE'], axis=1)
        dict_peso_macro = self.peso_macro()
        # Header
        header = ['ISIN', 'Asset class', 'Prodotto', 'Valore di mercato in euro', 'Peso']
        min_row = 1
        max_row = 2
        min_col = 1
        max_col = len(header)
        for col in ws_sintesi.iter_cols(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
            ws_sintesi[col[0].coordinate].value = header[col[0].column-min_col]
            ws_sintesi[col[0].coordinate].alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            ws_sintesi[col[0].coordinate].font = Font(name='Century Gothic', size=16, color='FFFFFF', bold=True)
            ws_sintesi[col[0].coordinate].fill = PatternFill(fill_type='solid', fgColor='808080')
            ws_sintesi[col[0].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_sintesi.row_dimensions[col[0].row].height = 30
            ws_sintesi.row_dimensions[col[1].row].height = 30
            ws_sintesi.merge_cells(start_row=min_row, end_row=max_row, start_column=col[0].column, end_column=col[0].column)
            # if ws_sintesi[col[0].coordinate].value == 'ISIN':
            #     ws_sintesi.column_dimensions[col[0].column_letter].width = 25
            # elif ws_sintesi[col[0].coordinate].value == 'Asset class':
            #     ws_sintesi.column_dimensions[col[0].column_letter].width = 56
            # elif ws_sintesi[col[0].coordinate].value == 'Prodotto':
            #     ws_sintesi.column_dimensions[col[0].column_letter].width = max([len(nome) for nome in df_m['nome'].values])*1.7
            # elif ws_sintesi[col[0].coordinate].value == 'Valore di mercato in euro':
            #     ws_sintesi.column_dimensions[col[0].column_letter].width = max(24.3, max([len(str(round(controvalore_in_euro,2))) for controvalore_in_euro in df_p['controvalore_in_euro'].values])*2.5)
            # elif ws_sintesi[col[0].coordinate].value == 'Peso':
            #     ws_sintesi.column_dimensions[col[0].column_letter].width = 13
        ws_sintesi.cell(row=min_row, column=len(header)+1, value='Warning')
        ws_sintesi.cell(row=min_row, column=len(header)+1).alignment = Alignment(horizontal='center', vertical='center')
        ws_sintesi.cell(row=min_row, column=len(header)+1).font = Font(name='Century Gothic', size=16, color='FFFFFF', bold=True)
        ws_sintesi.cell(row=min_row, column=len(header)+1).fill = PatternFill(fill_type='solid', fgColor='595959')
        ws_sintesi.cell(row=min_row, column=len(header)+1).border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
        ws_sintesi.merge_cells(start_row=min_row, end_row=min_row, start_column=len(header)+1, end_column=len(header)+3)
        ws_sintesi.cell(row=min_row+1, column=len(header)+1, value='C/R')
        ws_sintesi.cell(row=min_row+1, column=len(header)+1).alignment = Alignment(horizontal='center', vertical='center')
        ws_sintesi.cell(row=min_row+1, column=len(header)+1).font = Font(name='Century Gothic', size=16, color='FFFFFF', bold=True)
        ws_sintesi.cell(row=min_row+1, column=len(header)+1).fill = PatternFill(fill_type='solid', fgColor='595959')
        ws_sintesi.cell(row=min_row+1, column=len(header)+1).border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
        ws_sintesi.cell(row=min_row+1, column=len(header)+2, value='L')
        ws_sintesi.cell(row=min_row+1, column=len(header)+2).alignment = Alignment(horizontal='center', vertical='center')
        ws_sintesi.cell(row=min_row+1, column=len(header)+2).font = Font(name='Century Gothic', size=16, color='FFFFFF', bold=True)
        ws_sintesi.cell(row=min_row+1, column=len(header)+2).fill = PatternFill(fill_type='solid', fgColor='595959')
        ws_sintesi.cell(row=min_row+1, column=len(header)+2).border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
        ws_sintesi.cell(row=min_row+1, column=len(header)+3).fill = PatternFill(fill_type='solid', fgColor='595959')
        # Body
        df_m['asset_class'] = df_m.apply(lambda x : x[self.micro_asset_class].index[x[self.micro_asset_class] == 1.00].values[0] if any(x[self.micro_asset_class]==1.00) else 'Prodotto multi asset', axis=1)
        custom_dict = {value : num for num, value in enumerate(self.micro_asset_class)}
        custom_dict['Prodotto multi asset'] = 14
        # ordina il dataframe portfolio_valori per tipo di strumento così come è stato mappato (dalla liquidità al prodotto composto)
        df_m.sort_values(by=['asset_class'], inplace=True, key=lambda x : x.map(custom_dict))
        # print(df_m)
        min_row = min_row + 2
        max_row = min_row + len(df_m) - 1
        df_p.fillna('', inplace=True)
        df_m.fillna('', inplace=True)
        for row in ws_sintesi.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col+3):
            ws_sintesi[row[0].coordinate].value = df_m['ISIN'].values[row[0].row-min_row]
            ws_sintesi[row[0].coordinate].font = Font(name='Century Gothic', size=16, color='000000')
            ws_sintesi[row[0].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_sintesi[row[1].coordinate].value = df_m['asset_class'].values[row[0].row-min_row].replace('Obbligazionario Euro Governativo All Maturities', 'Obbligazionario Euro Governativo')
            ws_sintesi[row[1].coordinate].font = Font(name='Century Gothic', size=16, color='000000')
            ws_sintesi[row[1].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_sintesi[row[2].coordinate].value = df_m['nome'].values[row[0].row-min_row]
            ws_sintesi[row[2].coordinate].font = Font(name='Century Gothic', size=16, color='000000')
            ws_sintesi[row[2].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_sintesi[row[3].coordinate].value = df_p.loc[(df_p['ISIN']==ws_sintesi[row[0].coordinate].value) & (df_p['nome']==ws_sintesi[row[2].coordinate].value), 'controvalore_in_euro'].values[0]
            ws_sintesi[row[3].coordinate].font = Font(name='Century Gothic', size=16, color='000000')
            ws_sintesi[row[3].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_sintesi[row[3].coordinate].number_format = '€ #,0.00'
            ws_sintesi[row[4].coordinate].value = df_p.loc[(df_p['ISIN']==ws_sintesi[row[0].coordinate].value) & (df_p['nome']==ws_sintesi[row[2].coordinate].value), 'controvalore_in_euro'].values[0] / df_p['controvalore_in_euro'].sum()
            ws_sintesi[row[4].coordinate].font = Font(name='Century Gothic', size=16, color='000000')
            ws_sintesi[row[4].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_sintesi[row[4].coordinate].number_format = FORMAT_PERCENTAGE_00
            if df_p.loc[(df_p['ISIN']==ws_sintesi[row[0].coordinate].value) & (df_p['nome']==ws_sintesi[row[2].coordinate].value), 'strumento'].values[0] == 'gov_bond' or df_p.loc[(df_p['ISIN']==ws_sintesi[row[0].coordinate].value) & (df_p['nome']==ws_sintesi[row[2].coordinate].value), 'strumento'].values[0] == 'corp_bond':
                if ws_sintesi[row[4].coordinate].value / dict_peso_macro['Obbligazionario'] > 0.30:
                    ws_sintesi[row[5].coordinate].value = '!!!C'
                    ws_sintesi[row[5].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
                elif ws_sintesi[row[4].coordinate].value / dict_peso_macro['Obbligazionario'] > 0.20:
                    ws_sintesi[row[5].coordinate].value = '!!C'
                    ws_sintesi[row[5].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
                elif ws_sintesi[row[4].coordinate].value / dict_peso_macro['Obbligazionario'] > 0.10:
                    ws_sintesi[row[5].coordinate].value = '!C'
                    ws_sintesi[row[5].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
            elif df_p.loc[(df_p['ISIN']==ws_sintesi[row[0].coordinate].value) & (df_p['nome']==ws_sintesi[row[2].coordinate].value), 'strumento'].values[0] == 'equity':
                if ws_sintesi[row[4].coordinate].value / dict_peso_macro['Azionario'] > 0.30:
                    ws_sintesi[row[5].coordinate].value = '!!!C'
                    ws_sintesi[row[5].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
                elif ws_sintesi[row[4].coordinate].value / dict_peso_macro['Azionario'] > 0.20:
                    ws_sintesi[row[5].coordinate].value = '!!C'
                    ws_sintesi[row[5].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
                elif ws_sintesi[row[4].coordinate].value / dict_peso_macro['Azionario'] > 0.10:
                    ws_sintesi[row[5].coordinate].value = '!C'
                    ws_sintesi[row[5].coordinate].fill = PatternFill(fill_type='solid', fgColor='FFD700')
            ws_sintesi[row[5].coordinate].font = Font(name='Century Gothic', size=16, color='000000', bold=True)
            ws_sintesi[row[5].coordinate].alignment = Alignment(horizontal='center')
            ws_sintesi[row[5].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_sintesi[row[6].coordinate].font = Font(name='Century Gothic', size=16, color='FF0000', bold=True)
            ws_sintesi[row[6].coordinate].alignment = Alignment(horizontal='center')
            ws_sintesi[row[6].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_sintesi[row[7].coordinate].border = Border(right=Side(border_style='thin', color='000000'), left=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
        # Footer
        max_row = max_row + 1
        for row in ws_sintesi.iter_rows(min_row=max_row, max_row=max_row, min_col=min_col+1, max_col=max_col+3):
            ws_sintesi[row[0].coordinate].value = 'TOTALE PORTAFOGLIO'
            ws_sintesi[row[0].coordinate].font = Font(name='Century Gothic', size=16, color='000000', bold=True)
            ws_sintesi[row[0].coordinate].alignment = Alignment(horizontal='center')
            ws_sintesi[row[0].coordinate].border = Border(left=Side(border_style='thin', color='000000'), right=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_sintesi.merge_cells(start_row=max_row, end_row=max_row, start_column=min_col+1, end_column=min_col+2)
            ws_sintesi[row[2].coordinate].value = df_p['controvalore_in_euro'].sum()
            ws_sintesi[row[2].coordinate].font = Font(name='Century Gothic', size=16, bold=True)
            ws_sintesi[row[2].coordinate].number_format = '€ #,0.00'
            ws_sintesi[row[2].coordinate].border = Border(left=Side(border_style='thin', color='000000'), right=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            np.testing.assert_almost_equal(actual=sum([ws_sintesi.cell(row=_, column=min_col+3).value for _ in range(min_row, max_row)]), desired=ws_sintesi[row[2].coordinate].value, decimal=2, err_msg='La somma dei valori dei singoli strumenti non è uguale al controvalore totale di portafoglio', verbose=True)
            ws_sintesi[row[3].coordinate].value = sum([ws_sintesi.cell(row=_, column=min_col+4).value for _ in range(min_row, max_row)])
            ws_sintesi[row[3].coordinate].font = Font(name='Century Gothic', size=16, bold=True)
            np.testing.assert_almost_equal(actual=ws_sintesi[row[3].coordinate].value, desired=1.00, decimal=1, err_msg='La sommma dei pesi non fa 100', verbose=True)
            ws_sintesi[row[3].coordinate].number_format = FORMAT_PERCENTAGE_00
            ws_sintesi[row[3].coordinate].border = Border(left=Side(border_style='thin', color='000000'), right=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_sintesi[row[4].coordinate].border = Border(left=Side(border_style='thin', color='000000'), right=Side(border_style='thin', color='000000'), top=Side(border_style='thin', color='000000'), bottom=Side(border_style='thin', color='000000'))
            ws_sintesi.merge_cells(start_row=max_row, end_row=max_row, start_column=max_col+1, end_column=max_col+3)

    def salva_file_portafoglio(self):
        """Salva il file excel."""
        try:
            self.wb.save(self.file_elaborato)
        except PermissionError:
            for proc in psutil.process_iter():
                if proc.name() == "EXCEL.EXE":
                    proc.kill()
            self.wb.save(self.file_elaborato)

    def autofit(self, sheet, columns, min_width, max_width):
        """
        Imposta la miglior lunghezza per le colonne selezionate.
        # TODO : accetta più di un foglio
        # TODO: accetta anche lettere per selezionare le colonne.
        # TODO: se columns è vuoto, autofit tutte le colonne.

        Parameters:
            sheet {string} = foglio excel da formattare
            columns {list} = lista contenente il numero o le lettere delle colonne da formattare. if not columns: formatta tutte le colonne del foglio
            min_width {list} = lista contenente la lunghezza massima in pixels della colonna, che l'autofit potrebbe non superare (usa None se non serve su una data colonna)
            max_width {list} = lista contenente la lunghezza massima in pixels della colonna, che l'autofit potrebbe superare (usa None se non serve su una data colonna)
        """
        xls_file = win32com.client.Dispatch("Excel.Application")
        xls_file.visible = False
        wb = xls_file.Workbooks.Open(Filename=self.path.joinpath(self.file_elaborato).__str__())
        ws = wb.Worksheets(sheet)
        for num, value in enumerate(columns):
            if value > 0: # la colonna 0 e le negative non esistono
                ws.Columns(value).AutoFit()
                if max_width[num] is not None:
                    if ws.Columns(value).ColumnWidth > max_width[num]:
                        ws.Columns(value).ColumnWidth = max_width[num]
                if min_width[num] is not None:
                    if ws.Columns(value).ColumnWidth < min_width[num]:
                        ws.Columns(value).ColumnWidth = min_width[num]
            else:
                continue
        xls_file.DisplayAlerts = False
        wb.Close(SaveChanges=True, Filename=self.path.joinpath(self.file_elaborato).__str__())
        xls_file.Quit()


class Presentazione(Portfolio):
    """Tentativo di ricreare la presentazione. Inherits from Portfolio."""

    def __init__(self, intermediario, tipo_sap, file_portafoglio=None, **dimensioni):
        """
        Initialize the class.

        Parameters:
            file_portafoglio {str} = file da analizzare
            intermediario {str} = intermediario per cui fare l'analisi
            tipo_sap {str} = completo o light
            **dimensioni {dict} = dimensioni delle pagine word
        """
        super().__init__(intermediario=intermediario, file_portafoglio=file_portafoglio)

        self.tipo_sap = tipo_sap
        if tipo_sap != 'completo' and tipo_sap != 'light':
            print('Il tipo di SAP può essere completo o light!')
            quit()
        
        self.path_img = self.path.joinpath('img')
        self.path_img_default = self.path.joinpath('img', 'default')
        
        # Microsoft Excel
        if file_portafoglio is None: # se non viene inserito l'argument file_portafoglio
            ptf = self.file_portafoglio.name
        else:
            ptf = self.file_portafoglio
        self.file_elaborato = ptf[:-5] + '_elaborato.xlsx'
        self.wb = load_workbook(self.file_elaborato) # apre il file elaborato
        
        # Microsoft Word
        self.document = Document() # documento word per docx
        # proprietà documento
        self.document.core_properties.title = 'SAP'
        self.document.core_properties.subject = 'Analisi di portafoglio'
        self.document.core_properties.category = 'Financial analysis'
        self.document.core_properties.author = 'B&S'
        self.document.core_properties.comments = ''
        self.file_presentazione = ptf[:-5] + '.docx'
    
        self.page_height = dimensioni['page_height']
        self.page_width = dimensioni['page_width']
        self.top_margin = dimensioni['top_margin']
        self.bottom_margin = dimensioni['bottom_margin']
        self.left_margin = dimensioni['left_margin']
        self.right_margin = dimensioni['right_margin']
        self.larghezza_pagina = self.page_width - self.left_margin - self.right_margin
        self.altezza_pagina = self.page_height - self.top_margin - self.bottom_margin

    def copertina(self):
        """
        Copertina della presentazione
        """
        # 1.Copertina
        section = self.document.sections[0]
        # Imposta dimensioni A4
        section.page_height = shared.Cm(self.page_height)
        section.page_width = shared.Cm(self.page_width)
        # Imposta margini
        left_margin = 0.60
        right_margin = 0.60
        top_margin = 0.45
        bottom_margin = 0.45
        section.left_margin = shared.Cm(left_margin)
        section.right_margin = shared.Cm(right_margin)
        section.top_margin = shared.Cm(top_margin)
        section.bottom_margin = shared.Cm(bottom_margin)
        section.header_distance = shared.Cm(0)
        section.footer_distance = shared.Cm(0)
        # Image
        if self.intermediario == 'azimut':
            paragraph = self.document.add_paragraph()
            paragraph.alignment = 1
            copertina = 'copertina_completo.jpg' if self.tipo_sap=='completo' else 'copertina_light.jpg' if self.tipo_sap=='light' else print('Il tipo di SAP può essere completo o light!')
            paragraph.add_run().add_picture(self.path_img_default.joinpath(copertina).__str__(), height=shared.Cm(self.page_height-top_margin-bottom_margin), width=shared.Cm(self.page_width-left_margin-right_margin))
        elif self.intermediario == 'copernico':
            pass

    def indice(self):
        """
        Indice della presentazione
        """
        # 2.Indice
        paragraph_format = self.document.styles['Normal'].paragraph_format
        paragraph_format.space_after = 0 # Annulla lo spazio dopo il testo per tutte le stringhe di tipo normale.
        self.document.add_section()
        section = self.document.sections[1]
        # Imposta margini
        section.top_margin = shared.Cm(self.top_margin)
        section.bottom_margin = shared.Cm(self.bottom_margin)
        section.right_margin = shared.Cm(self.right_margin)
        section.left_margin = shared.Cm(self.left_margin)
        # Header
        header = section.header
        header.is_linked_to_previous = False # Se True crea l'header anche per la pagina precedente
        paragraph = header.paragraphs[0]
        if self.intermediario == 'azimut':
            paragraph.add_run().add_picture(self.path_img_default.joinpath('logo_azimut.bmp').__str__(), height=shared.Cm(1.4), width=shared.Cm(3.72))
            paragraph_0 = self.document.add_paragraph()
            run_0 = paragraph_0.add_run('\n')
            run_0.add_picture(self.path_img_default.joinpath('indice.bmp').__str__(), width=shared.Cm(12.5))
        elif self.intermediario == 'copernico':
            paragraph.paragraph_format.alignment = 2
            paragraph.add_run().add_picture(self.path_img_default.joinpath('logo_copernico.png').__str__(), height=shared.Cm(2), width=shared.Cm(4.2))
            # Title
            paragraph_0 = self.document.add_paragraph()
            run_0 = paragraph_0.add_run('\n')
            run_0 = paragraph_0.add_run(text='INDICE', style=None)
            run_0.bold = True
            run_0.font.name = 'Century Gothic'
            run_0.font.size = shared.Pt(24)
            run_0.font.color.rgb = shared.RGBColor(127, 127, 127)
            run_0 = paragraph_0.add_run('\n')
            # Body - list numbers
            paragraph_1 = self.document.add_paragraph('PORTAFOGLIO ATTUALE', style='List Number')
            paragraph_1.style.font.name = 'Century Gothic'
            paragraph_1.style.font.size = shared.Pt(14)
            paragraph_1.style.font.color.rgb = shared.RGBColor(127, 127, 127)
            paragraph_1.add_run('\n')
            paragraph_2 = self.document.add_paragraph('ANALISI DEL PORTAFOGLIO', style='List Number')
            run_2 = paragraph_2.add_run('\n\tPer macro asset class\n\tPer micro asset class\n\tPer tipologia di prodotto\n\tPer valuta')
            run_2.font.name = 'Century Gothic'
            run_2.font.size = shared.Pt(12)
            run_2.font.color.rgb = shared.RGBColor(127, 127, 127)
            paragraph_2.add_run('\n')
            paragraph_3 = self.document.add_paragraph('ANALISI DEI SINGOLI STRUMENTI', style='List Number')
            paragraph_3.add_run('\n')
            paragraph_4 = self.document.add_paragraph('ANALISI DEL RISCHIO', style='List Number')
            paragraph_4.add_run('\n')
            self.document.add_paragraph('NOTE METODOLOGICHE', style='List Number')

    def portafoglio_attuale(self, method):
        # TODO: con il 'method on top' e limite 60 sul portafoglio ptf_20.xlsx: C:\Users\Administrator\Desktop\Sbwkrq\SAP\bugs\agglomerato_2.png
        """
        Portafoglio complessivo diviso per strumenti.
        Metodo 1 (basic) : stampa i primi 57 senza riportare come prima riga dopo l'intestazione l'etichetta del primo strumento a comparire.
        Metodo 2 : (label on top) : stampa i primi 57 riportando sempre come prima riga dopo l'intestazione l'etichetta del primo strumento a comparire.

        Parameters
        method(str) = metodo con cui creare le immagini dell'agglomerato (basic, label_on_top)
        """
        df = self.df_portfolio
        if all(df['quantità'].isnull()):
            print("Mancano le quantità")
        if all(df['controvalore_iniziale_in_euro'].isnull()):
            print("Mancano i controvalori iniziali")
        if all(df['prezzo_di_carico'].isnull()):
            print("Mancano i prezzi di carico")
        sheet = self.wb['agglomerato']
        self.wb.active = sheet
        # Nascondi colonne vuote
        hidden_columns = 0
        if all(df['quantità'].isnull()):
            sheet.column_dimensions['C'].hidden= True
            hidden_columns += 1
        if all(df['controvalore_iniziale_in_euro'].isnull()):
            sheet.column_dimensions['D'].hidden= True
            hidden_columns += 1
        if all(df['prezzo_di_carico'].isnull()):
            sheet.column_dimensions['E'].hidden= True
            hidden_columns += 1
        self.wb.save(self.file_elaborato)
        if method == 'basic':
            c = Counter(list(df.loc[:, 'strumento']))
            strumenti_in_ptf = [strumento for strumento in self.strumenti if c[strumento] > 0]
            max_row = 1 + df['nome'].count() + len(strumenti_in_ptf) + 1
            LIMITE= 60
            if max_row <= LIMITE:
                tabelle_agglomerato = 1
            else:
                if max_row % LIMITE != 0:
                    tabelle_agglomerato = max_row // LIMITE + 1
                elif max_row % LIMITE == 0:
                    tabelle_agglomerato = max_row // LIMITE
            # # Librerie win32com + PIL
            # xls_file = win32com.client.Dispatch("Excel.Application")
            # # xls_file.Visible = False
            # # xls_file.ScreenUpdating = False
            # # xls_file.DisplayAlerts = False
            # # xls_file.EnableEvents = False
            # wb = xls_file.Workbooks.Open(Filename=self.path+"\\"+self.file_elaborato)
            # ws = wb.Worksheets("agglomerato")
            for tabella in range(1, tabelle_agglomerato+1):
                if tabella != tabelle_agglomerato:
                    # # Librerie win32com + PIL
                    # ws.Range(ws.Cells(1,1),ws.Cells(str(LIMITE*tabella),9)).CopyPicture(Format=2)
                    # img = ImageGrab.grabclipboard()
                    # img.save(self.path+'\img\\agglomerato_'+str(tabella-1)+'.png')
                    # ws.Range(ws.Cells(2,1),ws.Cells(LIMITE*tabella,9)).Rows.EntireRow.Hidden = True
                    # Libreria excel2img
                    excel2img.export_img(self.file_elaborato, self.path_img.joinpath('agglomerato_'+str(tabella-1)+'.png').__str__(), page='agglomerato', _range="A1:I"+str(LIMITE*tabella))
                    sheet.row_dimensions.group(2,LIMITE*tabella,hidden=True)
                    self.wb.save(self.file_elaborato)
                elif tabella == tabelle_agglomerato:
                    # Librerie win32com + PIL
                    # ws.Range(ws.Cells(1,1),ws.Cells(str(max_row),9)).CopyPicture(Format=2)
                    # img = ImageGrab.grabclipboard()
                    # img.save(self.path+'\img\\agglomerato_'+str(tabella-1)+'.png')
                    # Libreria excel2img
                    excel2img.export_img(self.file_elaborato, self.path_img.joinpath('agglomerato_'+str(tabella-1)+'.png').__str__(), page='agglomerato', _range="A1:I"+str(max_row))
            # wb.Close(SaveChanges=False, Filename=self.path+"\\"+self.file_elaborato)
            for tabella in range(1, tabelle_agglomerato+1):
                print(f"sto aggiungendo l'agglomerato {tabella-1} alla presentazione...")
                self.document.add_section()
                paragraph_0 = self.document.add_paragraph(text='', style=None)
                paragraph_0.paragraph_format.space_before = shared.Pt(6)
                paragraph_0.paragraph_format.space_after = shared.Pt(6)
                run_0 = paragraph_0.add_run(text='')
                run_0.add_picture(self.path_img_default.joinpath('1_portafoglio_attuale.bmp').__str__(), width=shared.Cm(8.5))
                paragraph_1 = self.document.add_paragraph(style=None)
                paragraph_1.paragraph_format.space_before = shared.Pt(6)
                paragraph_1.paragraph_format.space_after = shared.Pt(6)
                run_1 = paragraph_1.add_run()
                width = self.larghezza_pagina if hidden_columns==0 else self.larghezza_pagina - 1 if hidden_columns==1 else self.larghezza_pagina - 2 if hidden_columns==2 else self.larghezza_pagina - 3 if hidden_columns==3 else self.larghezza_pagina
                run_1.add_picture(self.path_img.joinpath('agglomerato_'+ str(tabella-1) +'.png').__str__(), width=shared.Cm(width))
            # sheet.row_dimensions.group(2,LIMITE*(tabelle_agglomerato),hidden=False)
        elif method == 'label_on_top':
            c = Counter(list(df.loc[:, 'strumento']))
            strumenti_in_ptf = [strumento for strumento in self.strumenti if c[strumento] > 0]
            max_row = 1 + df['nome'].count() + len(strumenti_in_ptf) + 1
            LIMITE = 61
            tabella = 0
            posizione_labels = {}
            numerosità_cumulata = 2 # una riga sotto la label
            for strumento in strumenti_in_ptf:
                if strumento == strumenti_in_ptf[0]:
                    posizione_labels[strumento] = numerosità_cumulata
                    numerosità_cumulata += c[strumento]
                else:
                    posizione_labels[strumento] = numerosità_cumulata + 1
                    numerosità_cumulata += c[strumento] + 1
            riga_cumulata = 1
            # # Librerie win32com + PIL
            # xls_file = win32com.client.Dispatch("Excel.Application")
            # # xls_file.Visible = False
            # # xls_file.ScreenUpdating = False
            # # xls_file.DisplayAlerts = False
            # # xls_file.EnableEvents = False
            # wb = xls_file.Workbooks.Open(Filename=self.path+"\\"+self.file_elaborato)
            # ws = wb.Worksheets("agglomerato")
            while strumenti_in_ptf: # prova quando ci sono più di 63 titoli nello stesso strumento
                riga = 1 # l'intestazione
                for num, strumento in enumerate(strumenti_in_ptf[:]):
                    numerosità_strumento = c[strumento]
                    _ = Counter({strumento : numerosità_strumento})
                    numerosità_strumento_più_label = numerosità_strumento + 1
                    riga += numerosità_strumento_più_label
                    if riga <= LIMITE:
                        c.subtract(_)
                        strumenti_in_ptf.remove(strumento)
                        riga_cumulata += numerosità_strumento_più_label
                    else:
                        scarto = riga - LIMITE
                        __ = Counter({strumento : numerosità_strumento-scarto})
                        c.subtract(__)
                        riga_cumulata += numerosità_strumento_più_label - scarto
                        # # Librerie win32com + PIL
                        # ws.Range(ws.Cells(1,1),ws.Cells(str(riga_cumulata),9)).CopyPicture(Format=2)
                        # img = ImageGrab.grabclipboard()
                        # img.save(self.path+'\img\\agglomerato_'+str(tabella)+'.png')
                        # ws.Range(ws.Cells(2,1),ws.Cells(riga_cumulata,9)).Rows.EntireRow.Hidden = True
                        # ws.Rows(posizione_labels[strumento]).EntireRow.Hidden = False
                        # Libreria excel2img
                        excel2img.export_img(self.file_elaborato, self.path_img.joinpath('agglomerato_'+str(tabella)+'.png').__str__(), page='agglomerato', _range="A1:I"+str(riga_cumulata))
                        sheet.row_dimensions.group(2, riga_cumulata, hidden=True)
                        sheet.row_dimensions[posizione_labels[strumento]].hidden = False
                        self.wb.save(self.file_elaborato)
                        print(f"sto aggiungendo l'agglomerato {tabella} alla presentazione...")
                        self.document.add_section()
                        paragraph_0 = self.document.add_paragraph(text='', style=None)
                        paragraph_0.paragraph_format.space_before = shared.Pt(6)
                        paragraph_0.paragraph_format.space_after = shared.Pt(6)
                        run_0 = paragraph_0.add_run(text='')
                        run_0.add_picture(self.path_img_default.joinpath('1_portafoglio_attuale.bmp').__str__(), width=shared.Cm(8.5))
                        paragraph_1 = self.document.add_paragraph(style=None)
                        paragraph_1.paragraph_format.space_before = shared.Pt(6)
                        paragraph_1.paragraph_format.space_after = shared.Pt(6)
                        run_1 = paragraph_1.add_run()
                        width = self.larghezza_pagina if hidden_columns==0 else self.larghezza_pagina - 1 if hidden_columns==1 else self.larghezza_pagina - 2 if hidden_columns==2 else self.larghezza_pagina - 3 if hidden_columns==3 else self.larghezza_pagina
                        run_1.add_picture(self.path_img.joinpath('agglomerato_'+ str(tabella) +'.png').__str__(), width=shared.Cm(width))
                        riga_cumulata -= 1
                        tabella += 1
                        break
            np.testing.assert_equal(riga_cumulata+1, max_row, err_msg="L'ultima riga cumulata non corrisponde all'ultima riga effettiva nel file excel", verbose=True)
            # # Librerie win32com + PIL
            # ws.Range(ws.Cells(1,1),ws.Cells(str(riga_cumulata+1),9)).CopyPicture(Format=2)
            # img = ImageGrab.grabclipboard()
            # img.save(self.path+'\img\\agglomerato_'+str(tabella)+'.png')
            # wb.Close(SaveChanges=False, Filename=self.path+"\\"+self.file_elaborato)
            # Libreria excel2img
            excel2img.export_img(self.file_elaborato, self.path_img.joinpath('agglomerato_'+str(tabella)+'.png').__str__(), page='agglomerato', _range="A1:I"+str(riga_cumulata+1))
            sheet.row_dimensions.group(2, riga_cumulata+1, hidden=False)
            print(f"sto aggiungendo l'agglomerato {tabella} alla presentazione...")
            self.document.add_section()
            paragraph_0 = self.document.add_paragraph(text='', style=None)
            paragraph_0.paragraph_format.space_before = shared.Pt(6)
            paragraph_0.paragraph_format.space_after = shared.Pt(6)
            run_0 = paragraph_0.add_run(text='')
            run_0.add_picture(self.path_img_default.joinpath('1_portafoglio_attuale.bmp').__str__(), width=shared.Cm(8.5))
            paragraph_1 = self.document.add_paragraph(style=None)
            paragraph_1.paragraph_format.space_before = shared.Pt(6)
            paragraph_1.paragraph_format.space_after = shared.Pt(6)
            run_1 = paragraph_1.add_run()
            width = self.larghezza_pagina if hidden_columns==0 else self.larghezza_pagina - 1 if hidden_columns==1 else self.larghezza_pagina - 2 if hidden_columns==2 else self.larghezza_pagina - 3 if hidden_columns==3 else self.larghezza_pagina
            run_1.add_picture(self.path_img.joinpath('agglomerato_'+ str(tabella) +'.png').__str__(), width=shared.Cm(width))

    def commento(self):
        """Commento alla composizione del portafoglio."""
        self.document.add_section()
        paragraph_0 = self.document.add_paragraph(text='', style=None)
        paragraph_0.paragraph_format.space_before = shared.Pt(6)
        run_0 = paragraph_0.add_run(text='')
        run_0.add_picture(self.path_img_default.joinpath('1_portafoglio_attuale.bmp').__str__(), width=shared.Cm(8.5))
        paragraph_1 = self.document.add_paragraph(text='\n', style=None)
        run_1 = paragraph_1.add_run('Commento generale sul portafoglio')
        run_1.bold = True
        run_1.font.name = 'Century Gothic'
        run_1.font.size = shared.Pt(14)
        run_1.font.color.rgb = shared.RGBColor(127, 127, 127)
        paragraph_2 = self.document.add_paragraph(style=None)
        paragraph_2.paragraph_format.space_after = shared.Pt(6)
        paragraph_2.paragraph_format.line_spacing_rule = 1

        # Carica il dataframe del portafoglio per estrane la composizione ed eventuali alert.
        df_portfolio = self.df_portfolio
        # Carica i dizionari delle macro e micro classi.
        dict_peso_macro = self.peso_macro()
        dict_peso_micro = self.peso_micro()
        # Crea il commento
        run_2 = paragraph_2.add_run(f'\nIl portafoglio attuale è investito ')
        run_2.font.name = 'Century Gothic'
        run_2.font.size = shared.Pt(10)
        # Trasformazione nome pesi in italiano
        dict_strumenti = self.peso_strumenti()
        dict_strumenti_attivi = {k : v * 100 for k, v in dict_strumenti.items() if v!=0}
        dict_strumenti_attivi = {k: v for k, v in sorted(dict_strumenti_attivi.items(), key=lambda item: item[1], reverse=True)}
        # Lista degli strumenti possibili nel file di input tradotti in italiano per il commento
        self.dict_str_comm = {'cash' : 'liquidità', 'gov_bond' : 'obbligazioni governative', 'corp_bond' : 'obbligazioni societarie',
            'certificate' : 'certificati', 'equity' : 'azioni', 'etf' : 'etf', 'fund' : 'fondi', 'real_estate' : 'real estate',
            'hedge_fund' : 'fondi hedge', 'private_equity' : 'private equity', 'venture_capital' : 'venture capital', 'private_debt' : 'private debt',
            'insurance' : 'polizze', 'gp' : 'gestioni patrimoniali', 'pip' : 'fondi pensione', 'alternative' : 'altro'}
        dict_peso_strumenti_attivi = {self.dict_str_comm[k] : v for k, v in dict_strumenti_attivi.items()}
        for strumento, peso in dict_peso_strumenti_attivi.items():
            articolo = 'il ' if int(str(peso)[0]) in (2, 3, 4, 5, 6, 7, 9) else 'lo ' if int(str(peso)[0]) == 0 else "l'" if int(str(peso)[0]) == 8 else "l'" if int(str(peso)[0]) == 1 and peso < 12 else "il "
            if  strumento not in list(dict_peso_strumenti_attivi.keys())[-1]:
                run = paragraph_2.add_run(f"per {articolo}{str(round(peso,2)).replace('.',',')}% in {strumento}, ")
                run.font.name = 'Century Gothic'
                run.font.size = shared.Pt(10)
            else:
                run = paragraph_2.add_run(f"e per {articolo}{str(round(peso,2)).replace('.',',')}% in {strumento}.")
                run.font.name = 'Century Gothic'
                run.font.size = shared.Pt(10)
        
        # Alert mercati azionari
        paragraph_3 = self.document.add_paragraph()
        paragraph_3.paragraph_format.space_after = shared.Pt(6)
        paragraph_3.paragraph_format.line_spacing_rule = 1

        dict_peso_micro_azionarie = {item : dict_peso_micro[item]/dict_peso_macro[key] for key, value in self.dict_macro_micro.items() for item in value if key=='Azionario'}
        pesi_limite_azionari = {'Azionario Europa' : 0.60, 'Azionario North America' : 0.60, 'Azionario Pacific' : 0.20, 'Azionario Emerging Markets' : 0.10}
        dict_alert_azionari = {k : True if v >= pesi_limite_azionari[k] else False for k, v in dict_peso_micro_azionarie.items()}
        nome_mercati_azionari = {'Azionario Europa' : 'europei', 'Azionario North America' : 'nordamericani', 'Azionario Pacific' : "nell'area del pacifico", 'Azionario Emerging Markets' : "emergenti"}
        
        if sum(dict_alert_azionari.values()) == 1:
            run_3 = paragraph_3.add_run(f"""Per la parte investita nel mercato azionario, si segnala l’eccessiva concentrazione verso i mercati {''.join([nome_mercati_azionari[k] for k, v in dict_alert_azionari.items() if v==True])}, che pesano {' '.join([str('il ' if str(dict_peso_micro_azionarie[k]*100)[0].startswith(('1','2','3','4','5','6','7','9')) and str(dict_peso_micro_azionarie[k]*100)[:2]!='11' else "l'")+str(round(dict_peso_micro_azionarie[k]*100, 2)).replace('.', ',')+str('%') for k, v in dict_alert_azionari.items() if v==True])} del comparto azionario, come indicato dal relativo warning. """)
            run_3.font.name = 'Century Gothic'
            run_3.font.size = shared.Pt(10)
        elif sum(dict_alert_azionari.values()) > 1:
            run_3 = paragraph_3.add_run(f"""Per la parte investita nel mercato azionario, si segnala l’eccessiva concentrazione verso i mercati {' e '.join([nome_mercati_azionari[k] for k, v in dict_alert_azionari.items() if v==True])} che pesano rispettivamente {', e '.join([str('il ' if str(dict_peso_micro_azionarie[k]*100)[0].startswith(('1','2','3','4','5','6','7','9')) and str(dict_peso_micro_azionarie[k]*100)[:2]!='11' else "l'")+str(round(dict_peso_micro_azionarie[k]*100, 2)).replace('.', ',')+str('%') for k, v in dict_alert_azionari.items() if v==True])} del comparto azionario, come indicato dai relativi warning. """)
            run_3.font.name = 'Century Gothic'
            run_3.font.size = shared.Pt(10)
        
        # Alert prodotti azionari
        quote_prodotti_azionari = df_portfolio.loc[df_portfolio['strumento']=='equity', ['nome', 'controvalore_in_euro']]
        quote_prodotti_azionari['peso_totale'] = quote_prodotti_azionari['controvalore_in_euro'] / df_portfolio['controvalore_in_euro'].sum()
        quote_prodotti_azionari['peso_su_azionario'] = quote_prodotti_azionari['peso_totale'] / dict_peso_macro['Azionario']
        if any(quote_prodotti_azionari['peso_su_azionario'] > 0.10):
            quote_prodotti_azionari_alert = dict(zip(quote_prodotti_azionari.loc[quote_prodotti_azionari['peso_su_azionario']>0.10, 'nome'].values, quote_prodotti_azionari.loc[quote_prodotti_azionari['peso_su_azionario']>0.10, 'peso_su_azionario']))
            run_3 = paragraph_3.add_run(f"""Riguardo agli strumenti azionari, si segnala il peso eccessivo di {' e '.join([k for k,v in quote_prodotti_azionari_alert.items()])} {'che pesa' if len(quote_prodotti_azionari_alert)==1 else 'che pesano rispettivamente'} {', e '.join([str('il ' if str(v*100)[0].startswith(('1','2','3','4','5','6','7','9')) and str(v*100)[:2]!='11' else "l'")+str(round(v*100, 2)).replace('.', ',')+str('%') for k,v in quote_prodotti_azionari_alert.items()])} dell’intero comparto azionario, come indicato {'dal relativo warning' if len(quote_prodotti_azionari_alert)==1 else 'dai relativi warning'} nella sezione di analisi del rischio dei singoli strumenti.""")
            run_3.font.name = 'Century Gothic'
            run_3.font.size = shared.Pt(10)
      
        # Alert mercati obbligazionari
        paragraph_4 = self.document.add_paragraph()
        paragraph_4.paragraph_format.space_after = shared.Pt(6)
        paragraph_4.paragraph_format.line_spacing_rule = 1

        dict_peso_micro_obbligazionarie = {item : dict_peso_micro[item]/dict_peso_macro[key] for key, value in self.dict_macro_micro.items() for item in value if key=='Obbligazionario'}
        dict_peso_micro_obbligazionarie['Obbligazionario High Yield'] = dict_peso_micro_obbligazionarie['Obbligazionario Euro High Yield'] + dict_peso_micro_obbligazionarie['Obbligazionario Globale High Yield']
        dict_peso_micro_obbligazionarie.pop('Obbligazionario Euro High Yield', None)
        dict_peso_micro_obbligazionarie.pop('Obbligazionario Globale High Yield', None)
        pesi_limite_obbligazionari = {'Obbligazionario Euro Corporate' : 0.40, 'Obbligazionario Paesi Emergenti' : 0.20, 'Obbligazionario High Yield' : 0.20}
        dict_alert_obbligazionari = {k : True if v >= pesi_limite_obbligazionari[k] else False for k, v in dict_peso_micro_obbligazionarie.items() if k in pesi_limite_obbligazionari.keys()}
        nome_mercati_obbligazionari = {'Obbligazionario Euro Corporate' : 'corporate europeo', 'Obbligazionario Paesi Emergenti' : 'emergente', 'Obbligazionario High Yield' : 'high yield'}
        
        if sum(dict_alert_obbligazionari.values()) == 1:
            run_4 = paragraph_4.add_run(f"""Per la parte investita nel mercato obbligazionario, si segnala l’eccessiva concentrazione verso il comparto {''.join([nome_mercati_obbligazionari[k] for k, v in dict_alert_obbligazionari.items() if v==True])}, che pesa {' '.join([str('il ' if str(dict_peso_micro_obbligazionarie[k]*100)[0].startswith(('1','2','3','4','5','6','7','9')) and str(dict_peso_micro_obbligazionarie[k]*100)[:2]!='11' else "l'")+str(round(dict_peso_micro_obbligazionarie[k]*100, 2)).replace('.', ',')+str('%') for k, v in dict_alert_obbligazionari.items() if v==True])} del comparto obbligazionario, come indicato dal relativo warning. """)
            run_4.font.name = 'Century Gothic'
            run_4.font.size = shared.Pt(10)
        elif sum(dict_alert_obbligazionari.values()) > 1:
            run_4 = paragraph_4.add_run(f"""Per la parte investita nel mercato obbligazionario, si segnala l’eccessiva concentrazione verso i comparti {' e '.join([nome_mercati_obbligazionari[k] for k, v in dict_alert_obbligazionari.items() if v==True])} che pesano rispettivamente {', e '.join([str('il ' if str(dict_peso_micro_obbligazionarie[k]*100)[0].startswith(('1','2','3','4','5','6','7','9')) and str(dict_peso_micro_obbligazionarie[k]*100)[:2]!='11' else "l'")+str(round(dict_peso_micro_obbligazionarie[k]*100, 2)).replace('.', ',')+str('%') for k, v in dict_alert_obbligazionari.items() if v==True])} del comparto obbligazionario, come indicato dai relativi warning. """)
            run_4.font.name = 'Century Gothic'
            run_4.font.size = shared.Pt(10)
        
        # Alert prodotti obbligazionari
        quote_prodotti_obbligazionari = df_portfolio.loc[(df_portfolio['strumento']=='gov_bond') | (df_portfolio['strumento']=='corp_bond'), ['nome', 'controvalore_in_euro']]
        quote_prodotti_obbligazionari['peso_totale'] = quote_prodotti_obbligazionari['controvalore_in_euro'] / df_portfolio['controvalore_in_euro'].sum()
        quote_prodotti_obbligazionari['peso_su_obbligazionario'] = quote_prodotti_obbligazionari['peso_totale'] / dict_peso_macro['Obbligazionario']
        if any(quote_prodotti_obbligazionari['peso_su_obbligazionario'] > 0.10):
            quote_prodotti_obbligazionari_alert = dict(zip(quote_prodotti_obbligazionari.loc[quote_prodotti_obbligazionari['peso_su_obbligazionario']>0.10, 'nome'].values, quote_prodotti_obbligazionari.loc[quote_prodotti_obbligazionari['peso_su_obbligazionario']>0.10, 'peso_su_obbligazionario']))
            run_4 = paragraph_4.add_run(f"""Riguardo agli strumenti obbligazionari, si segnala il peso eccessivo di {' e '.join([k for k,v in quote_prodotti_obbligazionari_alert.items()])} {'che pesa' if len(quote_prodotti_obbligazionari_alert)==1 else 'che pesano rispettivamente'} {', e '.join([str('il ' if str(v*100)[0].startswith(('1','2','3','4','5','6','7','9')) and str(v*100)[:2]!='11' else "l'")+str(round(v*100, 2)).replace('.', ',')+str('%') for k,v in quote_prodotti_obbligazionari_alert.items()])} dell’intero comparto obbligazionario, come indicato {'dal relativo warning' if len(quote_prodotti_obbligazionari_alert)==1 else 'dai relativi warning'} nella sezione di analisi del rischio dei singoli strumenti.""")
            run_4.font.name = 'Century Gothic'
            run_4.font.size = shared.Pt(10)

        # Alert valute
        dict_peso_valute = self.peso_valuta()
        if dict_peso_valute.get('EUR', None) < 0.40:
            paragraph_5 = self.document.add_paragraph()
            paragraph_5.paragraph_format.space_after = shared.Pt(6)
            paragraph_5.paragraph_format.line_spacing_rule = 1
            run_5 = paragraph_4.add_run(f"""Si segnala, infine, l’eccessiva esposizione del portafoglio a valute diverse dall’Euro, come indicato dal relativo warning.""")
            run_5.font.name = 'Century Gothic'
            run_5.font.size = shared.Pt(10)
        
    def analisi_di_portafoglio(self):
        """Incolla tabelle e grafici a torta."""
        self.document.add_section()
        paragraph_0 = self.document.add_paragraph(text='', style=None)
        paragraph_0.paragraph_format.space_before = shared.Pt(6)
        run_0 = paragraph_0.add_run(text='')
        run_0.add_picture(self.path_img_default.joinpath('2_analisi_del_portafoglio.bmp').__str__(), width=shared.Cm(8.5))
        table_0 = self.document.add_table(rows=9, cols=2)
        cell_1 = table_0.cell(0,0).merge(table_0.cell(0,1))
        paragraph_1 = cell_1.paragraphs[0]
        print('sto aggiungendo le macro categorie...')
        run_1_1 = paragraph_1.add_run('\n')
        run_1_1.font.size = shared.Pt(10)
        run_1_2 = paragraph_1.add_run('Analisi per Macro Asset Class')
        run_1_2.bold = True
        run_1_2.font.name = 'Century Gothic'
        run_1_2.font.size = shared.Pt(14)
        run_1_2.font.color.rgb = shared.RGBColor(127, 127, 127)
        cell_2 = table_0.cell(1,0).merge(table_0.cell(1,1))
        paragraph_2 = cell_2.paragraphs[0]
        paragraph_2.paragraph_format.line_spacing = shared.Cm(0.2)
        run_2 = paragraph_2.add_run()
        run_2.add_picture(self.path_img_default.joinpath('barra.png').__str__(), width=shared.Cm(self.larghezza_pagina+0.1))
        cell_3 = table_0.cell(2,0)
        paragraph_3 = cell_3.paragraphs[0]
        run_3 = paragraph_3.add_run()
        # # Librerie win32com + PIL
        # xls_file = win32com.client.Dispatch("Excel.Application")
        # # xls_file.Visible = False
        # # xls_file.ScreenUpdating = False
        # # xls_file.DisplayAlerts = False
        # # xls_file.EnableEvents = False
        # wb = xls_file.Workbooks.Open(Filename=self.path+"\\"+self.file_elaborato)
        # ws = wb.Worksheets("figure")
        # ws.Range(ws.Cells(1,1),ws.Cells(6,3)).CopyPicture(Format=2)
        # img = ImageGrab.grabclipboard()
        # img.save(self.path+r'\img\macro.png')
        # wb.Close(SaveChanges=False, Filename=self.path+"\\"+self.file_elaborato)
        # forse excel si aspetta un'azione come nascondere delle righe etc e non un continuo copiare range del file. prova a deselezionare il range precedente oppure selezionare il successivo
        # Librerie excel2img
        excel2img.export_img(self.file_elaborato, self.path_img.joinpath('macro.png').__str__(), page='figure', _range="A1:C6")
        run_3.add_picture(self.path_img.joinpath('macro.png').__str__(), width=shared.Cm(9.5))
        cell_4 = table_0.cell(2,1)
        paragraph_4 = cell_4.paragraphs[0]
        paragraph_4.paragraph_format.alignment = 2
        run_4 = paragraph_4.add_run()
        run_4.add_picture(self.path_img.joinpath('macro_pie.png').__str__(), height=shared.Cm(5), width=shared.Cm(5))
        cell_5 = table_0.cell(3,0).merge(table_0.cell(3,1))
        paragraph_5 = cell_5.paragraphs[0]
        run_5 = paragraph_5.add_run('\n')
        run_5.add_picture(self.path_img_default.joinpath('macro_info.bmp').__str__(), width=shared.Cm(self.larghezza_pagina))
        cell_6 = table_0.cell(4,0).merge(table_0.cell(4,1))
        paragraph_6 = cell_6.paragraphs[0]
        run_6 = paragraph_6.add_run('')
        run_6.font.size = shared.Pt(22)
        cell_7 = table_0.cell(5,0).merge(table_0.cell(5,1))
        paragraph_7 = cell_7.paragraphs[0]
        print('sto aggiungendo le micro categorie...')
        run_7 = paragraph_7.add_run('Analisi per Micro Asset Class')
        run_7.bold = True
        run_7.font.name = 'Century Gothic'
        run_7.font.size = shared.Pt(14)
        run_7.font.color.rgb = shared.RGBColor(127, 127, 127)
        cell_8 = table_0.cell(6,0).merge(table_0.cell(6,1))
        paragraph_8 = cell_8.paragraphs[0]
        paragraph_8.paragraph_format.line_spacing = shared.Cm(0.2)
        run_8 = paragraph_8.add_run()
        run_8.add_picture(self.path_img_default.joinpath('barra.png').__str__(), width=shared.Cm(self.larghezza_pagina+0.1))
        cell_9 = table_0.cell(7,0).merge(table_0.cell(7,1))
        paragraph_9 = cell_9.paragraphs[0]
        run_9 = paragraph_9.add_run()
        # # Librerie win32com + PIL
        # xls_file = win32com.client.Dispatch("Excel.Application")
        # # # xls_file.Visible = False
        # # # xls_file.ScreenUpdating = False
        # # # xls_file.DisplayAlerts = False
        # # # xls_file.EnableEvents = False
        # wb = xls_file.Workbooks.Open(Filename=self.path+"\\"+self.file_elaborato)
        # ws = wb.Worksheets("figure")
        # ws.Range(ws.Cells(1,9),ws.Cells(16,14)).CopyPicture(Format=2)
        # img = ImageGrab.grabclipboard()
        # img.save(self.path+r'\img\micro.png')
        # wb.Close(SaveChanges=False, Filename=self.path+"\\"+self.file_elaborato)
        # Librerie excel2img
        excel2img.export_img(self.file_elaborato, self.path_img.joinpath('micro.png').__str__(), page='figure', _range="E1:J16")
        run_9.add_picture(self.path_img.joinpath('micro.png').__str__(), height=shared.Cm(7), width=shared.Cm(self.larghezza_pagina))
        cell_10 = table_0.cell(8,0).merge(table_0.cell(8,1))
        paragraph_10 = cell_10.paragraphs[0]
        run_10 = paragraph_10.add_run()
        run_10.add_picture(self.path_img.joinpath('micro_bar.png').__str__(), height=shared.Cm(5), width=shared.Cm(self.larghezza_pagina))
        # Pagina nuova
        self.document.add_section()
        paragraph_10 = self.document.add_paragraph(text='', style=None)
        paragraph_10.paragraph_format.space_before = shared.Pt(6)
        run_10 = paragraph_10.add_run(text='')
        run_10.add_picture(self.path_img_default.joinpath('2_analisi_del_portafoglio.bmp').__str__(), width=shared.Cm(8.5))
        table_1 = self.document.add_table(rows=9, cols=2)
        cell_11 = table_1.cell(0,0).merge(table_1.cell(0,1))
        paragraph_11 = cell_11.paragraphs[0]
        print('sto aggiungendo gli strumenti...')
        run_11_1 = paragraph_11.add_run('\n')
        run_11_1.font.size = shared.Pt(10)
        run_11_2 = paragraph_11.add_run('Analisi per Strumenti')
        run_11_2.bold = True
        run_11_2.font.name = 'Century Gothic'
        run_11_2.font.size = shared.Pt(14)
        run_11_2.font.color.rgb = shared.RGBColor(127, 127, 127)
        cell_12 = table_1.cell(1,0).merge(table_1.cell(1,1))
        paragraph_12 = cell_12.paragraphs[0]
        paragraph_12.paragraph_format.line_spacing = shared.Cm(0.2)
        run_12 = paragraph_12.add_run()
        run_12.add_picture(self.path_img_default.joinpath('barra.png').__str__(), width=shared.Cm(self.larghezza_pagina+0.1))
        cell_13 = table_1.cell(2,0)
        paragraph_13 = cell_13.paragraphs[0]
        run_13 = paragraph_13.add_run()
        # # Librerie win32com + PIL
        # xls_file = win32com.client.Dispatch("Excel.Application")
        # # # xls_file.Visible = False
        # # # xls_file.ScreenUpdating = False
        # # # xls_file.DisplayAlerts = False
        # # # xls_file.EnableEvents = False
        # wb = xls_file.Workbooks.Open(Filename=self.path+"\\"+self.file_elaborato)
        # ws = wb.Worksheets("figure")
        # ws.Range(ws.Cells(18,1),ws.Cells(34,4)).CopyPicture(Format=2)
        # img = ImageGrab.grabclipboard()
        # img.save(self.path+r'\img\strumenti.png')
        # wb.Close(SaveChanges=False, Filename=self.path+"\\"+self.file_elaborato)
        # Librerie excel2img
        excel2img.export_img(self.file_elaborato, self.path_img.joinpath('strumenti.png').__str__(), page='figure', _range="L1:O17")
        run_13.add_picture(self.path_img.joinpath('strumenti.png').__str__(), width=shared.Cm(10.5))
        cell_14 = table_1.cell(2,1)
        paragraph_14 = cell_14.paragraphs[0]
        paragraph_14.paragraph_format.alignment = 2
        run_14 = paragraph_14.add_run()
        run_14.add_picture(self.path_img.joinpath('strumenti_pie.png').__str__(), height=shared.Cm(5), width=shared.Cm(5))
        cell_15 = table_1.cell(5,0).merge(table_1.cell(5,1))
        paragraph_15 = cell_15.paragraphs[0]
        print('sto aggiungendo le valute...')
        run_15 = paragraph_15.add_run('\n\n\nAnalisi per Valute')
        run_15.bold = True
        run_15.font.name = 'Century Gothic'
        run_15.font.size = shared.Pt(14)
        run_15.font.color.rgb = shared.RGBColor(127, 127, 127)
        cell_16 = table_1.cell(6,0).merge(table_1.cell(6,1))
        paragraph_16 = cell_16.paragraphs[0]
        paragraph_16.paragraph_format.line_spacing = shared.Cm(0.2)
        run_16 = paragraph_16.add_run()
        run_16.add_picture(self.path_img_default.joinpath('barra.png').__str__(), width=shared.Cm(self.larghezza_pagina+0.1))
        cell_17 = table_1.cell(7,0)
        paragraph_17 = cell_17.paragraphs[0]
        run_17 = paragraph_17.add_run()
        # # Librerie win32com + PIL
        # xls_file = win32com.client.Dispatch("Excel.Application")
        # # # xls_file.Visible = False
        # # # xls_file.ScreenUpdating = False
        # # # xls_file.DisplayAlerts = False
        # # # xls_file.EnableEvents = False
        # wb = xls_file.Workbooks.Open(Filename=self.path+"\\"+self.file_elaborato)
        # ws = wb.Worksheets("figure")
        # ws.Range(ws.Cells(1,16),ws.Cells(9,19)).CopyPicture(Format=2)
        # img = ImageGrab.grabclipboard()
        # img.save(self.path+r'\img\valute.png')
        # wb.Close(SaveChanges=False, Filename=self.path+"\\"+self.file_elaborato)
        # Librerie excel2img
        excel2img.export_img(self.file_elaborato, self.path_img.joinpath('valute.png').__str__(), page='figure', _range="Q1:T9")
        run_17.add_picture(self.path_img.joinpath('valute.png').__str__(), height=shared.Cm(3.7), width=shared.Cm(5))
        cell_18 = table_1.cell(7,1)
        paragraph_18 = cell_18.paragraphs[0]
        paragraph_18.paragraph_format.alignment = 2
        run_18 = paragraph_18.add_run()
        run_18.add_picture(self.path_img.joinpath('valute_pie.png').__str__(), height=shared.Cm(5), width=shared.Cm(5))
        cell_19 = table_1.cell(8,0).merge(table_1.cell(8,1))
        paragraph_19 = cell_19.paragraphs[0]
        run_19 = paragraph_19.add_run()
        run_19.add_picture(self.path_img_default.joinpath('valute_info_new.bmp').__str__(), width=shared.Cm(self.larghezza_pagina))
        # Pagina nuova
        self.document.add_section()
        paragraph_20 = self.document.add_paragraph(text='', style=None)
        paragraph_20.paragraph_format.space_before = shared.Pt(6)
        run_20 = paragraph_20.add_run(text='')
        run_20.add_picture(self.path_img_default.joinpath('2_analisi_del_portafoglio.bmp').__str__(), width=shared.Cm(8.5))
        table_2 = self.document.add_table(rows=9, cols=2)
        cell_21 = table_2.cell(0,0).merge(table_2.cell(0,1))
        paragraph_21 = cell_21.paragraphs[0]
        print('sto aggiungendo il risparmio...')
        run_21_1 = paragraph_21.add_run('\n')
        run_21_1.font.size = shared.Pt(10)
        run_21_2 = paragraph_21.add_run('Analisi per Risparmio')
        run_21_2.bold = True
        run_21_2.font.name = 'Century Gothic'
        run_21_2.font.size = shared.Pt(14)
        run_21_2.font.color.rgb = shared.RGBColor(127, 127, 127)
        cell_22 = table_2.cell(1,0).merge(table_2.cell(1,1))
        paragraph_22 = cell_22.paragraphs[0]
        paragraph_22.paragraph_format.line_spacing = shared.Cm(0.2)
        run_22 = paragraph_22.add_run()
        run_22.add_picture(self.path_img_default.joinpath('barra.png').__str__(), width=shared.Cm(self.larghezza_pagina+0.1))
        cell_23 = table_2.cell(2,0)
        paragraph_23 = cell_23.paragraphs[0]
        run_23 = paragraph_23.add_run()
        # # Librerie win32com + PIL
        # xls_file = win32com.client.Dispatch("Excel.Application")
        # # # xls_file.Visible = False
        # # # xls_file.ScreenUpdating = False
        # # # xls_file.DisplayAlerts = False
        # # # xls_file.EnableEvents = False
        # wb = xls_file.Workbooks.Open(Filename=self.path+"\\"+self.file_elaborato)
        # ws = wb.Worksheets("figure")
        # ws.Range(ws.Cells(18,1),ws.Cells(34,4)).CopyPicture(Format=2)
        # img = ImageGrab.grabclipboard()
        # img.save(self.path+r'\img\strumenti.png')
        # wb.Close(SaveChanges=False, Filename=self.path+"\\"+self.file_elaborato)
        # Librerie excel2img
        excel2img.export_img(self.file_elaborato, self.path_img.joinpath('risparmio.png').__str__(), page='figure', _range="V1:X17")
        run_23.add_picture(self.path_img.joinpath('risparmio.png').__str__(), width=shared.Cm(10.5))
        cell_24 = table_2.cell(2,1)
        paragraph_24 = cell_24.paragraphs[0]
        paragraph_24.paragraph_format.alignment = 2
        run_24 = paragraph_24.add_run()
        run_24.add_picture(self.path_img.joinpath('risparmio_pie.png').__str__(), height=shared.Cm(5), width=shared.Cm(5))
        
    def analisi_strumenti(self):
        """
        Incolla tabelle di obbligazioni e azioni.
        """
        # Obbligazioni #
        df_portfolio = self.df_portfolio
        prodotti_obbligazionari = df_portfolio.loc[(df_portfolio['strumento']=='gov_bond') | (df_portfolio['strumento']=='corp_bond')]
        numero_prodotti_obbligazionari = prodotti_obbligazionari.nome.count()
        # print('numero titoli obbligazionari:',numero_prodotti_obbligazionari)
        MAX_OBB_DES_PER_PAGINA = 52 # 52
        MAX_OBB_DATI_PER_PAGINA = 47 # 47
        MAX_AZIONI_PER_PAGINA = 62 # 62
        MAX_FONDI_PER_PAGINA = 53 # 53
        MAX_MAP_FONDI_PER_PAGINA = 90 # 90
        if numero_prodotti_obbligazionari > 0:
            # Carica il foglio obbligazioni
            obbligazioni = self.wb['obbligazioni']
            self.wb.active = obbligazioni
            # Nascondi le colonne del prezzo di carico e della variazione di prezzo
            hidden_columns = 0
            if all(df_portfolio['prezzo_di_carico'].isnull()):
                obbligazioni.column_dimensions['L'].hidden= True
                hidden_columns += 1
                obbligazioni.column_dimensions['O'].hidden= True
                hidden_columns += 1
            self.wb.save(self.file_elaborato)
            
            # Descrizione obbligazioni
            if numero_prodotti_obbligazionari > MAX_OBB_DES_PER_PAGINA and numero_prodotti_obbligazionari % MAX_OBB_DES_PER_PAGINA != 0:
                tabelle_des = int(numero_prodotti_obbligazionari // MAX_OBB_DES_PER_PAGINA + 1)
            elif numero_prodotti_obbligazionari > MAX_OBB_DES_PER_PAGINA and numero_prodotti_obbligazionari % MAX_OBB_DES_PER_PAGINA == 0:
                tabelle_des = int(numero_prodotti_obbligazionari // MAX_OBB_DES_PER_PAGINA)
            else:
                tabelle_des = 1
            # print('tabelle_des:',tabelle_des)

            # # Librerie win32com + PIL
            # xls_file = win32com.client.Dispatch("Excel.Application")
            # wb = xls_file.Workbooks.Open(Filename=self.path+"\\"+self.file_elaborato)
            # # xls_file.Visible = False
            # # xls_file.ScreenUpdating = False
            # # xls_file.DisplayAlerts = False
            # # xls_file.EnableEvents = False
            # ws = wb.Worksheets("obbligazioni")
            for tabella in range(1, tabelle_des+1):
                print(f'sto aggiungendo la tabella descrizione obbligazioni: {tabella} / {tabelle_des}')
                # print(tabella)
                if tabella != tabelle_des:
                    # ws.Range(ws.Cells(1,2),ws.Cells(MAX_OBB_DES_PER_PAGINA*tabella+1,8)).CopyPicture(Format=2)            
                    # img = ImageGrab.grabclipboard()
                    # img.save(self.path+r'\img\obbligazioni_des_' + str(tabella) + '.png')
                    # ws.Range(ws.Cells(2+MAX_OBB_DES_PER_PAGINA*(tabella-1),2),ws.Cells(MAX_OBB_DES_PER_PAGINA*tabella+1,8)).Rows.EntireRow.Hidden = True
                    # Libreria excel2img
                    excel2img.export_img(self.file_elaborato, self.path_img.joinpath('obbligazioni_des_' + str(tabella) + '.png').__str__(), page='obbligazioni', _range="B1:H"+str(MAX_OBB_DES_PER_PAGINA*tabella+1))
                    obbligazioni.row_dimensions.group(2+MAX_OBB_DES_PER_PAGINA*(tabella-1),MAX_OBB_DES_PER_PAGINA*tabella+1,hidden=True)
                    self.wb.save(self.file_elaborato)
                else:
                    # ws.Range(ws.Cells(1,2),ws.Cells(prodotti_obbligazionari.nome.count()+1,8)).CopyPicture(Format=2)            
                    # img = ImageGrab.grabclipboard()
                    # img.save(self.path+r'\img\obbligazioni_des_' + str(tabella) + '.png')
                    # Libreria excel2img
                    excel2img.export_img(self.file_elaborato, self.path_img.joinpath('obbligazioni_des_' + str(tabella) + '.png').__str__(), page='obbligazioni', _range="B1:H"+str(prodotti_obbligazionari.nome.count()+1))         
            obbligazioni.row_dimensions.group(1,MAX_OBB_DES_PER_PAGINA*tabelle_des,hidden=False)
            self.wb.save(self.file_elaborato)
            # wb.Close(SaveChanges=False, Filename=self.path+"\\"+self.file_elaborato)

            for tabella in range(1, tabelle_des+1):
                self.document.add_section()
                paragraph = self.document.add_paragraph(text='', style=None)
                paragraph.paragraph_format.space_before = shared.Pt(6)
                paragraph.paragraph_format.space_after = shared.Pt(0)
                run = paragraph.add_run(text='')
                run.add_picture(self.path_img_default.joinpath('3_analisi_dei_singoli_strumenti.bmp').__str__(), width=shared.Cm(8.5))
                paragraph = self.document.add_paragraph(text='', style=None)
                run_0 = paragraph.add_run('\n')
                run_0.font.size = shared.Pt(7)
                run = paragraph.add_run('Caratteristiche anagrafiche dei titoli obbligazionari')
                run.bold = True
                run.font.name = 'Century Gothic'
                run.font.size = shared.Pt(14)
                run.font.color.rgb = shared.RGBColor(127, 127, 127)
                paragraph = self.document.add_paragraph(text='', style=None)
                paragraph.paragraph_format.line_spacing = shared.Cm(0.2)
                run = paragraph.add_run()
                run.add_picture(self.path_img_default.joinpath('barra.png').__str__(), width=shared.Cm(self.larghezza_pagina+0.1))
                paragraph = self.document.add_paragraph(text='', style=None)
                run = paragraph.add_run()
                run.add_picture(self.path_img.joinpath('obbligazioni_des_'+str(tabella)+'.png').__str__(), width=shared.Cm(self.larghezza_pagina))

            # Dati obbligazioni
            # Calcolo numero titoli nell'ultima tabella
            if numero_prodotti_obbligazionari > MAX_OBB_DES_PER_PAGINA:
                if numero_prodotti_obbligazionari % MAX_OBB_DES_PER_PAGINA == 0:
                    num_obb_des_ultima_pagina = MAX_OBB_DES_PER_PAGINA
                elif numero_prodotti_obbligazionari % MAX_OBB_DES_PER_PAGINA != 0:
                    num_obb_des_ultima_pagina = numero_prodotti_obbligazionari % MAX_OBB_DES_PER_PAGINA
            elif numero_prodotti_obbligazionari <= MAX_OBB_DES_PER_PAGINA:
                num_obb_des_ultima_pagina = numero_prodotti_obbligazionari
            # print("prodotti nell'ultima pagina:",num_obb_des_ultima_pagina)
            # Calcolo numero titoli nell'eventuale tabella sotto l'ultima
            if MAX_OBB_DATI_PER_PAGINA - int(num_obb_des_ultima_pagina*MAX_OBB_DATI_PER_PAGINA/MAX_OBB_DES_PER_PAGINA) - 9 > 0: # se rimane spazio sufficiente sotto l'ultima tabella precedente
                if (MAX_OBB_DATI_PER_PAGINA - int(num_obb_des_ultima_pagina*MAX_OBB_DATI_PER_PAGINA/MAX_OBB_DES_PER_PAGINA) - 9) < numero_prodotti_obbligazionari:
                    numerosita_tabella_obb_dati_sotto_la_precedente = MAX_OBB_DATI_PER_PAGINA - int(num_obb_des_ultima_pagina*MAX_OBB_DATI_PER_PAGINA/MAX_OBB_DES_PER_PAGINA) - 9
                else: # se tutte le obbligazioni ci stanno in quello spazio rimasto
                    numerosita_tabella_obb_dati_sotto_la_precedente = numero_prodotti_obbligazionari
            else: # se non rimane spazio a sufficienza per una tabella sotto la precedente
                numerosita_tabella_obb_dati_sotto_la_precedente = 0
            # print("numerosità tabella obb dati sotto la precedente:",numerosita_tabella_obb_dati_sotto_la_precedente)
            
            # Inserisci l'eventuale tabella sotto l'ultima
            # Librerie win32com + PIL
            # xls_file = win32com.client.Dispatch("Excel.Application")
            # wb = xls_file.Workbooks.Open(Filename=self.path+"\\"+self.file_elaborato)
            # # xls_file.Visible = False
            # # xls_file.ScreenUpdating = False
            # # xls_file.DisplayAlerts = False
            # # xls_file.EnableEvents = False
            # ws = wb.Worksheets("obbligazioni")
            if numerosita_tabella_obb_dati_sotto_la_precedente > 0:
                # Prima tabella dati obbligazioni
                # ws.Range(ws.Cells(1,10),ws.Cells(numerosita_tabella_obb_dati_sotto_la_precedente+1,17)).CopyPicture(Format=2)
                # img = ImageGrab.grabclipboard()
                # img.save(self.path+r'\img\obbligazioni_dati_0.png')
                # ws.Range(ws.Cells(2,10),ws.Cells(numerosita_tabella_obb_dati_sotto_la_precedente+1,17)).Rows.EntireRow.Hidden = True
                # Libreria excel2img
                excel2img.export_img(self.file_elaborato, self.path_img.joinpath('obbligazioni_dati_0.png').__str__(), page='obbligazioni', _range="J1:Q"+str(numerosita_tabella_obb_dati_sotto_la_precedente+1))
                obbligazioni.row_dimensions.group(2,numerosita_tabella_obb_dati_sotto_la_precedente+1,hidden=True)
                self.wb.save(self.file_elaborato)
                # print(0)
                paragraph = self.document.add_paragraph(text='', style=None)
                run = paragraph.add_run('\nCaratteristiche finanziarie dei titoli obbligazionari')
                run.bold = True
                run.font.name = 'Century Gothic'
                run.font.size = shared.Pt(14)
                run.font.color.rgb = shared.RGBColor(127, 127, 127)
                paragraph = self.document.add_paragraph(text='', style=None)
                paragraph.paragraph_format.line_spacing = shared.Cm(0.2)
                run = paragraph.add_run()
                run.add_picture(self.path_img_default.joinpath('barra.png').__str__(), width=shared.Cm(self.larghezza_pagina+0.1))
                paragraph = self.document.add_paragraph(text='', style=None)
                run = paragraph.add_run()
                run.add_picture(self.path_img.joinpath('obbligazioni_dati_0.png').__str__(), width=shared.Cm(self.larghezza_pagina) if hidden_columns==0 else shared.Cm(self.larghezza_pagina-4))
            
            # Inserisci le tabelle rimanenti
            if numero_prodotti_obbligazionari - numerosita_tabella_obb_dati_sotto_la_precedente == 0: # tutti i titoli sono contenuti nella tabella sotto l'ultima
                tabelle_dati = 1
                # print('tabelle_dati:',tabelle_dati)
            else:
                if numero_prodotti_obbligazionari - numerosita_tabella_obb_dati_sotto_la_precedente <= MAX_OBB_DATI_PER_PAGINA:
                    tabelle_dati = 1
                    # numerosita_tabella_obb_dati_sotto_la_precedente = numero_prodotti_obbligazionari
                elif numero_prodotti_obbligazionari - numerosita_tabella_obb_dati_sotto_la_precedente > MAX_OBB_DATI_PER_PAGINA:
                    if (numero_prodotti_obbligazionari - numerosita_tabella_obb_dati_sotto_la_precedente) % MAX_OBB_DATI_PER_PAGINA != 0:
                        tabelle_dati = int((numero_prodotti_obbligazionari - numerosita_tabella_obb_dati_sotto_la_precedente) // MAX_OBB_DATI_PER_PAGINA) + 1
                    else:
                        tabelle_dati = int((numero_prodotti_obbligazionari - numerosita_tabella_obb_dati_sotto_la_precedente) // MAX_OBB_DATI_PER_PAGINA)
                    # numerosita_tabella_obb_dati_sotto_la_precedente = MAX_OBB_DATI_PER_PAGINA
                # print('tabelle_dati:',tabelle_dati if numerosita_tabella_obb_dati_sotto_la_precedente == 0 else tabelle_dati+1)
                print(f'sto aggiungendo la tabella dati obbligazioni sotto la precedente: 0 / {tabelle_dati}') if numerosita_tabella_obb_dati_sotto_la_precedente > 0 else None
                for tabella in range(1, tabelle_dati+1):
                    print(f'sto aggiungendo la tabella dati obbligazioni: {tabella} / {tabelle_dati}')
                    # print(tabella)
                    if tabella != tabelle_dati:
                        # ws.Range(ws.Cells(1,10),ws.Cells(numerosita_tabella_obb_dati_sotto_la_precedente+MAX_OBB_DATI_PER_PAGINA*tabella+1,17)).CopyPicture(Format=2)
                        # img = ImageGrab.grabclipboard()
                        # img.save(self.path+'\img\obbligazioni_dati_' + str(tabella) + '.png')
                        # ws.Range(ws.Cells(2+MAX_OBB_DATI_PER_PAGINA*(tabella-1),10),ws.Cells(numerosita_tabella_obb_dati_sotto_la_precedente+MAX_OBB_DATI_PER_PAGINA*tabella+1,17)).Rows.EntireRow.Hidden = True
                        # Libreria excel2img
                        excel2img.export_img(self.file_elaborato, self.path_img.joinpath('obbligazioni_dati_' + str(tabella) + '.png').__str__(), page='obbligazioni', _range="J1:Q"+str(numerosita_tabella_obb_dati_sotto_la_precedente+MAX_OBB_DATI_PER_PAGINA*tabella+1))
                        obbligazioni.row_dimensions.group(2+MAX_OBB_DATI_PER_PAGINA*(tabella-1),numerosita_tabella_obb_dati_sotto_la_precedente+MAX_OBB_DATI_PER_PAGINA*tabella+1,hidden=True)
                        self.wb.save(self.file_elaborato)
                    else:
                        # ws.Range(ws.Cells(1,10),ws.Cells(numero_prodotti_obbligazionari+1,17)).CopyPicture(Format=2)
                        # img = ImageGrab.grabclipboard()
                        # img.save(self.path+'\img\obbligazioni_dati_' + str(tabella) + '.png')
                        # Libreria excel2img
                        excel2img.export_img(self.file_elaborato, self.path_img.joinpath('obbligazioni_dati_' + str(tabella) + '.png').__str__(), page='obbligazioni', _range="J1:Q"+str(numero_prodotti_obbligazionari+1))
                obbligazioni.row_dimensions.group(1,MAX_OBB_DATI_PER_PAGINA*(tabelle_dati+1),hidden=False)
                self.wb.save(self.file_elaborato)

                for tabella in range(1, tabelle_dati+1):
                    self.document.add_section()
                    paragraph = self.document.add_paragraph(text='', style=None)
                    paragraph.paragraph_format.space_before = shared.Pt(6)
                    paragraph.paragraph_format.space_after = shared.Pt(0)
                    run = paragraph.add_run(text='')
                    run.add_picture(self.path_img_default.joinpath('3_analisi_dei_singoli_strumenti.bmp').__str__(), width=shared.Cm(8.5))
                    paragraph = self.document.add_paragraph(text='', style=None)
                    run_0 = paragraph.add_run('\n')
                    run_0.font.size = shared.Pt(7)
                    run = paragraph.add_run('Caratteristiche finanziarie dei titoli obbligazionari')
                    run.bold = True
                    run.font.name = 'Century Gothic'
                    run.font.size = shared.Pt(14)
                    run.font.color.rgb = shared.RGBColor(127, 127, 127)
                    paragraph = self.document.add_paragraph(text='', style=None)
                    paragraph.paragraph_format.line_spacing = shared.Cm(0.2)
                    run = paragraph.add_run()
                    run.add_picture(self.path_img_default.joinpath('barra.png').__str__(), width=shared.Cm(self.larghezza_pagina+0.1))
                    paragraph = self.document.add_paragraph(text='', style=None)
                    run = paragraph.add_run()
                    run.add_picture(self.path_img.joinpath('obbligazioni_dati_'+str(tabella)+'.png').__str__(), width=shared.Cm(self.larghezza_pagina) if hidden_columns==0 else shared.Cm(self.larghezza_pagina-4))
            # wb.Close(SaveChanges=False, Filename=self.path+"\\"+self.file_elaborato)
        
        elif numero_prodotti_obbligazionari == 0:
            # tabelle_dati = 0
            numerosita_tabella_obb_dati_sotto_la_precedente = 0


        # Azioni
        prodotti_azionari = df_portfolio.loc[df_portfolio['strumento']=='equity']
        numero_prodotti_azionari = prodotti_azionari.nome.count()
        # print('numero titoli azionari:',numero_prodotti_azionari)
        if numero_prodotti_azionari > 0:
            # Carica il foglio azioni
            azioni = self.wb['azioni']
            self.wb.active = azioni
            # Nascondi le colonne del prezzo di carico e della variazione di prezzo
            hidden_columns = 0
            if all(df_portfolio['prezzo_di_carico'].isnull()):
                azioni.column_dimensions['F'].hidden= True
                hidden_columns += 1
                azioni.column_dimensions['H'].hidden= True
                hidden_columns += 1
                azioni.column_dimensions['I'].hidden= True
                hidden_columns += 1
            self.wb.save(self.file_elaborato)

            # Calcola il numero di prodotti nell'ultima pagina
            if numero_prodotti_obbligazionari > MAX_OBB_DES_PER_PAGINA:
                num_prodotti_ultima_pagina = (numero_prodotti_obbligazionari - numerosita_tabella_obb_dati_sotto_la_precedente) % MAX_OBB_DATI_PER_PAGINA
            elif numero_prodotti_obbligazionari <= MAX_OBB_DES_PER_PAGINA:
                if numerosita_tabella_obb_dati_sotto_la_precedente == 0:
                    if numero_prodotti_obbligazionari <= MAX_OBB_DATI_PER_PAGINA:
                        num_prodotti_ultima_pagina = numero_prodotti_obbligazionari
                    elif numero_prodotti_obbligazionari > MAX_OBB_DATI_PER_PAGINA:
                        num_prodotti_ultima_pagina = (numero_prodotti_obbligazionari - numerosita_tabella_obb_dati_sotto_la_precedente) % MAX_OBB_DATI_PER_PAGINA
                elif numerosita_tabella_obb_dati_sotto_la_precedente > 0:
                    if numerosita_tabella_obb_dati_sotto_la_precedente == numero_prodotti_obbligazionari:
                        num_prodotti_ultima_pagina = numero_prodotti_obbligazionari + numero_prodotti_obbligazionari
                    elif numerosita_tabella_obb_dati_sotto_la_precedente < numero_prodotti_obbligazionari:
                        num_prodotti_ultima_pagina = (numero_prodotti_obbligazionari - numerosita_tabella_obb_dati_sotto_la_precedente) % MAX_OBB_DATI_PER_PAGINA
            elif numero_prodotti_obbligazionari == 0:
                num_prodotti_ultima_pagina = 0
            # print("prodotti nell'ultima pagina:",num_prodotti_ultima_pagina)

            # Calcolo il numero titoli nell'eventuale tabella sotto l'ultima
            if num_prodotti_ultima_pagina == 0: # se non ci sono obbligazioni
                numerosita_tabella_azioni_sotto_la_precedente = 0
            elif num_prodotti_ultima_pagina == (numero_prodotti_obbligazionari * 2): # se le tabelle des e dati sono sulla stessa pagina
                if MAX_AZIONI_PER_PAGINA - int(numero_prodotti_obbligazionari*MAX_AZIONI_PER_PAGINA*(MAX_OBB_DES_PER_PAGINA+MAX_OBB_DATI_PER_PAGINA)/(MAX_OBB_DES_PER_PAGINA*MAX_OBB_DATI_PER_PAGINA)) - 22 > 0: # se rimane spazio sufficiente sotto le due tabelle precedenti
                    if (MAX_AZIONI_PER_PAGINA - int(numero_prodotti_obbligazionari*MAX_AZIONI_PER_PAGINA*(MAX_OBB_DES_PER_PAGINA+MAX_OBB_DATI_PER_PAGINA)/(MAX_OBB_DES_PER_PAGINA*MAX_OBB_DATI_PER_PAGINA)) - 22) < numero_prodotti_azionari: # ma non ce nè abbastanza per tutte le azioni
                        numerosita_tabella_azioni_sotto_la_precedente = MAX_AZIONI_PER_PAGINA - int(numero_prodotti_obbligazionari*MAX_AZIONI_PER_PAGINA*(MAX_OBB_DES_PER_PAGINA+MAX_OBB_DATI_PER_PAGINA)/(MAX_OBB_DES_PER_PAGINA*MAX_OBB_DATI_PER_PAGINA)) - 22
                    else: # e tutte le azioni ci stanno in quello spazio rimasto
                        numerosita_tabella_azioni_sotto_la_precedente = numero_prodotti_azionari
                else: # se non rimane spazio a sufficienza per una tabella sotto la precedente
                    numerosita_tabella_azioni_sotto_la_precedente = 0
            else: # ci sono obbligazioni ma le tabelle des e dati non sono sulla stessa pagina
                if MAX_AZIONI_PER_PAGINA - int(num_prodotti_ultima_pagina*MAX_AZIONI_PER_PAGINA/MAX_OBB_DATI_PER_PAGINA) - 11 > 0: # se rimane spazio sufficiente sotto la tabella precedente 
                    if MAX_AZIONI_PER_PAGINA - int(num_prodotti_ultima_pagina*MAX_AZIONI_PER_PAGINA/MAX_OBB_DATI_PER_PAGINA) - 11 < numero_prodotti_azionari: # ma non ce nè abbastanza per tutte le azioni
                        numerosita_tabella_azioni_sotto_la_precedente = MAX_AZIONI_PER_PAGINA - int(num_prodotti_ultima_pagina*MAX_AZIONI_PER_PAGINA/MAX_OBB_DATI_PER_PAGINA) - 11
                    else: # e tutte le azioni ci stanno in quello spazio rimasto
                        numerosita_tabella_azioni_sotto_la_precedente = numero_prodotti_azionari
                else: # se non rimane spazio a sufficienza per una tabella sotto la precedente
                    numerosita_tabella_azioni_sotto_la_precedente = 0
            # print("numerosità tabella azioni sotto la precedente:",numerosita_tabella_azioni_sotto_la_precedente)
            
            # Inserisci l'eventuale tabella sotto l'ultima
            # # Librerie win32com + PIL
            # xls_file = win32com.client.Dispatch("Excel.Application")
            # wb = xls_file.Workbooks.Open(Filename=self.path+"\\"+self.file_elaborato)
            # # xls_file.Visible = False
            # # xls_file.ScreenUpdating = False
            # # xls_file.DisplayAlerts = False
            # # xls_file.EnableEvents = False
            # ws = wb.Worksheets("azioni")
            if numerosita_tabella_azioni_sotto_la_precedente > 0:
                # Prima tabella dati azioni
                # ws.Range(ws.Cells(1,2),ws.Cells(numerosita_tabella_azioni_sotto_la_precedente+1,11)).CopyPicture(Format=2)
                # img = ImageGrab.grabclipboard()
                # img.save(self.path+r'\img\azioni_0.png')
                # ws.Range(ws.Cells(2,2),ws.Cells(numerosita_tabella_azioni_sotto_la_precedente+1,11)).Rows.EntireRow.Hidden = True
                # Libreria excel2img
                excel2img.export_img(self.file_elaborato, self.path_img.joinpath('azioni_0.png').__str__(), page='azioni', _range="B1:K"+str(numerosita_tabella_azioni_sotto_la_precedente+1))
                azioni.row_dimensions.group(2,numerosita_tabella_azioni_sotto_la_precedente+1,hidden=True)
                self.wb.save(self.file_elaborato)
                # print(0)
                paragraph = self.document.add_paragraph(text='', style=None)
                run = paragraph.add_run('\nCaratteristiche dei titoli azionari')
                run.bold = True
                run.font.name = 'Century Gothic'
                run.font.size = shared.Pt(14)
                run.font.color.rgb = shared.RGBColor(127, 127, 127)
                paragraph = self.document.add_paragraph(text='', style=None)
                paragraph.paragraph_format.line_spacing = shared.Cm(0.2)
                run = paragraph.add_run()
                run.add_picture(self.path_img_default.joinpath('barra.png').__str__(), width=shared.Cm(self.larghezza_pagina+0.1))
                paragraph = self.document.add_paragraph(text='', style=None)
                run = paragraph.add_run()
                run.add_picture(self.path_img.joinpath('azioni_0.png').__str__(), width=shared.Cm(self.larghezza_pagina) if hidden_columns==0 else shared.Cm(self.larghezza_pagina-3.5))
        
            # Inserisci le tabelle rimanenti
            if numero_prodotti_azionari - numerosita_tabella_azioni_sotto_la_precedente == 0: # tutti i titoli sono contenuti nella tabella sotto l'ultima
                tabelle_azioni = 1
                # print('tabelle_azioni:',tabelle_azioni)
            else:
                if numero_prodotti_azionari - numerosita_tabella_azioni_sotto_la_precedente <= MAX_AZIONI_PER_PAGINA:
                    tabelle_azioni = 1
                    # numerosita_tabella_azioni_sotto_la_precedente = numero_prodotti_azionari
                elif numero_prodotti_azionari - numerosita_tabella_azioni_sotto_la_precedente > MAX_AZIONI_PER_PAGINA:
                    if (numero_prodotti_azionari - numerosita_tabella_azioni_sotto_la_precedente) % MAX_AZIONI_PER_PAGINA != 0:
                        tabelle_azioni = int((numero_prodotti_azionari - numerosita_tabella_azioni_sotto_la_precedente) // MAX_AZIONI_PER_PAGINA) + 1
                    else:
                        tabelle_azioni = int((numero_prodotti_azionari - numerosita_tabella_azioni_sotto_la_precedente) // MAX_AZIONI_PER_PAGINA)
                    # numerosita_tabella_azioni_sotto_la_precedente = MAX_AZIONI_PER_PAGINA
                # print('tabelle_azioni:',tabelle_azioni if numerosita_tabella_azioni_sotto_la_precedente == 0 else tabelle_azioni+1)
                print(f'sto aggiungendo la tabella azioni sotto la precedente: 0 / {tabelle_azioni}') if numerosita_tabella_azioni_sotto_la_precedente > 0 else None

                for tabella in range(1, tabelle_azioni+1):
                    print(f'sto aggiungendo la tabella azioni: {tabella} / {tabelle_azioni}')
                    # print(tabella)
                    if tabella != tabelle_azioni:
                        # ws.Range(ws.Cells(1,2),ws.Cells(numerosita_tabella_azioni_sotto_la_precedente+MAX_AZIONI_PER_PAGINA*tabella+1,11)).CopyPicture(Format=2)
                        # img = ImageGrab.grabclipboard()
                        # img.save(self.path+r'\img\azioni_' + str(tabella) + '.png')
                        # ws.Range(ws.Cells(2+MAX_AZIONI_PER_PAGINA*(tabella-1),2),ws.Cells(numerosita_tabella_azioni_sotto_la_precedente+MAX_AZIONI_PER_PAGINA*tabella+1,11)).Rows.EntireRow.Hidden = True
                        # Libreria excel2img
                        excel2img.export_img(self.file_elaborato, self.path_img.joinpath('azioni_' + str(tabella) + '.png').__str__(), page='azioni', _range="B1:K"+str(numerosita_tabella_azioni_sotto_la_precedente+MAX_AZIONI_PER_PAGINA*tabella+1))
                        azioni.row_dimensions.group(2+MAX_AZIONI_PER_PAGINA*(tabella-1),numerosita_tabella_azioni_sotto_la_precedente+MAX_AZIONI_PER_PAGINA*tabella+1,hidden=True)
                        self.wb.save(self.file_elaborato)
                    else:
                        # ws.Range(ws.Cells(1,2),ws.Cells(numero_prodotti_azionari+1,11)).CopyPicture(Format=2)
                        # img = ImageGrab.grabclipboard()
                        # img.save(self.path+r'\img\azioni_' + str(tabella) + '.png')
                        # Libreria excel2img
                        excel2img.export_img(self.file_elaborato, self.path_img.joinpath('azioni_' + str(tabella) + '.png').__str__(), page='azioni', _range="B1:K"+str(numero_prodotti_azionari+1))
                azioni.row_dimensions.group(1,MAX_AZIONI_PER_PAGINA*(tabelle_azioni+1),hidden=False)
                self.wb.save(self.file_elaborato)

                for tabella in range(1, tabelle_azioni+1):
                    self.document.add_section()
                    paragraph = self.document.add_paragraph(text='', style=None)
                    paragraph.paragraph_format.space_before = shared.Pt(6)
                    paragraph.paragraph_format.space_after = shared.Pt(0)
                    run = paragraph.add_run(text='')
                    run.add_picture(self.path_img_default.joinpath('3_analisi_dei_singoli_strumenti.bmp').__str__(), width=shared.Cm(8.5))
                    paragraph = self.document.add_paragraph(text='', style=None)
                    run_0 = paragraph.add_run('\n')
                    run_0.font.size = shared.Pt(7)
                    run = paragraph.add_run('Caratteristiche dei titoli azionari')
                    run.bold = True
                    run.font.name = 'Century Gothic'
                    run.font.size = shared.Pt(14)
                    run.font.color.rgb = shared.RGBColor(127, 127, 127)
                    paragraph = self.document.add_paragraph(text='', style=None)
                    paragraph.paragraph_format.line_spacing = shared.Cm(0.2)
                    run = paragraph.add_run()
                    run.add_picture(self.path_img_default.joinpath('barra.png').__str__(), width=shared.Cm(self.larghezza_pagina+0.1))
                    paragraph = self.document.add_paragraph(text='', style=None)
                    run = paragraph.add_run()
                    run.add_picture(self.path_img.joinpath('azioni_'+str(tabella)+'.png').__str__(), width=shared.Cm(self.larghezza_pagina) if hidden_columns==0 else shared.Cm(self.larghezza_pagina-3.5))
            # wb.Close(SaveChanges=False, Filename=self.path+"\\"+self.file_elaborato)

        elif numero_prodotti_azionari == 0:
            numerosita_tabella_azioni_sotto_la_precedente = 0

        # Fondi
        prodotti_gestiti = df_portfolio.loc[df_portfolio['strumento']=='fund']
        numero_prodotti_gestiti = prodotti_gestiti.nome.count()
        # print('numero fondi:',numero_prodotti_gestiti)
        if numero_prodotti_gestiti > 0:
            # Carica il foglio fondi
            fondi = self.wb['fondi']
            self.wb.active = fondi
            # Nascondi le colonne del prezzo di carico e della variazione di prezzo
            hidden_columns = 0
            if all(df_portfolio['prezzo_di_carico'].isnull()):
                fondi.column_dimensions['F'].hidden= True
                hidden_columns += 1
                fondi.column_dimensions['H'].hidden= True
                hidden_columns += 1
                fondi.column_dimensions['I'].hidden= True
                hidden_columns += 1
            self.wb.save(self.file_elaborato)

            # Calcola il numero di prodotti nell'ultima pagina
            if numero_prodotti_obbligazionari == 0 and numero_prodotti_azionari == 0: # non ci sono obbligazioni nè azioni
                num_prodotti_ultima_pagina = 0
            elif numero_prodotti_obbligazionari == 0 and numero_prodotti_azionari > 0: # ci sono azioni ma non obbligazioni
                if numero_prodotti_azionari % MAX_AZIONI_PER_PAGINA == 0:
                    num_prodotti_ultima_pagina = MAX_AZIONI_PER_PAGINA
                elif numero_prodotti_azionari % MAX_AZIONI_PER_PAGINA != 0:
                    num_prodotti_ultima_pagina = numero_prodotti_azionari % MAX_AZIONI_PER_PAGINA
            elif numero_prodotti_obbligazionari > 0 and numero_prodotti_azionari == 0: # ci sono obbligazioni ma non azioni
                if numero_prodotti_obbligazionari > MAX_OBB_DES_PER_PAGINA:
                    num_prodotti_ultima_pagina = (numero_prodotti_obbligazionari - numerosita_tabella_obb_dati_sotto_la_precedente) % MAX_OBB_DATI_PER_PAGINA
                elif numero_prodotti_obbligazionari <= MAX_OBB_DES_PER_PAGINA:
                    if numerosita_tabella_obb_dati_sotto_la_precedente == 0:
                        if numero_prodotti_obbligazionari <= MAX_OBB_DATI_PER_PAGINA:
                            num_prodotti_ultima_pagina = numero_prodotti_obbligazionari
                        elif numero_prodotti_obbligazionari > MAX_OBB_DATI_PER_PAGINA:
                            num_prodotti_ultima_pagina = (numero_prodotti_obbligazionari - numerosita_tabella_obb_dati_sotto_la_precedente) % MAX_OBB_DATI_PER_PAGINA
                    elif numerosita_tabella_obb_dati_sotto_la_precedente > 0:
                        if numerosita_tabella_obb_dati_sotto_la_precedente == numero_prodotti_obbligazionari:
                            num_prodotti_ultima_pagina = numero_prodotti_obbligazionari + numero_prodotti_obbligazionari
                        elif numerosita_tabella_obb_dati_sotto_la_precedente < numero_prodotti_obbligazionari:
                            num_prodotti_ultima_pagina = (numero_prodotti_obbligazionari - numerosita_tabella_obb_dati_sotto_la_precedente) % MAX_OBB_DATI_PER_PAGINA
            elif numero_prodotti_obbligazionari > 0 and numero_prodotti_azionari > 0: # ci sono obbligazioni e azioni
                if numerosita_tabella_azioni_sotto_la_precedente == numero_prodotti_azionari: # le azioni stanno tutte sotto l'ultima tabella precedente
                    num_prodotti_ultima_pagina = num_prodotti_ultima_pagina + numerosita_tabella_azioni_sotto_la_precedente
                elif numerosita_tabella_azioni_sotto_la_precedente == 0: # le tabelle delle azioni cominciano da pagina nuova
                    num_prodotti_ultima_pagina = numero_prodotti_azionari % MAX_AZIONI_PER_PAGINA
                elif numerosita_tabella_azioni_sotto_la_precedente < numero_prodotti_azionari: # le azioni non stanno tutte sotto l'ultima tabella precedente
                    if (numero_prodotti_azionari - numerosita_tabella_azioni_sotto_la_precedente) % MAX_AZIONI_PER_PAGINA == 0: # l'ultima pagina contiene il numero massimo di azioni
                        num_prodotti_ultima_pagina = MAX_AZIONI_PER_PAGINA
                    else: # l'ultima pagina contiene un numero di azioni inferiore al numero massimo di azioni per pagina
                        num_prodotti_ultima_pagina = (numero_prodotti_azionari - numerosita_tabella_azioni_sotto_la_precedente) % MAX_AZIONI_PER_PAGINA
            # print("prodotti nell'ultima pagina:",num_prodotti_ultima_pagina)

            # Calcolo il numero dei fondi da inserire nell'eventuale tabella sotto l'ultima
            if num_prodotti_ultima_pagina == 0: # se non ci sono obbligazioni nè azioni
                numerosita_tabella_fondi_sotto_la_precedente = 0
            elif numero_prodotti_obbligazionari == 0: # non ci sono obbligazioni
                if MAX_FONDI_PER_PAGINA - int(num_prodotti_ultima_pagina*MAX_FONDI_PER_PAGINA/MAX_AZIONI_PER_PAGINA) - 9 > 0: # se rimane spazio sufficiente sotto la tabella precedente
                    if (MAX_FONDI_PER_PAGINA - int(num_prodotti_ultima_pagina*MAX_FONDI_PER_PAGINA/MAX_AZIONI_PER_PAGINA) - 9) <= numero_prodotti_gestiti: # ma non ce nè abbastanza per tutti i fondi
                        numerosita_tabella_fondi_sotto_la_precedente = MAX_FONDI_PER_PAGINA - int(num_prodotti_ultima_pagina*MAX_FONDI_PER_PAGINA/MAX_AZIONI_PER_PAGINA) - 9
                    else: # e tutti i fondi ci stanno in quello spazio rimasto
                        numerosita_tabella_fondi_sotto_la_precedente = numero_prodotti_gestiti
                else: # se non rimane spazio a sufficienza per una tabella sotto la precedente
                    numerosita_tabella_fondi_sotto_la_precedente = 0
            elif numero_prodotti_azionari == 0: # non ci sono azioni
                if num_prodotti_ultima_pagina == (numero_prodotti_obbligazionari * 2): # se le tabelle des e dati sono sulla stessa pagina
                    if MAX_FONDI_PER_PAGINA - int(numero_prodotti_obbligazionari*MAX_FONDI_PER_PAGINA*(MAX_OBB_DES_PER_PAGINA+MAX_OBB_DATI_PER_PAGINA)/(MAX_OBB_DES_PER_PAGINA*MAX_OBB_DATI_PER_PAGINA)) - 18 > 0: # se rimane spazio sufficiente sotto le due tabelle precedenti
                        if (MAX_FONDI_PER_PAGINA - int(numero_prodotti_obbligazionari*MAX_FONDI_PER_PAGINA*(MAX_OBB_DES_PER_PAGINA+MAX_OBB_DATI_PER_PAGINA)/(MAX_OBB_DES_PER_PAGINA*MAX_OBB_DATI_PER_PAGINA)) - 18) < numero_prodotti_gestiti: # ma non ce nè abbastanza per tutti i fondi
                            numerosita_tabella_fondi_sotto_la_precedente = MAX_FONDI_PER_PAGINA - int(numero_prodotti_obbligazionari*MAX_FONDI_PER_PAGINA*(MAX_OBB_DES_PER_PAGINA+MAX_OBB_DATI_PER_PAGINA)/(MAX_OBB_DES_PER_PAGINA*MAX_OBB_DATI_PER_PAGINA)) - 18
                        else: # e tutti i fondi ci stanno in quello spazio rimasto
                            numerosita_tabella_fondi_sotto_la_precedente = numero_prodotti_gestiti
                    else: # se non rimane spazio a sufficienza per una tabella sotto la precedente
                        numerosita_tabella_fondi_sotto_la_precedente = 0
                else: # ci sono obbligazioni ma le tabelle des e dati non sono sulla stessa pagina
                    if MAX_FONDI_PER_PAGINA - int(num_prodotti_ultima_pagina*MAX_FONDI_PER_PAGINA/MAX_OBB_DATI_PER_PAGINA) - 9> 0: # se rimane spazio sufficiente sotto le due tabelle precedenti
                        if (MAX_FONDI_PER_PAGINA - int(num_prodotti_ultima_pagina*MAX_FONDI_PER_PAGINA/MAX_OBB_DATI_PER_PAGINA) - 9) < numero_prodotti_gestiti: # ma non ce nè abbastanza per tutti i fondi
                            numerosita_tabella_fondi_sotto_la_precedente = MAX_FONDI_PER_PAGINA - int(num_prodotti_ultima_pagina*MAX_FONDI_PER_PAGINA/MAX_OBB_DATI_PER_PAGINA) - 9
                        else: # e tutti i fondi ci stanno in quello spazio rimasto
                            numerosita_tabella_fondi_sotto_la_precedente = numero_prodotti_gestiti
                    else: # se non rimane spazio a sufficienza per una tabella sotto la precedente
                        numerosita_tabella_fondi_sotto_la_precedente = 0
            elif num_prodotti_ultima_pagina == (numero_prodotti_obbligazionari * 2) + numero_prodotti_azionari: # se le tabelle delle obbligazioni e delle azioni sono sulla stessa pagina
                if MAX_FONDI_PER_PAGINA - int(numero_prodotti_obbligazionari*MAX_FONDI_PER_PAGINA*(MAX_OBB_DES_PER_PAGINA+MAX_OBB_DATI_PER_PAGINA)/(MAX_OBB_DES_PER_PAGINA*MAX_OBB_DATI_PER_PAGINA)) - int(numero_prodotti_azionari*MAX_FONDI_PER_PAGINA/MAX_AZIONI_PER_PAGINA) - 26 > 0: # se rimane spazio sufficiente sotto le tre tabelle precedenti
                    if (MAX_FONDI_PER_PAGINA - int(numero_prodotti_obbligazionari*MAX_FONDI_PER_PAGINA*(MAX_OBB_DES_PER_PAGINA+MAX_OBB_DATI_PER_PAGINA)/(MAX_OBB_DES_PER_PAGINA*MAX_OBB_DATI_PER_PAGINA)) - int(numero_prodotti_azionari*MAX_FONDI_PER_PAGINA/MAX_AZIONI_PER_PAGINA) - 26) < numero_prodotti_gestiti: # ma non ce nè abbastanza per tutti i fondi
                        numerosita_tabella_fondi_sotto_la_precedente = MAX_FONDI_PER_PAGINA - int(numero_prodotti_obbligazionari*MAX_FONDI_PER_PAGINA*(MAX_OBB_DES_PER_PAGINA+MAX_OBB_DATI_PER_PAGINA)/(MAX_OBB_DES_PER_PAGINA*MAX_OBB_DATI_PER_PAGINA)) - int(numero_prodotti_azionari*MAX_FONDI_PER_PAGINA/MAX_AZIONI_PER_PAGINA) - 26
                    else: # e tutti i fondi ci stanno in quello spazio rimasto
                        numerosita_tabella_fondi_sotto_la_precedente = numero_prodotti_gestiti
                else: # se non rimane spazio a sufficienza per una tabella sotto la precedente
                    numerosita_tabella_fondi_sotto_la_precedente = 0
            else: # ci sono obbligazioni e/o azioni, ma le tabelle non sono sulla stessa pagina
                if numero_prodotti_azionari <= numerosita_tabella_azioni_sotto_la_precedente: # l'ultima pagina ha la tabella obbligazioni dati e la tabella azioni
                    if MAX_FONDI_PER_PAGINA - int((num_prodotti_ultima_pagina-numero_prodotti_azionari)*MAX_FONDI_PER_PAGINA/MAX_OBB_DATI_PER_PAGINA) - int(numero_prodotti_azionari*MAX_FONDI_PER_PAGINA/MAX_AZIONI_PER_PAGINA) - 18 > 0: # se rimane spazio sufficiente sotto le tabelle precedenti
                        if (MAX_FONDI_PER_PAGINA - int((num_prodotti_ultima_pagina-numero_prodotti_azionari)*MAX_FONDI_PER_PAGINA/MAX_OBB_DATI_PER_PAGINA) - int(numero_prodotti_azionari*MAX_FONDI_PER_PAGINA/MAX_AZIONI_PER_PAGINA) - 18) < numero_prodotti_gestiti: # ma non ce nè abbastanza per tutti i fondi
                            numerosita_tabella_fondi_sotto_la_precedente = MAX_FONDI_PER_PAGINA - int((num_prodotti_ultima_pagina-numero_prodotti_azionari)*MAX_FONDI_PER_PAGINA/MAX_OBB_DATI_PER_PAGINA) - int(numero_prodotti_azionari*MAX_FONDI_PER_PAGINA/MAX_AZIONI_PER_PAGINA) - 18
                        else: # e tutti i fondi ci stanno in quello spazio rimasto
                            numerosita_tabella_fondi_sotto_la_precedente = numero_prodotti_gestiti
                    else: # se non rimane spazio a sufficienza per una tabella sotto la precedente
                        numerosita_tabella_fondi_sotto_la_precedente = 0
                elif numero_prodotti_azionari > numerosita_tabella_azioni_sotto_la_precedente: # l'ultima pagina ha una sola tabella di azioni
                    if MAX_FONDI_PER_PAGINA - int(num_prodotti_ultima_pagina*MAX_FONDI_PER_PAGINA/MAX_AZIONI_PER_PAGINA) - 9 > 0: # se rimane spazio sufficiente sotto la tabella precedente
                        if (MAX_FONDI_PER_PAGINA - int(num_prodotti_ultima_pagina*MAX_FONDI_PER_PAGINA/MAX_AZIONI_PER_PAGINA) - 9) < numero_prodotti_gestiti: # ma non ce nè abbastanza per tutti i fondi
                            numerosita_tabella_fondi_sotto_la_precedente = MAX_FONDI_PER_PAGINA - int(num_prodotti_ultima_pagina*MAX_FONDI_PER_PAGINA/MAX_AZIONI_PER_PAGINA) - 9
                        else: # e tutti i fondi ci stanno in quello spazio rimasto
                            numerosita_tabella_fondi_sotto_la_precedente = numero_prodotti_gestiti
                    else: # se non rimane spazio a sufficienza per una tabella sotto la precedente
                        numerosita_tabella_fondi_sotto_la_precedente = 0
            # print("numerosità tabella fondi sotto la precedente:",numerosita_tabella_fondi_sotto_la_precedente)
            
            # Inserisci l'eventuale tabella sotto l'ultima
            # # Librerie win32com + PIL
            # xls_file = win32com.client.Dispatch("Excel.Application")
            # wb = xls_file.Workbooks.Open(Filename=self.path+"\\"+self.file_elaborato)
            # # xls_file.Visible = False
            # # xls_file.ScreenUpdating = False
            # # xls_file.DisplayAlerts = False
            # # xls_file.EnableEvents = False
            # ws = wb.Worksheets("fondi")
            if numerosita_tabella_fondi_sotto_la_precedente > 0:
                # Prima tabella dati fondi
                # ws.Range(ws.Cells(1,2),ws.Cells(numerosita_tabella_fondi_sotto_la_precedente+1,9)).CopyPicture(Format=2)
                # img = ImageGrab.grabclipboard()
                # img.save(self.path+r'\img\fondi_0.png')
                # ws.Range(ws.Cells(2,2),ws.Cells(numerosita_tabella_fondi_sotto_la_precedente+1,9)).Rows.EntireRow.Hidden = True
                # Libreria excel2img
                excel2img.export_img(self.file_elaborato, self.path_img.joinpath('fondi_0.png').__str__(), page='fondi', _range="B1:I"+str(numerosita_tabella_fondi_sotto_la_precedente+1))
                fondi.row_dimensions.group(2,numerosita_tabella_fondi_sotto_la_precedente+1,hidden=True)
                self.wb.save(self.file_elaborato)
                # print(0)
                paragraph = self.document.add_paragraph(text='', style=None)
                run = paragraph.add_run('\nCaratteristiche finanziarie dei fondi comuni di investimento')
                run.bold = True
                run.font.name = 'Century Gothic'
                run.font.size = shared.Pt(14)
                run.font.color.rgb = shared.RGBColor(127, 127, 127)
                paragraph = self.document.add_paragraph(text='', style=None)
                paragraph.paragraph_format.line_spacing = shared.Cm(0.2)
                run = paragraph.add_run()
                run.add_picture(self.path_img_default.joinpath('barra.png').__str__(), width=shared.Cm(self.larghezza_pagina+0.1))
                paragraph = self.document.add_paragraph(text='', style=None)
                run = paragraph.add_run()
                run.add_picture(self.path_img.joinpath('fondi_0.png').__str__(), width=shared.Cm(self.larghezza_pagina) if hidden_columns==0 else shared.Cm(self.larghezza_pagina-4.5))
            

            # Inserisci le tabelle rimanenti
            if numero_prodotti_gestiti - numerosita_tabella_fondi_sotto_la_precedente == 0: # tutti i titoli sono contenuti nella tabella sotto l'ultima
                tabelle_fondi = 1
                # print('tabelle_fondi:',tabelle_fondi)
            else:
                if numero_prodotti_gestiti - numerosita_tabella_fondi_sotto_la_precedente <= MAX_FONDI_PER_PAGINA:
                    tabelle_fondi = 1
                    # numerosita_tabella_fondi_sotto_la_precedente = numero_prodotti_gestiti
                elif numero_prodotti_gestiti - numerosita_tabella_fondi_sotto_la_precedente > MAX_FONDI_PER_PAGINA:
                    if (numero_prodotti_gestiti - numerosita_tabella_fondi_sotto_la_precedente) % MAX_FONDI_PER_PAGINA != 0:
                        tabelle_fondi = int((numero_prodotti_gestiti - numerosita_tabella_fondi_sotto_la_precedente) // MAX_FONDI_PER_PAGINA) + 1
                    else:
                        tabelle_fondi = int((numero_prodotti_gestiti - numerosita_tabella_fondi_sotto_la_precedente) // MAX_FONDI_PER_PAGINA)
                    # numerosita_tabella_fondi_sotto_la_precedente = MAX_FONDI_PER_PAGINA
                # print('tabelle_fondi:',tabelle_fondi if numerosita_tabella_fondi_sotto_la_precedente == 0 else tabelle_fondi+1)
                print(f'sto aggiungendo la tabella fondi sotto la precedente: 0 / {tabelle_fondi}') if numerosita_tabella_fondi_sotto_la_precedente > 0 else None
                for tabella in range(1, tabelle_fondi+1):
                    print(f'sto aggiungendo la tabella fondi: {tabella} / {tabelle_fondi}')
                    # print(tabella)
                    if tabella != tabelle_fondi:
                        # ws.Range(ws.Cells(1,2),ws.Cells(numerosita_tabella_fondi_sotto_la_precedente+MAX_FONDI_PER_PAGINA*tabella+1,9)).CopyPicture(Format=2)
                        # img = ImageGrab.grabclipboard()
                        # img.save(self.path+r'\img\fondi_' + str(tabella) + '.png')
                        # ws.Range(ws.Cells(2+MAX_FONDI_PER_PAGINA*(tabella-1),2),ws.Cells(numerosita_tabella_fondi_sotto_la_precedente+MAX_FONDI_PER_PAGINA*tabella+1,9)).Rows.EntireRow.Hidden = True
                        # Libreria excel2img
                        excel2img.export_img(self.file_elaborato, self.path_img.joinpath('fondi_' + str(tabella) + '.png').__str__(), page='fondi', _range="B1:I"+str(numerosita_tabella_fondi_sotto_la_precedente+MAX_FONDI_PER_PAGINA*tabella+1))
                        fondi.row_dimensions.group(2+MAX_FONDI_PER_PAGINA*(tabella-1),numerosita_tabella_fondi_sotto_la_precedente+MAX_FONDI_PER_PAGINA*tabella+1,hidden=True)
                        self.wb.save(self.file_elaborato)
                    else:
                        # ws.Range(ws.Cells(1,2),ws.Cells(numero_prodotti_gestiti+1,9)).CopyPicture(Format=2)
                        # img = ImageGrab.grabclipboard()
                        # img.save(self.path+r'\img\fondi_' + str(tabella) + '.png')
                        # Libreria excel2img
                        excel2img.export_img(self.file_elaborato, self.path_img.joinpath('fondi_' + str(tabella) + '.png').__str__(), page='fondi', _range="B1:I"+str(numero_prodotti_gestiti+1))
                fondi.row_dimensions.group(1,MAX_FONDI_PER_PAGINA*(tabelle_fondi+1),hidden=False)
                self.wb.save(self.file_elaborato)
                
                for tabella in range(1, tabelle_fondi+1):
                    self.document.add_section()
                    paragraph = self.document.add_paragraph(text='', style=None)
                    paragraph.paragraph_format.space_before = shared.Pt(6)
                    paragraph.paragraph_format.space_after = shared.Pt(0)
                    run = paragraph.add_run(text='')
                    run.add_picture(self.path_img_default.joinpath('3_analisi_dei_singoli_strumenti.bmp').__str__(), width=shared.Cm(8.5))
                    paragraph = self.document.add_paragraph(text='', style=None)
                    run_0 = paragraph.add_run('\n')
                    run_0.font.size = shared.Pt(7)
                    run = paragraph.add_run('Caratteristiche finanziarie dei fondi comuni di investimento')
                    run.bold = True
                    run.font.name = 'Century Gothic'
                    run.font.size = shared.Pt(14)
                    run.font.color.rgb = shared.RGBColor(127, 127, 127)
                    paragraph = self.document.add_paragraph(text='', style=None)
                    paragraph.paragraph_format.line_spacing = shared.Cm(0.2)
                    run = paragraph.add_run()
                    run.add_picture(self.path_img_default.joinpath('barra.png').__str__(), width=shared.Cm(self.larghezza_pagina+0.1))
                    paragraph = self.document.add_paragraph(text='', style=None)
                    run = paragraph.add_run()
                    run.add_picture(self.path_img.joinpath('fondi_'+str(tabella)+'.png').__str__(), width=shared.Cm(self.larghezza_pagina) if hidden_columns==0 else shared.Cm(self.larghezza_pagina-4.5))
            # wb.Close(SaveChanges=False, Filename=self.path+"\\"+self.file_elaborato)


            # Mappatura fondi #
            numero_prodotti_gestiti_map = numero_prodotti_gestiti + 2
            if numero_prodotti_gestiti_map > MAX_MAP_FONDI_PER_PAGINA and numero_prodotti_gestiti_map % MAX_MAP_FONDI_PER_PAGINA != 0:
                tabelle_map_fondi = int(numero_prodotti_gestiti_map // MAX_MAP_FONDI_PER_PAGINA + 1)
            elif numero_prodotti_gestiti_map > MAX_MAP_FONDI_PER_PAGINA and numero_prodotti_gestiti_map % MAX_MAP_FONDI_PER_PAGINA == 0:
                tabelle_map_fondi = int(numero_prodotti_gestiti_map // MAX_MAP_FONDI_PER_PAGINA)
            else:
                tabelle_map_fondi = 1
            # print('tabelle_map_fondi:',tabelle_map_fondi)
            # # Librerie win32com + PIL
            # xls_file = win32com.client.Dispatch("Excel.Application")
            # wb = xls_file.Workbooks.Open(Filename=self.path+"\\"+self.file_elaborato)
            # # xls_file.Visible = False
            # # xls_file.ScreenUpdating = False
            # # xls_file.DisplayAlerts = False
            # # xls_file.EnableEvents = False
            # ws = wb.Worksheets("fondi")
            for tabella in range(1, tabelle_map_fondi+1):
                print(f'sto aggiungendo la tabella mappatura fondi: {tabella} / {tabelle_map_fondi}')
                # print(tabella)
                if tabella != tabelle_map_fondi:
                    # ws.Range(ws.Cells(1,12),ws.Cells(MAX_MAP_FONDI_PER_PAGINA*tabella+1,26)).CopyPicture(Format=2)
                    # img = ImageGrab.grabclipboard()
                    # img.save(self.path+r'\img\map_fondi_' + str(tabella) + '.png')
                    # ws.Range(ws.Cells(2+MAX_MAP_FONDI_PER_PAGINA*(tabella-1),12),ws.Cells(MAX_MAP_FONDI_PER_PAGINA*tabella+1,26)).Rows.EntireRow.Hidden = True
                    # Libreria excel2img
                    excel2img.export_img(self.file_elaborato, self.path_img.joinpath('map_fondi_' + str(tabella) + '.png').__str__(), page='fondi', _range="L1:Z"+str(MAX_MAP_FONDI_PER_PAGINA*tabella+1))
                    fondi.row_dimensions.group(2+MAX_MAP_FONDI_PER_PAGINA*(tabella-1),MAX_MAP_FONDI_PER_PAGINA*tabella+1,hidden=True)
                    self.wb.save(self.file_elaborato)
                else:
                    # ws.Range(ws.Cells(1,12),ws.Cells(numero_prodotti_gestiti_map+1,26)).CopyPicture(Format=2)
                    # img = ImageGrab.grabclipboard()
                    # img.save(self.path+r'\img\map_fondi_' + str(tabella) + '.png')
                    # Libreria excel2img
                    excel2img.export_img(self.file_elaborato, self.path_img.joinpath('ap_fondi_' + str(tabella) + '.png').__str__(), page='fondi', _range="L1:Z"+str(numero_prodotti_gestiti_map+1))      
            fondi.row_dimensions.group(1,MAX_MAP_FONDI_PER_PAGINA*tabelle_map_fondi,hidden=False)
            self.wb.save(self.file_elaborato)
            # wb.Close(SaveChanges=False, Filename=self.path+"\\"+self.file_elaborato)

            for tabella in range(1, tabelle_map_fondi+1):
                self.document.add_section()
                paragraph = self.document.add_paragraph(text='', style=None)
                paragraph.paragraph_format.space_before = shared.Pt(6)
                paragraph.paragraph_format.space_after = shared.Pt(0)
                run = paragraph.add_run(text='')
                run.add_picture(self.path_img_default.joinpath('3_analisi_dei_singoli_strumenti.bmp').__str__(), width=shared.Cm(8.5))
                paragraph = self.document.add_paragraph(text='', style=None)
                run_0 = paragraph.add_run('\n')
                run_0.font.size = shared.Pt(7)
                run = paragraph.add_run('Mappatura dei fondi comuni di investimento')
                run.bold = True
                run.font.name = 'Century Gothic'
                run.font.size = shared.Pt(14)
                run.font.color.rgb = shared.RGBColor(127, 127, 127)
                paragraph = self.document.add_paragraph(text='', style=None)
                paragraph.paragraph_format.line_spacing = shared.Cm(0.2)
                run = paragraph.add_run()
                run.add_picture(self.path_img_default.joinpath('barra.png').__str__(), width=shared.Cm(self.larghezza_pagina+0.1))
                paragraph = self.document.add_paragraph(text='', style=None)
                run = paragraph.add_run()
                run.add_picture(self.path_img_default.joinpath('map_fondi_info.bmp').__str__(), width=shared.Cm(self.larghezza_pagina))
                paragraph = self.document.add_paragraph(text='', style=None)
                run = paragraph.add_run()
                run.add_picture(self.path_img.joinpath('map_fondi_'+str(tabella)+'.png').__str__(), width=shared.Cm(self.larghezza_pagina))

            # Calcola numero fondi mappati nell'ultima pagina
            if numero_prodotti_gestiti_map <= MAX_MAP_FONDI_PER_PAGINA:
                num_prodotti_ultima_pagina = numero_prodotti_gestiti_map
            elif numero_prodotti_gestiti_map > MAX_MAP_FONDI_PER_PAGINA:
                if numero_prodotti_gestiti_map % MAX_MAP_FONDI_PER_PAGINA != 0:
                    num_prodotti_ultima_pagina = numero_prodotti_gestiti_map % MAX_MAP_FONDI_PER_PAGINA
                elif numero_prodotti_gestiti_map % MAX_MAP_FONDI_PER_PAGINA == 0:
                    num_prodotti_ultima_pagina = MAX_MAP_FONDI_PER_PAGINA
            # print("numerosità ultima tabella mappatura fondi:",num_prodotti_ultima_pagina)

            if MAX_MAP_FONDI_PER_PAGINA - num_prodotti_ultima_pagina - 29 > 0: # c'è spazio per inserire il grafico a barre
                paragraph = self.document.add_paragraph(text='', style=None)
                run = paragraph.add_run()
                run.add_picture(self.path_img.joinpath('map_fondi_bar.png').__str__(), width=shared.Cm(self.larghezza_pagina))
            else: # non c'è spazio per inserire il grafico a barre
                self.document.add_section()
                paragraph = self.document.add_paragraph(text='', style=None)
                paragraph.paragraph_format.space_before = shared.Pt(6)
                paragraph.paragraph_format.space_after = shared.Pt(0)
                run = paragraph.add_run(text='')
                run.add_picture(self.path_img_default.joinpath('3_analisi_dei_singoli_strumenti.bmp').__str__(), width=shared.Cm(8.5))
                paragraph = self.document.add_paragraph(text='', style=None)
                run_0 = paragraph.add_run('\n')
                run_0.font.size = shared.Pt(7)
                run = paragraph.add_run('Mappatura dei fondi comuni di investimento')
                run.bold = True
                run.font.name = 'Century Gothic'
                run.font.size = shared.Pt(14)
                run.font.color.rgb = shared.RGBColor(127, 127, 127)
                paragraph = self.document.add_paragraph(text='', style=None)
                paragraph.paragraph_format.line_spacing = shared.Cm(0.2)
                run = paragraph.add_run()
                run.add_picture(self.path_img_default.joinpath('barra.png').__str__(), width=shared.Cm(self.larghezza_pagina+0.1))
                paragraph = self.document.add_paragraph(text='', style=None)
                run = paragraph.add_run()
                run.add_picture(self.path_img_default.joinpath('map_fondi_info.bmp').__str__(), width=shared.Cm(self.larghezza_pagina))
                paragraph = self.document.add_paragraph(text='', style=None)
                run = paragraph.add_run()
                run.add_picture(self.path_img.joinpath('map_fondi_bar.png').__str__(), width=shared.Cm(self.larghezza_pagina))
            
            # Emittenti #
            # Se gli emittenti sono tanti la tabella e il grafico diventano illeggibili.
            # Meglio fare il top 20 emittenti nel portafoglio.
            dict_emittenti = self.peso_emittente()
            if dict_emittenti is None:
                pass
            else:
                print('sto aggiungendo la tabella emittenti...')
                self.document.add_section()
                paragraph = self.document.add_paragraph(text='', style=None)
                paragraph.paragraph_format.space_before = shared.Pt(6)
                paragraph.paragraph_format.space_after = shared.Pt(0)
                run = paragraph.add_run(text='')
                run.add_picture(self.path_img_default.joinpath('3_analisi_dei_singoli_strumenti.bmp').__str__(), width=shared.Cm(8.5))
                paragraph = self.document.add_paragraph(text='', style=None)
                paragraph.paragraph_format.alignment = 1
                run_0 = paragraph.add_run('\n')
                run_0.font.size = shared.Pt(7)
                run = paragraph.add_run('Concentrazione dei maggiori 20 emittenti nel portafoglio')
                run.bold = True
                run.font.name = 'Century Gothic'
                run.font.size = shared.Pt(14)
                run.font.color.rgb = shared.RGBColor(127, 127, 127)
                paragraph = self.document.add_paragraph(text='', style=None)
                paragraph.paragraph_format.alignment = 1
                run = paragraph.add_run()
                excel2img.export_img(self.file_elaborato, self.path_img.joinpath('emittenti' + '.png').__str__(), page='figure', _range="Z1:AB22") # dipende dal top n nuemittenti scelti
                run.add_picture(self.path_img.joinpath('emittenti.png').__str__(), width=shared.Cm(10))
                paragraph = self.document.add_paragraph(text='', style=None)
                run = paragraph.add_run('\n\n\n')
                paragraph = self.document.add_paragraph(text='', style=None)
                paragraph.paragraph_format.alignment = 1
                run = paragraph.add_run()
                run.add_picture(self.path_img.joinpath('emittenti_bar.png').__str__(), width=shared.Cm(self.larghezza_pagina))

            # Matrice di correlazione dei fondi attivi e passivi #
            # matr_corr = self.matrice_correlazioni()
            # if matr_corr is None:
            #     pass
            # else:
            #     print('sto aggiungendo la matrice delle correlazioni...')
            #     self.document.add_section()
            #     paragraph = self.document.add_paragraph(text='', style=None)
            #     paragraph.paragraph_format.space_before = shared.Pt(6)
            #     paragraph.paragraph_format.space_after = shared.Pt(0)
            #     run = paragraph.add_run(text='')
            #     run.add_picture(self.path_img_default.joinpath('3_analisi_dei_singoli_strumenti.bmp').__str__(), width=shared.Cm(8.5))
            #     paragraph = self.document.add_paragraph(text='', style=None)
            #     paragraph.paragraph_format.alignment = 1
            #     run_0 = paragraph.add_run('\n')
            #     run_0.font.size = shared.Pt(7)
            #     run = paragraph.add_run('Matrice di correlazione dei fondi attivi e passivi')
            #     run.bold = True
            #     run.font.name = 'Century Gothic'
            #     run.font.size = shared.Pt(14)
            #     run.font.color.rgb = shared.RGBColor(127, 127, 127)
            #     paragraph = self.document.add_paragraph(text='', style=None)
            #     paragraph.paragraph_format.alignment = 1
            #     run = paragraph.add_run('\n\n')
            #     run.add_picture(self.path_img.joinpath('matr_corr.png').__str__(), width=shared.Cm(self.larghezza_pagina), height=shared.Cm(12))              

    def analisi_del_rischio(self):
        """Inserisci la parte di rischio"""
        # Prima pagina
        self.document.add_section()
        paragraph_0 = self.document.add_paragraph(text='', style=None)
        paragraph_0.paragraph_format.space_before = shared.Pt(6)
        paragraph_0.paragraph_format.space_after = shared.Pt(6)
        run_0 = paragraph_0.add_run(text='')
        run_0.add_picture(self.path_img_default.joinpath('4_analisi_del_rischio.bmp').__str__(), width=shared.Cm(8.5))
        paragraph_1 = self.document.add_paragraph(text='', style=None)
        run_1 = paragraph_1.add_run('')
        run_1.add_picture(self.path_img_default.joinpath('rischio_info_1.bmp').__str__(), width=shared.Cm(self.larghezza_pagina))
        paragraph_2 = self.document.add_paragraph(text='\n\n\n\n\n\n\n', style=None)
        run_2 = paragraph_2.add_run('')
        vol = self.risk()
        if vol <= 0.03:
            profilo = 'basso'
        elif vol > 0.03 and vol <= 0.07:
            profilo = 'medio_basso'
        elif vol > 0.07 and vol <= 0.12:
            profilo = 'medio'
        elif vol > 0.12 and vol <= 0.15:
            profilo = 'medio_alto'
        elif vol > 0.15:
            profilo = 'alto'
        run_2.add_picture(self.path_img_default.joinpath('rischio_profilo_'+profilo+'.bmp').__str__(), width=shared.Cm(self.larghezza_pagina))
        # Seconda pagina
        self.document.add_section()
        paragraph_0 = self.document.add_paragraph(text='', style=None)
        paragraph_0.paragraph_format.space_before = shared.Pt(6)
        paragraph_0.paragraph_format.space_after = shared.Pt(6)
        run_0 = paragraph_0.add_run(text='')
        run_0.add_picture(self.path_img_default.joinpath('4_analisi_del_rischio.bmp').__str__(), width=shared.Cm(8.5))
        paragraph_1 = self.document.add_paragraph(text='', style=None)
        run_1 = paragraph_1.add_run('')
        run_1.add_picture(self.path_img_default.joinpath('rischio_info_2.bmp').__str__(), width=shared.Cm(self.larghezza_pagina))
        # Terza pagina
        self.document.add_section()
        paragraph_0 = self.document.add_paragraph(text='', style=None)
        paragraph_0.paragraph_format.space_before = shared.Pt(6)
        paragraph_0.paragraph_format.space_after = shared.Pt(6)
        run_0 = paragraph_0.add_run(text='')
        run_0.add_picture(self.path_img_default.joinpath('4_analisi_del_rischio.bmp').__str__(), width=shared.Cm(8.5))
        paragraph_1 = self.document.add_paragraph(text='', style=None)
        run_1 = paragraph_1.add_run('')
        run_1.add_picture(self.path_img_default.joinpath('rischio_info_3.bmp').__str__(), width=shared.Cm(self.larghezza_pagina))
        # Quarta pagina
        self.document.add_section()
        paragraph_0 = self.document.add_paragraph(text='', style=None)
        paragraph_0.paragraph_format.space_before = shared.Pt(6)
        paragraph_0.paragraph_format.space_after = shared.Pt(6)
        run_0 = paragraph_0.add_run(text='')
        run_0.add_picture(self.path_img_default.joinpath('4_analisi_del_rischio.bmp').__str__(), width=shared.Cm(8.5))
        paragraph_1 = self.document.add_paragraph(text='', style=None)
        run_1 = paragraph_1.add_run('')
        run_1.add_picture(self.path_img_default.joinpath('rischio_info_4.bmp').__str__(), width=shared.Cm(self.larghezza_pagina))
        # Quinta pagina
        self.document.add_section()
        paragraph_0 = self.document.add_paragraph(text='', style=None)
        paragraph_0.paragraph_format.space_before = shared.Pt(6)
        paragraph_0.paragraph_format.space_after = shared.Pt(6)
        run_0 = paragraph_0.add_run(text='')
        run_0.add_picture(self.path_img_default.joinpath('4_analisi_del_rischio.bmp').__str__(), width=shared.Cm(8.5))
        paragraph_1 = self.document.add_paragraph(text='', style=None)
        run_1 = paragraph_1.add_run('')
        run_1.add_picture(self.path_img_default.joinpath('rischio_info_5.bmp').__str__(), width=shared.Cm(self.larghezza_pagina))
        # Sesta pagina
        self.document.add_section()
        paragraph_0 = self.document.add_paragraph(text='', style=None)
        paragraph_0.paragraph_format.space_before = shared.Pt(6)
        paragraph_0.paragraph_format.space_after = shared.Pt(6)
        run_0 = paragraph_0.add_run(text='')
        run_0.add_picture(self.path_img_default.joinpath('4_analisi_del_rischio.bmp').__str__(), width=shared.Cm(8.5))
        paragraph_1 = self.document.add_paragraph(text='', style=None)
        run_1 = paragraph_1.add_run('')
        run_1.add_picture(self.path_img_default.joinpath('rischio_info_6.bmp').__str__(), width=shared.Cm(self.larghezza_pagina))
        # Settima pagina
        self.document.add_section()
        paragraph_0 = self.document.add_paragraph(text='', style=None)
        paragraph_0.paragraph_format.space_before = shared.Pt(6)
        paragraph_0.paragraph_format.space_after = shared.Pt(6)
        run_0 = paragraph_0.add_run(text='')
        run_0.add_picture(self.path_img_default.joinpath('4_analisi_del_rischio.bmp').__str__(), width=shared.Cm(8.5))
        paragraph_1 = self.document.add_paragraph(text='', style=None)
        run_1 = paragraph_1.add_run('')
        run_1.add_picture(self.path_img_default.joinpath('rischio_info_7.bmp').__str__(), width=shared.Cm(self.larghezza_pagina))
        paragraph_2 = self.document.add_paragraph(text='\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n', style=None)
        run_2 = paragraph_2.add_run(text='')
        run_2.add_picture(self.path_img_default.joinpath('rischio_info_7_footer.bmp').__str__(), width=shared.Cm(self.larghezza_pagina))
        # Ottava pagina
        self.document.add_section()
        paragraph_0 = self.document.add_paragraph(text='', style=None)
        paragraph_0.paragraph_format.space_before = shared.Pt(6)
        paragraph_0.paragraph_format.space_after = shared.Pt(6)
        run_0 = paragraph_0.add_run(text='')
        run_0.add_picture(self.path_img_default.joinpath('4_analisi_del_rischio.bmp').__str__(), width=shared.Cm(8.5))
        paragraph_1 = self.document.add_paragraph(text='', style=None)
        run_1_1 = paragraph_1.add_run('\n')
        run_1_1.font.size = shared.Pt(10)
        run_1_2 = paragraph_1.add_run('Analisi del rischio dei singoli strumenti')
        run_1_2.bold = True
        run_1_2.font.name = 'Century Gothic'
        run_1_2.font.size = shared.Pt(14)
        run_1_2.font.color.rgb = shared.RGBColor(127, 127, 127)
        paragraph_2 = self.document.add_paragraph(text='', style=None)
        paragraph_2.paragraph_format.line_spacing = shared.Cm(0.2)
        run_2 = paragraph_2.add_run()
        run_2.add_picture(self.path_img_default.joinpath('barra.png').__str__(), width=shared.Cm(self.larghezza_pagina+0.1))

    def note_metodologiche(self):
        """Inserisci le note metodologiche e le avvertenze più la pagina di chiusura."""
        # Note metodologiche 1
        print('sto aggiungendo le note metodologiche...')
        self.document.add_section()
        paragraph_0 = self.document.add_paragraph(text='', style=None)
        paragraph_0.paragraph_format.space_before = shared.Pt(6)
        paragraph_0.paragraph_format.space_after = shared.Pt(0)
        run_0 = paragraph_0.add_run(text='')
        run_0.add_picture(self.path_img_default.joinpath('5_note_metodologiche.bmp').__str__(), width=shared.Cm(8.5))
        paragraph_1 = self.document.add_paragraph(text='\n', style=None)
        paragraph_1.paragraph_format.alignment = 3
        paragraph_1.paragraph_format.space_after = shared.Pt(6)
        if self.intermediario == 'azimut':
            run_1 = paragraph_1.add_run('Nello svolgimento di questa analisi ci siamo avvalsi della documentazione fornitaci da Azimut Wealth Management. Tali informazioni saranno assunte come attendibili da Benchmark&Style. Sono inoltre stati analizzati i dati di mercato tratti da MorningStar e Bloomberg.')
        elif self.intermediario == 'copernico':
            run_1 = paragraph_1.add_run("La Società Copernico, nell'analisi dei dati ed elaborazione del report, si è avvalsa del supporto della Società Benchmark&Style e sono stati analizzati i dati di mercato tratti da Bloomberg e MorningStar.")
        run_1.bold = True
        run_1.font.name = 'Century Gothic'
        run_1.font.size = shared.Pt(10)
        paragraph_2 = self.document.add_paragraph(text='', style=None)
        paragraph_2.paragraph_format.alignment = 3
        paragraph_2.paragraph_format.space_after = shared.Pt(6)
        run_2_1 = paragraph_2.add_run('Sezione A): ')
        run_2_1.bold = True
        run_2_1.font.name = 'Century Gothic'
        run_2_1.font.size = shared.Pt(10)
        run_2_2 = paragraph_2.add_run('nella sezione A viene riportata la composizione del portafoglio. Le quantità dei titoli/prodotti in portafoglio fanno riferimento ai dati di reportistica cliente ricevuta, mentre il controvalore fa riferimento al valore degli stessi alla data di analisi del portafoglio.')
        run_2_2.font.name = 'Century Gothic'
        run_2_2.font.size = shared.Pt(10)
        paragraph_3 = self.document.add_paragraph(text='', style=None)
        paragraph_3.paragraph_format.alignment = 3
        paragraph_3.paragraph_format.space_after = shared.Pt(6)
        run_3_1 = paragraph_3.add_run('Sezione B): ')
        run_3_1.bold = True
        run_3_1.font.name = 'Century Gothic'
        run_3_1.font.size = shared.Pt(10)
        run_3_2 = paragraph_3.add_run('le segnalazioni riportate nelle pagine dedicate all’analisi del portafoglio evidenziano eventuali concen-trazioni del portafoglio stesso su specifiche tipologie di prodotto/strumento, su particolari asset class e sulle valute; per le concentrazioni in valuta si distingue tra investimenti in Euro e in valute diverse dall’Euro.')
        run_3_2.font.name = 'Century Gothic'
        run_3_2.font.size = shared.Pt(10)
        paragraph_4 = self.document.add_paragraph(text='', style=None)
        paragraph_4.paragraph_format.alignment = 0
        paragraph_4.paragraph_format.space_after = shared.Pt(6)
        run_4_1 = paragraph_4.add_run('Sezione C): ')
        run_4_1.bold = True
        run_4_1.font.name = 'Century Gothic'
        run_4_1.font.size = shared.Pt(10)
        run_4_2 = paragraph_4.add_run('per la descrizione sintetica dei titoli obbligazionari sono stati adottati i seguenti criteri:')
        run_4_2.font.name = 'Century Gothic'
        run_4_2.font.size = shared.Pt(10)
        paragraph_5 = self.document.add_paragraph(text='', style=None)
        paragraph_5.paragraph_format.alignment = 0
        paragraph_5.paragraph_format.space_after = shared.Pt(6)
        run_5 = paragraph_5.add_run('- i rating indicano il giudizio espresso dalle principali società di rating (Moodys, Standard and Poor’s e Fitch);')
        run_5.font.name = 'Century Gothic'
        run_5.font.size = shared.Pt(10)
        paragraph_6 = self.document.add_paragraph(text='', style=None)
        paragraph_6.paragraph_format.alignment = 0
        paragraph_6.paragraph_format.space_after = shared.Pt(6)
        run_6 = paragraph_6.add_run('- con la tipologia FIXED facciamo riferimento ad obbligazioni a tasso fisso, mentre VARIABLE indica obbligazioni a tasso variabile;')
        run_6.font.name = 'Century Gothic'
        run_6.font.size = shared.Pt(10)
        paragraph_7 = self.document.add_paragraph(text='', style=None)
        paragraph_7.paragraph_format.alignment = 0
        paragraph_7.paragraph_format.space_after = shared.Pt(6)
        run_7 = paragraph_7.add_run('- la duration ponderata fa riferimento alla sola componente obbligazionaria del portafoglio.')
        run_7.font.name = 'Century Gothic'
        run_7.font.size = shared.Pt(10)
        paragraph_8 = self.document.add_paragraph(text='', style=None)
        paragraph_8.paragraph_format.alignment = 3
        paragraph_8.paragraph_format.space_after = shared.Pt(6)
        run_8 = paragraph_8.add_run('La duration di un portafoglio titoli (o di un singolo titolo) è la vita finanziaria media ponderata dei titoli presenti nel portafoglio (o di un singolo titolo); si tratta di un valore espresso in anni ed indica l’arco temporale entro il quale l’investitore rientrerà in possesso del capitale inizialmente investito, tenendo conto sia delle cedole, sia del rimborso finale.')
        run_8.font.name = 'Century Gothic'
        run_8.font.size = shared.Pt(10)
        paragraph_9 = self.document.add_paragraph(text='', style=None)
        paragraph_9.paragraph_format.alignment = 3
        paragraph_9.paragraph_format.space_after = shared.Pt(6)
        run_9_1 = paragraph_9.add_run('Sezione D): ')
        run_9_1.bold = True
        run_9_1.font.name = 'Century Gothic'
        run_9_1.font.size = shared.Pt(10)
        run_9_2 = paragraph_9.add_run('per il calcolo del portafoglio di rischio si è proceduto come segue: tutti gli strumenti e i prodotti presenti in portafoglio sono stati ricondotti a delle specifiche asset class. Conoscendo il grado di rischio delle singole asset class e il grado di correlazione che le lega fra loro, si procede al calcolo del grado di rischio complessivo del portafoglio.')
        run_9_2.font.name = 'Century Gothic'
        run_9_2.font.size = shared.Pt(10)
        paragraph_10 = self.document.add_paragraph(text='', style=None)
        paragraph_10.paragraph_format.alignment = 3
        paragraph_10.paragraph_format.space_after = shared.Pt(6)
        run_10_1 = paragraph_10.add_run('Sezione E): ')
        run_10_1.bold = True
        run_10_1.font.name = 'Century Gothic'
        run_10_1.font.size = shared.Pt(10)
        run_10_2 = paragraph_10.add_run('nell’analisi di dettaglio si riportano le segnalazioni a livello di singolo prodotto/strumento distinguendo tra segnalazioni di concentrazione e di liquidabilità. Per quanto riguarda la concentrazione, si evidenzia sia l’eventuale eccesso di concentrazione in termini di peso sulla rispettiva asset class, sia quella sul singolo emittente.')
        run_10_2.font.name = 'Century Gothic'
        run_10_2.font.size = shared.Pt(10)
        paragraph_11 = self.document.add_paragraph(text='', style=None)
        paragraph_11.paragraph_format.alignment = 3
        paragraph_11.paragraph_format.space_after = shared.Pt(6)
        run_11 = paragraph_11.add_run('Con riferimento alla liquidabilità degli strumenti, si evidenzia l’eventuale basso grado di liquidabilità dell’investimento, dovuto ai tempi lunghi sia a costi elevati di smobilizzo. In questa sezione, inoltre, viene fornito un giudizio a livello di singolo prodotto/strumento:')
        run_11.font.name = 'Century Gothic'
        run_11.font.size = shared.Pt(10)
        paragraph_12 = self.document.add_paragraph(text='', style=None)
        paragraph_12.paragraph_format.alignment = 3
        paragraph_12.paragraph_format.space_after = shared.Pt(6)
        run_12_1 = paragraph_12.add_run('- “poco efficiente” ')
        run_12_1.bold = True
        run_12_1.font.name = 'Century Gothic'
        run_12_1.font.size = shared.Pt(10)
        run_12_2 = paragraph_12.add_run('(in termini di rapporto rendimento/rischio);')
        run_12_2.font.name = 'Century Gothic'
        run_12_2.font.size = shared.Pt(10)
        paragraph_13 = self.document.add_paragraph(text='', style=None)
        paragraph_13.paragraph_format.alignment = 3
        paragraph_13.paragraph_format.space_after = shared.Pt(6)
        run_13_1 = paragraph_13.add_run('- “neutro” ')
        run_13_1.bold = True
        run_13_1.font.name = 'Century Gothic'
        run_13_1.font.size = shared.Pt(10)
        run_13_2 = paragraph_13.add_run('(in termini di rapporto rendimento/rischio);')
        run_13_2.font.name = 'Century Gothic'
        run_13_2.font.size = shared.Pt(10)
        paragraph_14 = self.document.add_paragraph(text='', style=None)
        paragraph_14.paragraph_format.alignment = 3
        paragraph_14.paragraph_format.space_after = shared.Pt(6)
        run_14 = paragraph_14.add_run('- “default”;')
        run_14.bold = True
        run_14.font.name = 'Century Gothic'
        run_14.font.size = shared.Pt(10)
        paragraph_15 = self.document.add_paragraph(text='', style=None)
        paragraph_15.paragraph_format.alignment = 3
        paragraph_15.paragraph_format.space_after = shared.Pt(6)
        run_15 = paragraph_15.add_run('- “scaduto”.')
        run_15.bold = True
        run_15.font.name = 'Century Gothic'
        run_15.font.size = shared.Pt(10)
        # Note metodologiche 2
        self.document.add_section()
        paragraph_0 = self.document.add_paragraph(text='', style=None)
        paragraph_0.paragraph_format.space_before = shared.Pt(6)
        paragraph_0.paragraph_format.space_after = shared.Pt(0)
        run_0 = paragraph_0.add_run(text='')
        run_0.add_picture(self.path_img_default.joinpath('5_note_metodologiche.bmp').__str__(), width=shared.Cm(8.5))
        paragraph_1 = self.document.add_paragraph(text='\n', style=None)
        paragraph_1.paragraph_format.alignment = 0
        paragraph_1.paragraph_format.line_spacing_rule = 1
        paragraph_1.paragraph_format.space_after = shared.Pt(6)
        run_1_1 = paragraph_1.add_run('In particolare, per quanto riguarda le sezioni B) ed E), di seguito viene riportata una tabella riassuntiva con gli «alert» di concentrazione analizzati:\n\n')
        run_1_1.font.name = 'Century Gothic'
        run_1_1.font.size = shared.Pt(10)
        run_1_2 = paragraph_1.add_run()
        run_1_2.add_picture(self.path_img_default.joinpath('note_metodologiche.jpg').__str__(), width=shared.Cm(self.larghezza_pagina))
        paragraph_2 = self.document.add_paragraph(text='\n', style=None)
        paragraph_2.paragraph_format.alignment = 3
        paragraph_2.paragraph_format.line_spacing_rule = 1
        paragraph_2.paragraph_format.space_after = shared.Pt(6)
        run_2 = paragraph_2.add_run('Nel processo di analisi del portafoglio, quindi, vengono dapprima calcolate le esposizioni delle singole micro asset class rispetto alla macro asset class di riferimento. Per esempio, per quanto riguarda il compar-to azionario, se la componente europea pesa più del 60% dell’intera esposizione azionaria, scatterà il warning «!C», se pesa più del 70%, l’alert «!!C», altrimenti se rappresenta più dell’80%, il warning «!!!C».')
        run_2.font.name = 'Century Gothic'
        run_2.font.size = shared.Pt(10)
        paragraph_3 = self.document.add_paragraph(text='', style=None)
        paragraph_3.paragraph_format.alignment = 3
        paragraph_3.paragraph_format.line_spacing_rule = 1
        paragraph_3.paragraph_format.space_after = shared.Pt(6)
        run_3 = paragraph_3.add_run('Per quanto concerne l’esposizione dei singoli strumenti (azioni ed obbligazioni) sulla macro asset class di riferimento, ad esempio se un titolo azionario pesa più del 10%, del 20% o del 30% dell’intero comparto equity, allora scatteranno rispettivamente uno dei seguenti warning: «!C», «!!C», «!!!C». Per le obbligazioni tale controllo non viene svolto nel caso di bond governativi emessi dai paesi sviluppati.')
        run_3.font.name = 'Century Gothic'
        run_3.font.size = shared.Pt(10)
        paragraph_4 = self.document.add_paragraph(text='', style=None)
        paragraph_4.paragraph_format.alignment = 3
        paragraph_4.paragraph_format.line_spacing_rule = 1
        paragraph_4.paragraph_format.space_after = shared.Pt(6)
        run_4 = paragraph_4.add_run('Per le classi di strumenti più complessi (obbligazioni strutturate, Hedge Fund, Altri) si confronta, da un lato, l’esposizione rispetto all’intero portafoglio, dall’altro i livelli percentuali critici indicati nella terza tabella.')
        run_4.font.name = 'Century Gothic'
        run_4.font.size = shared.Pt(10)
        paragraph_5 = self.document.add_paragraph(text='', style=None)
        paragraph_5.paragraph_format.alignment = 3
        paragraph_5.paragraph_format.line_spacing_rule = 1
        paragraph_5.paragraph_format.space_after = shared.Pt(6)
        run_5 = paragraph_5.add_run('Per quanto riguarda l’esposizione valutaria, infine, scatterà un warning quando il peso delle valute diverse dall’Euro è maggiore del 40%, 50% o 60% dell’intero portafoglio.')
        run_5.font.name = 'Century Gothic'
        run_5.font.size = shared.Pt(10)
        # Note metodologiche 3
        if self.intermediario == 'azimut':
            self.document.add_section()
            paragraph_0 = self.document.add_paragraph(text='', style=None)
            paragraph_0.paragraph_format.space_before = shared.Pt(6)
            paragraph_0.paragraph_format.space_after = shared.Pt(0)
            run_0 = paragraph_0.add_run(text='')
            run_0.add_picture(self.path_img_default.joinpath('6_avvertenze.bmp').__str__(), width=shared.Cm(8.5))
            paragraph_1 = self.document.add_paragraph(text='\n', style=None)
            paragraph_1.paragraph_format.alignment = 3
            paragraph_1.paragraph_format.line_spacing_rule = 1
            paragraph_1.paragraph_format.space_after = shared.Pt(6)
            run_1 = paragraph_1.add_run('Questo documento è stato prodotto a solo scopo informativo. Di conseguenza non è fornita alcuna garanzia circa la completezza, l’accuratezza, l’affidabilità delle informazioni in esso contenute.')
            run_1.font.name = 'Century Gothic'
            run_1.font.size = shared.Pt(10)
            paragraph_2 = self.document.add_paragraph(text='', style=None)
            paragraph_2.paragraph_format.alignment = 3
            paragraph_2.paragraph_format.line_spacing_rule = 1
            paragraph_2.paragraph_format.space_after = shared.Pt(6)
            run_2 = paragraph_2.add_run('Di conseguenza nessuna garanzia, esplicita o implicita è fornita da parte o per conto della Società o di alcuno dei suoi membri, dirigenti, funzionari o impiegati o altre persone. Né la Società né alcuno dei suoi membri, dirigenti, funzionari o impiegati o altre persone che agiscano per conto della Società accetta alcuna responsabilità per qualsiasi perdita potesse derivare dall’uso di questa presentazione o dei suoi contenuti o altrimenti connesso con la presentazione e i suoi contenuti.')
            run_2.font.name = 'Century Gothic'
            run_2.font.size = shared.Pt(10)
            paragraph_3 = self.document.add_paragraph(text='', style=None)
            paragraph_3.paragraph_format.alignment = 3
            paragraph_3.paragraph_format.line_spacing_rule = 1
            paragraph_3.paragraph_format.space_after = shared.Pt(6)
            run_3 = paragraph_3.add_run('Le informazioni e opinioni contenute in questa presentazione sono aggiornate alla data indicata sulla presentazione e possono essere cambiate senza preavviso.')
            run_3.font.name = 'Century Gothic'
            run_3.font.size = shared.Pt(10)
            paragraph_4 = self.document.add_paragraph(text='', style=None)
            paragraph_4.paragraph_format.alignment = 3
            paragraph_4.paragraph_format.line_spacing_rule = 1
            paragraph_4.paragraph_format.space_after = shared.Pt(6)
            run_4 = paragraph_4.add_run('Questo documento non costituisce una sollecitazione o un’offerta e nessuna parte di esso può costituire la base o il riferimento per qualsivoglia contratto o impegno.')
            run_4.font.name = 'Century Gothic'
            run_4.font.size = shared.Pt(10)
            paragraph_5 = self.document.add_paragraph(text='', style=None)
            paragraph_5.paragraph_format.alignment = 3
            paragraph_5.paragraph_format.line_spacing_rule = 1
            paragraph_5.paragraph_format.space_after = shared.Pt(6)
            run_5 = paragraph_5.add_run('All’investimento descritto è associato il rischio di andamento dei tassi di interesse nominali e reali, dell’inflazione, dei cambi e dei mercati azionari e il rischio legato al possibile deterioramento del merito di credito degli emittenti.')
            run_5.font.name = 'Century Gothic'
            run_5.font.size = shared.Pt(10)
            paragraph_6 = self.document.add_paragraph(text='', style=None)
            paragraph_6.paragraph_format.alignment = 3
            paragraph_6.paragraph_format.line_spacing_rule = 1
            paragraph_6.paragraph_format.space_after = shared.Pt(6)
            run_6 = paragraph_6.add_run('Relativamente all’investimento in AZ Fund e Azimut Fondi si rimanda ai prospetti informativi dei relativi fondi che raccomandiamo di leggere prima della sottoscrizione.')
            run_6.font.name = 'Century Gothic'
            run_6.font.size = shared.Pt(10)
            paragraph_7 = self.document.add_paragraph(text='', style=None)
            paragraph_7.paragraph_format.alignment = 3
            paragraph_7.paragraph_format.line_spacing_rule = 1
            paragraph_7.paragraph_format.space_after = shared.Pt(6)
            run_7 = paragraph_7.add_run('L’investimento descritto non assicura il mantenimento del capitale, né offre garanzie di rendimento.')
            run_7.font.name = 'Century Gothic'
            run_7.font.size = shared.Pt(10)
        elif self.intermediario == 'copernico':
            self.document.add_section()
            paragraph_0 = self.document.add_paragraph(text='', style=None)
            paragraph_0.paragraph_format.space_before = shared.Pt(6)
            paragraph_0.paragraph_format.space_after = shared.Pt(0)
            run_0 = paragraph_0.add_run(text='')
            run_0.add_picture(self.path_img_default.joinpath('6_avvertenze.bmp').__str__(), width=shared.Cm(8.5))
            paragraph_1 = self.document.add_paragraph(text='\n', style=None)
            paragraph_1.paragraph_format.alignment = 3
            paragraph_1.paragraph_format.line_spacing_rule = 1
            paragraph_1.paragraph_format.space_after = shared.Pt(6)
            run_1 = paragraph_1.add_run('Questo documento non costituisce una sollecitazione o un’offerta né una raccomandazione ad effettuare investimenti di qualsiasi natura e nessuna parte di esso può costituire la base o il riferimento per qualsivoglia contratto o impegno.')
            run_1.font.name = 'Century Gothic'
            run_1.font.size = shared.Pt(10)
            paragraph_2 = self.document.add_paragraph(text='', style=None)
            paragraph_2.paragraph_format.alignment = 3
            paragraph_2.paragraph_format.line_spacing_rule = 1
            paragraph_2.paragraph_format.space_after = shared.Pt(6)
            run_2 = paragraph_2.add_run('La presente analisi è condotta tenendo conto del rischio di andamento dei tassi di interesse nominali e reali, dell’inflazione, dei cambi e dei mercati azionari e il rischio legato al possibile deterioramento del merito di credito degli emittenti.')
            run_2.font.name = 'Century Gothic'
            run_2.font.size = shared.Pt(10)
            paragraph_3 = self.document.add_paragraph(text='', style=None)
            paragraph_3.paragraph_format.alignment = 3
            paragraph_3.paragraph_format.line_spacing_rule = 1
            paragraph_3.paragraph_format.space_after = shared.Pt(6)
            run_3 = paragraph_3.add_run('La presente analisi non assicura il mantenimento del capitale, né offre garanzie di rendimento.')
            run_3.font.name = 'Century Gothic'
            run_3.font.size = shared.Pt(10)
        # Pagina di chiusura
        if self.intermediario == 'azimut':
            self.document.add_section()
            header = self.document.sections[-1].header
            header.is_linked_to_previous = False
            section = self.document.sections[-1]
            left_margin = 0.60
            right_margin = 0.60
            top_margin = 0.45
            bottom_margin = 0.45
            section.left_margin = shared.Cm(left_margin)
            section.right_margin = shared.Cm(right_margin)
            section.top_margin = shared.Cm(top_margin)
            section.bottom_margin = shared.Cm(bottom_margin)
            section.header_distance = shared.Cm(0)
            section.footer_distance = shared.Cm(0)
            paragraph_0 = self.document.add_paragraph(text='', style=None)
            paragraph_0.alignment = 1
            paragraph_0.add_run().add_picture(self.path_img_default.joinpath('pagina_di_chiusura.jpg').__str__(), height=shared.Cm(28.8), width=shared.Cm(19.8))
        elif self.intermediario == 'copernico':
            pass
        
    def pagine_numerate(self):
        """
        Numera le pagine a partire dall'indice
        """
        pagina = 1
        for num_section, section in enumerate(self.document.sections):
            if num_section != 0:
                self.document.sections[num_section].footer.is_linked_to_previous = False
                self.document.sections[num_section].footer.add_paragraph(str(pagina)).paragraph_format.alignment = 2
                self.document.sections[num_section].footer_distance = shared.Cm(1)
                pagina += 1

    def salva_file_portafoglio(self):
        """Salva il file excel."""
        try:
            self.wb.save(self.file_elaborato)
        except PermissionError:
            for proc in psutil.process_iter():
                if proc.name() == "EXCEL.EXE":
                    proc.kill()
            self.wb.save(self.file_elaborato)

    def salva_file_presentazione(self):
        """Salva il file della presentazione con nome."""
        try:
            self.document.save(self.file_presentazione)
        except PermissionError:
            print(f'\nChiudi il file {self.file_presentazione}')

if __name__ == "__main__":
    start = time.perf_counter()
    PTF = 'ptf_20.xlsx'
    INTERMEDIARIO = 'azimut'

    __ = Elaborazione(intermediario=INTERMEDIARIO)
    __.agglomerato()
    __.figure()
    __.mappatura_fondi()
    __.volatilità()
    __.sintesi()
    __.salva_file_portafoglio()
    __.autofit(sheet='agglomerato', columns=[1, 2, 3, 4, 5, 6, 7, 8, 9], min_width=[22, 50, 16, 22.5, 12, 10.5, 15, 10.5, 22.5], max_width=[26.5, None, None, None, None, None, None, None, None])
    # __.autofit(sheet='sintesi', columns=[1, 2, 3, 4, 5], min_width=[21, 34.5, None, 18.5, None], max_width=[23.5, None, None, 29.5, None])

    ___ = Presentazione(intermediario=INTERMEDIARIO, tipo_sap='completo', page_height = 29.7, page_width = 21, top_margin = 2.5, bottom_margin = 2.5, left_margin = 1.5, right_margin = 1.5)
    ___.copertina()
    ___.indice()
    # ___.portafoglio_attuale(method='label_on_top')
    # ___.commento()
    # ___.analisi_di_portafoglio()
    # ___.analisi_strumenti()
    # ___.analisi_del_rischio()
    ___.note_metodologiche()

    ___.pagine_numerate()
    ___.salva_file_portafoglio()
    ___.salva_file_presentazione()
    
    end = time.perf_counter()
    print("Elapsed time: ", round(end - start, 2), 'seconds')