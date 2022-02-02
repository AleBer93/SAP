from docx import Document, shared
from SAP import Portfolio
import dataframe_image as dfi
import pandas as pd
ptf = Portfolio(path=r'C:\Users\Administrator\Desktop\Sbwkrq\SAP', file_portafoglio='ptf_20.xlsx', intermediario='azimut')
document = Document()
dict_strumenti = ptf.peso_strumenti()
print(dict_strumenti)
df = pd.DataFrame.from_dict(dict_strumenti, orient='index', columns=['peso_strumento'])
dfi.export(df, 'lol.png')
section = document.sections[0]
paragraph = document.add_paragraph()
paragraph.add_run().add_picture('lol.png')
document.save('lol.docx')